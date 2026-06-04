#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import urllib.request
import urllib.parse
import re
import os
import hashlib
import time
import sqlite3
import threading
import subprocess
import platform
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError as FuturesTimeoutError
from pathlib import Path
from collections import defaultdict
from flask import Flask, render_template, request, jsonify, Response, stream_with_context, send_file
import io
import logging
import requests   # ← dodane dla lepszej obsługi rozłączeń klienta w streamach SSE
from qdrant_client import QdrantClient
from qdrant_client.models import PointStruct
from prompts import SEARCH_MODES, MODEL_REGISTRY, VERIFY_SYSTEM_PROMPT, VERIFY_PROMPT_TEMPLATE, DETECTIVE_PROMPT_TEMPLATE, DEFAULT_PROMPT_TEMPLATE, CHAT_CONTEXT_BLOCK_TEMPLATE
from sql_safety import _is_sql_safe, _extract_sql_table_refs, _validate_sql_table_refs, DANGEROUS_SQL_KEYWORDS
from models_fleet import _load_custom_endpoint_stats, _save_custom_endpoint_stats, _get_or_init_endpoint_stats, _load_model_live_stats, _save_model_live_stats, _model_score, _custom_endpoint_score, _set_model_live, _set_MODEL_REGISTRY
import llm_client
from llm_client import (
    call_llm,
    stream_llm_tokens,
    get_llm_provider,
    _is_rate_limit_error,
    _get_retry_after,
    _redact_api_keys,
)

# Wczytaj .env jeśli istnieje
_env_path = Path(__file__).parent / ".env"
if _env_path.exists():
    for _line in _env_path.read_text().splitlines():
        _line = _line.strip()
        if _line and not _line.startswith("#") and "=" in _line:
            _k, _v = _line.split("=", 1)
            os.environ.setdefault(_k.strip(), _v.strip())

# Bezpieczne importy opcjonalne
try: import docx
except ImportError: docx = None
try: import pdfplumber
except ImportError: pdfplumber = None
try: import openpyxl
except ImportError: openpyxl = None
try: import pytesseract
except ImportError: pytesseract = None
try: from pdf2image import convert_from_path
except ImportError: convert_from_path = None
try:
    import pymssql
    PYMSSQL_AVAILABLE = True
except ImportError:
    pymssql = None
    PYMSSQL_AVAILABLE = False


def _format_table_key(schema_name: str, table_name: str) -> str:
    if schema_name:
        return f"{schema_name}.{table_name}"
    return table_name


def _sql_dialect(cfg: dict) -> str:
    t = (cfg.get("type") or cfg.get("db_type") or "mssql").strip().lower()
    return "sqlite" if t == "sqlite" else "mssql"


def _sql_conn_configured(cfg: dict) -> bool:
    if not cfg:
        return False
    if _sql_dialect(cfg) == "sqlite":
        raw = cfg.get("db_path") or cfg.get("path") or cfg.get("database") or ""
        return bool(_normalize_sqlite_path_input(raw))
    return bool(cfg.get("server") and cfg.get("database"))


def _normalize_sqlite_path_input(raw: str) -> str:
    """Normalizuje ścieżkę SQLite (WSL: G:\\foo → /mnt/g/foo)."""
    raw = (raw or "").strip().strip('"').strip("'")
    if not raw:
        return raw
    win = re.match(r"^([A-Za-z]):[\\/](.*)$", raw)
    if win:
        drive = win.group(1).lower()
        rest = win.group(2).replace("\\", "/")
        return f"/mnt/{drive}/{rest}"
    return raw


def _resolve_sqlite_path(cfg: dict) -> Path:
    raw = (
        cfg.get("db_path") or cfg.get("path") or cfg.get("database") or ""
    ).strip()
    raw = _normalize_sqlite_path_input(raw)
    if not raw:
        raise ValueError(
            "Podaj ścieżkę do pliku SQLite (.db) — np. /mnt/g/dane/eksport.db"
        )
    p = Path(raw).expanduser()
    if not p.is_file():
        raise ValueError(f"Plik bazy SQLite nie istnieje: {raw}")
    if not _path_is_allowed(p):
        raise ValueError("Ścieżka poza dozwolonymi katalogami (SEARCH_ROOTS)")
    if p.suffix.lower() not in (".db", ".sqlite", ".sqlite3"):
        raise ValueError("Dozwolone rozszerzenia pliku SQLite: .db, .sqlite, .sqlite3")
    return p.resolve()


def _get_mssql_conn(cfg: dict):
    """Tworzy połączenie z MS SQL Server przez pymssql."""
    if not PYMSSQL_AVAILABLE:
        raise RuntimeError("pymssql nie jest zainstalowany")
    server = cfg.get("server", "127.0.0.1")
    port   = str(cfg.get("port", 1433))
    if ":" not in server and "\\" not in server:
        server = f"{server}:{port}"
    db = cfg.get("database", "")
    return pymssql.connect(
        server   = server,
        user     = cfg.get("user", ""),
        password = cfg.get("password", ""),
        database = db if db else None,
        timeout  = 15,
        charset  = "UTF-8"
    )


def _get_sql_conn(cfg: dict, readonly: bool = True):
    """Połączenie z MS SQL Server lub SQLite (domyślnie read-only dla SQLite)."""
    if _sql_dialect(cfg) == "sqlite":
        path = _resolve_sqlite_path(cfg)
        mode = "ro" if readonly else "rw"
        conn = sqlite3.connect(f"file:{path}?mode={mode}", uri=True, timeout=15)
        conn.row_factory = sqlite3.Row
        return conn
    return _get_mssql_conn(cfg)


def _require_sql_backend(cfg: dict) -> None:
    if _sql_dialect(cfg) == "mssql" and not PYMSSQL_AVAILABLE:
        raise RuntimeError("Moduł pymssql nie jest zainstalowany")


def _sql_quote_ident(name: str, dialect: str) -> str:
    if dialect == "mssql":
        return f"[{name}]"
    return f'"{name}"'


def _sql_rows_to_dicts(cur, limit: int = 500) -> list[dict]:
    rows = cur.fetchmany(limit)
    out: list[dict] = []
    for row in rows:
        if isinstance(row, dict):
            item = row
        elif hasattr(row, "keys"):
            item = {k: row[k] for k in row.keys()}
        else:
            cols = [d[0] for d in (cur.description or [])]
            item = {cols[i]: row[i] for i in range(len(cols))}
        out.append({k: (str(v) if v is not None else "") for k, v in item.items()})
    return out


def _sql_llm_dialect_label(dialect: str) -> str:
    return "SQLite" if dialect == "sqlite" else "T-SQL (MS SQL Server)"


def _fetch_sql_known_tables(cfg: dict) -> set[str]:
    """Zwraca znane nazwy tabel do walidacji SQL z LLM."""
    known: set[str] = set()
    dialect = _sql_dialect(cfg)
    conn = _get_sql_conn(cfg)
    try:
        cur = conn.cursor()
        if dialect == "sqlite":
            cur.execute("""
                SELECT name FROM sqlite_master
                WHERE type IN ('table','view') AND name NOT LIKE 'sqlite_%'
            """)
            for (table_name,) in cur.fetchall():
                known.add(table_name.lower())
        else:
            cur.execute("""
                SELECT TABLE_SCHEMA, TABLE_NAME
                FROM INFORMATION_SCHEMA.TABLES
                WHERE TABLE_TYPE IN ('BASE TABLE','VIEW')
            """)
            for schema_name, table_name in cur.fetchall():
                key = _format_table_key(schema_name, table_name)
                known.add(key.lower())
                known.add(table_name.lower())
    finally:
        conn.close()
    return known


SQL_ROW_LIMIT = 100
SQL_CORRECTION_MAX_RETRIES = 2
SQL_SCHEMA_MAX_TABLES = 50
SQL_SCHEMA_MAX_SAMPLES = 5
SQL_SAMPLE_VALUE_TYPES = frozenset({
    "date", "datetime", "datetime2", "smalldatetime", "timestamp", "time",
    "datetimeoffset", "text", "ntext", "varchar", "nvarchar", "char", "nchar",
})


def _sql_limit_instruction(dialect: str) -> str:
    if dialect == "mssql":
        return f"ZAWSZE używaj TOP {SQL_ROW_LIMIT} (np. SELECT TOP {SQL_ROW_LIMIT} ...)."
    return f"ZAWSZE dodaj LIMIT {SQL_ROW_LIMIT} na końcu zapytania."


def _ensure_sql_row_limit(sql_query: str, dialect: str, limit: int = SQL_ROW_LIMIT) -> str:
    sql = (sql_query or "").strip()
    if not sql:
        return sql
    upper = sql.upper()
    if dialect == "mssql":
        if re.search(r"\bTOP\s+\d+", upper):
            return sql
        idx = upper.rfind("SELECT")
        if idx < 0:
            return sql
        prefix, rest = sql[:idx], sql[idx:]
        rest = re.sub(
            r"(?i)^SELECT\s+(DISTINCT\s+)?",
            lambda m: f"SELECT {m.group(1) or ''}TOP {limit} ",
            rest,
            count=1,
        )
        return prefix + rest
    if re.search(r"\bLIMIT\s+\d+", upper):
        return sql
    return f"{sql.rstrip(';').strip()} LIMIT {limit}"


def _sql_fetch_column_sample(
    conn, table: str, col: str, col_type: str, dialect: str,
) -> str | None:
    ct = (col_type or "").lower().split("(")[0].strip()
    should_sample = (
        ct in SQL_SAMPLE_VALUE_TYPES
        or "date" in ct
        or "time" in ct
    )
    if not should_sample:
        return None
    qtable = _sql_quote_ident(table, dialect)
    qcol = _sql_quote_ident(col, dialect)
    cur = conn.cursor()
    try:
        if dialect == "mssql":
            cur.execute(
                f"SELECT TOP 1 CAST({qcol} AS NVARCHAR(200)) AS v "
                f"FROM {qtable} WHERE {qcol} IS NOT NULL"
            )
        else:
            cur.execute(
                f"SELECT CAST({qcol} AS TEXT) AS v "
                f"FROM {qtable} WHERE {qcol} IS NOT NULL LIMIT 1"
            )
        row = cur.fetchone()
        if not row or row[0] is None:
            return None
        val = str(row[0]).strip()
        return val[:80] if val else None
    except Exception:
        return None
    finally:
        cur.close()


def _sql_table_column_samples(
    conn, table: str, col_parts: list[str], dialect: str,
    max_samples: int = SQL_SCHEMA_MAX_SAMPLES,
) -> str:
    """Próbki wartości tylko dla kolumn dat/tekst (max N na tabelę — bez skanowania wszystkich kolumn)."""
    samples: list[str] = []
    for part in col_parts:
        if len(samples) >= max_samples:
            break
        m = re.match(r"([\w]+)\s+\(([^)]+)\)", part)
        if not m:
            continue
        col, ctype = m.group(1), m.group(2)
        val = _sql_fetch_column_sample(conn, table, col, ctype, dialect)
        if val:
            samples.append(f'{col} np. "{val}"')
    return "; ".join(samples)


def _build_sql_schema_entries(
    cfg: dict,
    max_tables: int = SQL_SCHEMA_MAX_TABLES,
    include_samples: bool = True,
) -> list[dict]:
    """Pojedyncze wpisy schematu (tabela → tekst) z opcjonalnymi próbkami wartości."""
    dialect = _sql_dialect(cfg)
    conn = _get_sql_conn(cfg)
    entries: list[dict] = []
    try:
        cur = conn.cursor()
        if dialect == "sqlite":
            cur.execute("""
                SELECT name, sql FROM sqlite_master
                WHERE type='table' AND name NOT LIKE 'sqlite_%'
                ORDER BY name
                LIMIT ?
            """, (max_tables,))
            table_rows = cur.fetchall()
            if not table_rows:
                cur.execute("""
                    SELECT name, type FROM sqlite_master
                    WHERE type IN ('table','view') AND name NOT LIKE 'sqlite_%'
                    ORDER BY name LIMIT ?
                """, (max_tables,))
                for name, ttype in cur.fetchall():
                    entries.append({
                        "table": name,
                        "text": f"Tabela: {name} ({ttype})",
                    })
            else:
                for name, ddl in table_rows:
                    cur.execute(f'PRAGMA table_info("{name}")')
                    col_rows = cur.fetchall()
                    cols = [f"{r[1]} ({r[2]})" for r in col_rows]
                    col_txt = ", ".join(cols[:25]) + (" …" if len(cols) > 25 else "")
                    text = f"Tabela: {name}\nKolumny: {col_txt}" if col_txt else f"Tabela: {name}"
                    if include_samples and col_txt:
                        samp = _sql_table_column_samples(conn, name, cols[:25], dialect)
                        if samp:
                            text += f"\nPrzykładowe wartości: {samp}"
                    elif ddl and not col_txt:
                        text = f"Tabela: {name}\nDDL: {ddl[:400]}"
                    entries.append({"table": name, "text": text})
        else:
            row_limit = max(100, min(max_tables * 40, 4000))
            cur.execute(f"""
                SELECT TOP {row_limit} c.TABLE_SCHEMA, c.TABLE_NAME, c.COLUMN_NAME, c.DATA_TYPE
                FROM INFORMATION_SCHEMA.COLUMNS c
                INNER JOIN INFORMATION_SCHEMA.TABLES t
                  ON c.TABLE_SCHEMA = t.TABLE_SCHEMA AND c.TABLE_NAME = t.TABLE_NAME
                WHERE t.TABLE_TYPE IN ('BASE TABLE','VIEW')
                ORDER BY c.TABLE_SCHEMA, c.TABLE_NAME, c.ORDINAL_POSITION
            """)
            grouped: dict[str, list[str]] = defaultdict(list)
            tables_seen: set[str] = set()
            for schema_name, table_name, col_name, data_type in cur.fetchall():
                key = _format_table_key(schema_name, table_name)
                if key not in tables_seen and len(tables_seen) >= max_tables:
                    continue
                tables_seen.add(key)
                grouped[key].append(f"{col_name} ({data_type})")

            if not grouped:
                cur.execute("""
                    SELECT TABLE_SCHEMA, TABLE_NAME, TABLE_TYPE
                    FROM INFORMATION_SCHEMA.TABLES
                    WHERE TABLE_TYPE IN ('BASE TABLE','VIEW')
                    ORDER BY TABLE_NAME
                """)
                for schema_name, table_name, table_type in cur.fetchall()[:max_tables]:
                    key = _format_table_key(schema_name, table_name)
                    entries.append({
                        "table": key,
                        "text": f"Tabela: {key} ({table_type})",
                    })
            else:
                for key in sorted(grouped.keys()):
                    cols = grouped[key]
                    col_txt = ", ".join(cols[:25]) + (" …" if len(cols) > 25 else "")
                    text = f"Tabela: {key}\nKolumny: {col_txt}"
                    if include_samples:
                        table_only = key.split(".")[-1]
                        samp = _sql_table_column_samples(conn, table_only, cols[:25], dialect)
                        if samp:
                            text += f"\nPrzykładowe wartości: {samp}"
                    entries.append({"table": key, "text": text})
    finally:
        conn.close()
    return entries


def _format_sql_schema_text(entries: list[dict]) -> str:
    header = (
        "UWAGA: Używaj WYŁĄCZNIE poniższych tabel i kolumn. "
        "Nie wymyślaj nazw, jeśli ich nie ma na liście.\n"
    )
    return header + "\n".join(e["text"] for e in entries)


def _build_auto_sql_schema(cfg: dict, max_tables: int = SQL_SCHEMA_MAX_TABLES) -> str:
    entries = _build_sql_schema_entries(cfg, max_tables=max_tables, include_samples=True)
    return _format_sql_schema_text(entries)


def _resolve_sql_schema(cfg: dict, user_schema: str) -> tuple[str, set[str]]:
    """Łączy schemat z UI z automatycznym pobraniem z bazy (INFORMATION_SCHEMA + próbki)."""
    known = _fetch_sql_known_tables(cfg)
    user_schema = (user_schema or "").strip()
    if user_schema and "Kolumny:" in user_schema:
        return user_schema, known
    if user_schema:
        auto = _build_auto_sql_schema(cfg)
        merged = f"{user_schema}\n\n--- Pełny schemat z bazy ---\n{auto}"
        return merged, known
    return _build_auto_sql_schema(cfg), known


def _is_numeric_sql_value(val) -> bool:
    s = str(val or "").strip().replace(",", ".").replace(" ", "")
    if not s:
        return False
    try:
        float(s)
        return True
    except ValueError:
        return False


def _build_sql_chart_config(columns: list[str], rows: list[dict]) -> dict | None:
    """Heurystyka: wykres słupkowy/kołowy/liniowy dla wyników agregacji."""
    if not rows or len(columns) < 2 or len(rows) > 50:
        return None

    label_col = None
    value_col = None
    for col in columns:
        samples = [rows[i].get(col, "") for i in range(min(5, len(rows)))]
        non_empty = [s for s in samples if str(s).strip()]
        if non_empty and all(_is_numeric_sql_value(v) for v in non_empty):
            if value_col is None:
                value_col = col
        elif label_col is None:
            label_col = col

    if (not label_col or not value_col) and len(columns) >= 2:
        if all(_is_numeric_sql_value(rows[i].get(columns[1], "")) for i in range(min(5, len(rows)))):
            label_col, value_col = columns[0], columns[1]

    if not label_col or not value_col or label_col == value_col:
        return None

    labels = [str(r.get(label_col, ""))[:40] for r in rows]
    data: list[float] = []
    for r in rows:
        try:
            data.append(float(str(r.get(value_col, "")).replace(",", ".").replace(" ", "")))
        except ValueError:
            data.append(0.0)

    chart_type = "bar"
    if 2 <= len(labels) <= 8:
        chart_type = "pie"
    label_blob = " ".join(labels[:3]).lower()
    if any(x in label_blob for x in ("202", "201", "-01", "-02", "sty", "lut", "mar", "kwi")):
        chart_type = "line"

    return {
        "type": chart_type,
        "title": f"{value_col} wg {label_col}",
        "labels": labels,
        "datasets": [{"label": value_col, "data": data}],
    }


def _run_text_to_sql_pipeline(
    cfg: dict,
    question: str,
    schema_str: str = "",
    fetch_limit: int = 500,
    with_chart: bool = True,
) -> dict:
    """Text-to-SQL z auto-korektą, LIMIT/TOP, schematem z próbkami i opcjonalnym wykresem."""
    dialect = _sql_dialect(cfg)
    dialect_label = _sql_llm_dialect_label(dialect)
    effective_schema, known_tables = _resolve_sql_schema(cfg, schema_str)

    limit_hint = _sql_limit_instruction(dialect)
    system_sql = (
        f"Jesteś ekspertem SQL ({dialect_label}). "
        "Generujesz wyłącznie zapytania SELECT (nigdy DELETE/UPDATE/DROP). "
        f"{limit_hint} "
        "Używaj formatów dat widocznych w przykładowych wartościach schematu. "
        "Odpowiadasz WYŁĄCZNIE samym SQL — bez wyjaśnień, bez markdown."
    )

    sql_query = ""
    last_error = ""
    corrections = 0

    for attempt in range(SQL_CORRECTION_MAX_RETRIES + 1):
        if attempt == 0:
            prompt_sql = (
                f"SCHEMAT BAZY:\n{effective_schema}\n\n"
                f"PYTANIE UŻYTKOWNIKA: {question}\n\n"
                f"Wygeneruj zapytanie SQL ({dialect_label}):"
            )
        else:
            corrections = attempt
            prompt_sql = (
                f"SCHEMAT BAZY:\n{effective_schema}\n\n"
                f"PYTANIE: {question}\n\n"
                f"Twoje poprzednie zapytanie:\n{sql_query}\n\n"
                f"Serwer bazy zwrócił błąd:\n{last_error}\n\n"
                "Popraw zapytanie SQL. Zwróć WYŁĄCZNIE poprawiony kod SQL:"
            )

        sql_response = call_llm(
            prompt_sql, system_sql, stream=False, max_tokens=2000, temperature=0.1,
        )
        sql_query = _clean_llm_sql(_llm_response_text(sql_response))
        sql_query = _ensure_sql_row_limit(sql_query, dialect)

        first_word = sql_query.split()[0].upper() if sql_query.split() else ""
        if first_word not in ("SELECT", "WITH"):
            last_error = f"Niedozwolone polecenie: {first_word}"
            if attempt >= SQL_CORRECTION_MAX_RETRIES:
                break
            continue

        is_safe, safe_err = _is_sql_safe(sql_query, ("SELECT", "WITH"))
        if not is_safe:
            last_error = safe_err or "Niedozwolone zapytanie SQL"
            if attempt >= SQL_CORRECTION_MAX_RETRIES:
                break
            continue

        ok_tables, table_err = _validate_sql_table_refs(sql_query, known_tables)
        if not ok_tables:
            last_error = table_err or "Nieznane tabele w SQL"
            if attempt >= SQL_CORRECTION_MAX_RETRIES:
                break
            continue

        try:
            conn = _get_sql_conn(cfg, readonly=True)
            cur = conn.cursor(as_dict=True) if dialect == "mssql" else conn.cursor()
            cur.execute(sql_query)
            result_rows = _sql_rows_to_dicts(cur, fetch_limit)
            cols = list(result_rows[0].keys()) if result_rows else []
            conn.close()

            out: dict = {
                "success": True,
                "sql": sql_query,
                "columns": cols,
                "rows": result_rows,
                "total": len(result_rows),
                "truncated": len(result_rows) == fetch_limit,
                "sql_corrections": corrections,
            }
            if with_chart:
                chart = _build_sql_chart_config(cols, result_rows)
                if chart:
                    out["chart"] = chart
            return out
        except Exception as e:
            last_error = str(e)
            if attempt >= SQL_CORRECTION_MAX_RETRIES:
                break

    return {
        "success": False,
        "error": last_error or "Nie udało się wygenerować poprawnego SQL",
        "sql": sql_query,
        "sql_corrections": corrections,
    }


# ============================================================
# KONIEC PRZYWRÓCONYCH FUNKCJI BEZPIECZEŃSTWA SQL
# ============================================================

SQL_CONFIG_PATH = Path(__file__).parent / ".sql_config.json"


def _load_sql_config() -> dict:
    """Pełna konfiguracja SQL z pliku (w tym hasło) — tylko backend."""
    if SQL_CONFIG_PATH.exists():
        try:
            return json.loads(SQL_CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def _save_sql_config(cfg: dict) -> None:
    SQL_CONFIG_PATH.write_text(json.dumps(cfg, indent=2), encoding="utf-8")


def _llm_response_text(result) -> str:
    if isinstance(result, dict):
        return (result.get("response") or "").strip()
    return str(result).strip()


def _clean_llm_sql(raw: str) -> str:
    sql = re.sub(r"^```\w*\n?", "", raw or "", flags=re.MULTILINE)
    sql = re.sub(r"```$", "", sql.strip()).strip()
    return sql


def _validate_sql_table_name(name: str) -> bool:
    return bool(name and re.match(r"^[\w]+$", name))


app = Flask(__name__)

# === Podstawowe logowanie ===
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)
logger = logging.getLogger("ai_analiza")

# === Walidacja wymaganych zmiennych środowiskowych ===
required_env = ["QDRANT_URL", "QDRANT_KEY"]
missing = [key for key in required_env if not os.environ.get(key)]
if missing:
    logger.critical(f"Brak wymaganych zmiennych środowiskowych: {missing}")
    logger.critical("Upewnij się, że plik .env istnieje i zawiera poprawne wartości.")
    raise RuntimeError(f"Brakujące zmienne środowiskowe: {missing}")

QDRANT_URL        = os.environ["QDRANT_URL"]
QDRANT_KEY        = os.environ["QDRANT_KEY"]
QDRANT_LOCAL_URL  = os.environ.get("QDRANT_LOCAL_URL", "http://127.0.0.1:6333")
QDRANT_LOCAL_KEY  = os.environ.get("QDRANT_LOCAL_KEY", "dev-local-key")
QDRANT_CLOUD_URL  = os.environ.get("QDRANT_CLOUD_URL", "")
QDRANT_CLOUD_KEY  = os.environ.get("QDRANT_CLOUD_KEY", "")
ACTIVE_COLLECTION = os.environ.get("ACTIVE_COLLECTION", "dokumenty")
OLLAMA_URL        = os.environ.get("OLLAMA_URL", "http://127.0.0.1:11434")
LLM_MODEL         = os.environ.get("LLM_MODEL", "llama3")
EMBED_MODEL       = os.environ.get("EMBED_MODEL", "nomic-embed-text")
OCR_LANG          = os.environ.get("OCR_LANG", "pol+eng")

# OpenRouter
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
OPENROUTER_MODEL   = os.environ.get("OPENROUTER_MODEL", "meta-llama/llama-3.3-70b-instruct:free")
OPENROUTER_MODEL_VERIFY = os.environ.get("OPENROUTER_MODEL_VERIFY", "google/gemini-2.0-flash-exp:free")
OPENROUTER_FALLBACK_TO_OLLAMA = os.environ.get("OPENROUTER_FALLBACK_TO_OLLAMA", "true").lower() in ("1", "true", "yes")
try:
    OPENROUTER_MAX_RETRIES = max(1, int(os.environ.get("OPENROUTER_MAX_RETRIES", "3")))
except ValueError:
    logger.warning("Nieprawidłowa wartość OPENROUTER_MAX_RETRIES w .env, używam domyślnej: 3")
    OPENROUTER_MAX_RETRIES = 3

DEFAULT_LLM_PROVIDER = os.environ.get("LLM_PROVIDER", "openrouter").lower()   # openrouter (domyślnie dla LLM) | ollama (tylko dla importu/embeddings)
GROQ_MODEL = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile").strip()

# Google Gemini (bezpośrednie API — audyt dużych plików / logów)
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()
# gemini-1.5-* wycofany w API (404) — domyślnie 2.0-flash (tańszy niż 2.5, ~1M ctx)
GEMINI_MODEL_DEFAULT = "gemini-2.0-flash"
GEMINI_DEPRECATED_MODELS = frozenset({
    "gemini-1.5-flash", "gemini-1.5-flash-8b", "gemini-1.5-pro",
    "gemini-1.5-flash-latest", "gemini-1.5-pro-latest",
})
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", GEMINI_MODEL_DEFAULT).strip() or GEMINI_MODEL_DEFAULT
try:
    GEMINI_AUDIT_MAX_BYTES = max(1024 * 1024, int(os.environ.get("GEMINI_AUDIT_MAX_BYTES", str(12 * 1024 * 1024))))
except ValueError:
    GEMINI_AUDIT_MAX_BYTES = 12 * 1024 * 1024

GEMINI_AUDIT_EXTENSIONS = {".log", ".txt", ".json", ".out", ".err", ".md", ".csv", ".xml"}
GEMINI_API_BASE = "https://generativelanguage.googleapis.com/v1beta"
try:
    GEMINI_MAX_RETRIES = max(1, int(os.environ.get("GEMINI_MAX_RETRIES", "5")))
except ValueError:
    GEMINI_MAX_RETRIES = 5
try:
    GEMINI_RETRY_INITIAL_DELAY = max(0.5, float(os.environ.get("GEMINI_RETRY_INITIAL_DELAY", "2.5")))
except ValueError:
    GEMINI_RETRY_INITIAL_DELAY = 2.5
GEMINI_RETRY_MAX_WAIT_SEC = 60.0

SEARCH_ROOTS      = [p.strip() for p in os.environ.get("SEARCH_ROOTS", "").split(':') if p.strip()]

# ---- Swarm Reports (persystencja wyników roju) ----
_SWARM_REPORTS_DIR = Path(__file__).parent / ".swarm_reports"

def _ensure_swarm_reports_dir() -> None:
    _SWARM_REPORTS_DIR.mkdir(exist_ok=True)

# ---- Provider Pool (wiele kluczy / adresów serwerów) ----
_PROVIDERS_FILE = Path(__file__).parent / ".providers.json"
_provider_pool_lock = threading.Lock()

def _load_provider_pool() -> dict:
    """Wczytuje pulę dostawców z pliku JSON."""
    try:
        if _PROVIDERS_FILE.exists():
            return json.loads(_PROVIDERS_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {"openrouter_keys": [], "ollama_urls": [], "custom_endpoints": []}

def _save_provider_pool(pool: dict) -> None:
    with _provider_pool_lock:
        _PROVIDERS_FILE.write_text(json.dumps(pool, indent=2, ensure_ascii=False), encoding="utf-8")

# Wskaźnik rotacji round-robin (klucze OpenRouter)
_provider_rr_idx: int = 0

def _pick_openrouter_key() -> str:
    """Zwraca kolejny aktywny klucz OpenRouter z puli (round-robin).
    Jeśli pula pusta — używa głównego OPENROUTER_API_KEY z .env."""
    global _provider_rr_idx
    pool = _load_provider_pool()
    active = [e for e in pool.get("openrouter_keys", []) if e.get("active") and e.get("key")]
    if not active:
        return OPENROUTER_API_KEY
    with _provider_pool_lock:
        idx = _provider_rr_idx % len(active)
        _provider_rr_idx = (idx + 1) % len(active)
    return active[idx]["key"]

# ---- Custom Endpoint Health & Stats ----


APP_API_KEY_ENV = os.environ.get("APP_API_KEY", "").strip()
APP_API_KEY = APP_API_KEY_ENV
APP_HOST    = os.environ.get("APP_HOST", "127.0.0.1")
# Ustaw TRUST_PROXY=true tylko jeśli aplikacja stoi za zaufanym reverse proxy (nginx/traefik)
# który nadpisuje X-Forwarded-For — NIGDY nie włączaj gdy app jest dostępna z zewnątrz bez proxy
TRUST_PROXY = os.environ.get("TRUST_PROXY", "false").lower() in ("1", "true", "yes")

# Inicjalizuj llm_client z bieżącymi zmiennymi konfiguracyjnymi
llm_client._set_load_provider_pool(_load_provider_pool)
llm_client._sync_llm_client_config(
    GEMINI_API_KEY=GEMINI_API_KEY,
    GEMINI_MODEL=GEMINI_MODEL,
    OPENROUTER_API_KEY=OPENROUTER_API_KEY,
    OPENROUTER_MODEL=OPENROUTER_MODEL,
    OPENROUTER_MODEL_VERIFY=OPENROUTER_MODEL_VERIFY,
    LLM_MODEL=LLM_MODEL,
    GROQ_MODEL=GROQ_MODEL,
    DEFAULT_LLM_PROVIDER=DEFAULT_LLM_PROVIDER,
)

_HEAVY_TASK_HISTORY_LIMIT = 20
_heavy_task_mutex = threading.Lock()
_heavy_task_state_lock = threading.Lock()
_heavy_task_current: dict | None = None
_heavy_task_history: list[dict] = []


def _new_task_id(kind: str) -> str:
    raw = f"{kind}:{time.time_ns()}:{threading.get_ident()}"
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()[:12]


def _task_snapshot(task: dict | None) -> dict | None:
    if not task:
        return None
    snapshot = dict(task)
    meta = snapshot.get("meta")
    if isinstance(meta, dict):
        snapshot["meta"] = dict(meta)
    return snapshot


def _tasks_snapshot() -> tuple[dict | None, list[dict]]:
    with _heavy_task_state_lock:
        current = _task_snapshot(_heavy_task_current)
        history = [_task_snapshot(item) for item in _heavy_task_history]
    return current, history


def _start_heavy_task(kind: str, label: str, meta: dict | None = None) -> tuple[bool, dict | None]:
    if not _heavy_task_mutex.acquire(blocking=False):
        current, _ = _tasks_snapshot()
        return False, current

    now = time.time()
    task = {
        "id": _new_task_id(kind),
        "kind": kind,
        "label": label,
        "status": "running",
        "started_at": now,
        "updated_at": now,
        "finished_at": None,
        "progress_pct": 0,
        "done": 0,
        "total": None,
        "current_item": None,
        "meta": meta or {},
        "error": None,
    }
    with _heavy_task_state_lock:
        global _heavy_task_current
        _heavy_task_current = task
    return True, _task_snapshot(task)


def _update_heavy_task(task_id: str, **updates) -> None:
    if not task_id:
        return
    safe_updates = {k: v for k, v in updates.items() if k not in {"id", "started_at"}}
    with _heavy_task_state_lock:
        if not _heavy_task_current or _heavy_task_current.get("id") != task_id:
            return
        _heavy_task_current.update(safe_updates)
        _heavy_task_current["updated_at"] = time.time()


def _finish_heavy_task(task_id: str, status: str = "done", error: str | None = None) -> None:
    if not task_id:
        return
    should_release = False
    with _heavy_task_state_lock:
        global _heavy_task_current
        if not _heavy_task_current or _heavy_task_current.get("id") != task_id:
            return
        now = time.time()
        _heavy_task_current["status"] = status
        _heavy_task_current["updated_at"] = now
        _heavy_task_current["finished_at"] = now
        _heavy_task_current["error"] = error
        _heavy_task_history.insert(0, _task_snapshot(_heavy_task_current))
        del _heavy_task_history[_HEAVY_TASK_HISTORY_LIMIT:]
        _heavy_task_current = None
        should_release = True
    if should_release:
        _heavy_task_mutex.release()


def _heavy_task_busy_payload(task: dict | None) -> dict:
    return {
        "msg": "Inna ciężka operacja jest już uruchomiona. Poczekaj na zakończenie zadania.",
        "error": "heavy_task_already_running",
        "task": task,
    }

# Wersja aplikacji — automatycznie odczytywana z git tagów w trybie deweloperskim
# (git describe --tags --dirty). W releasach produkcyjnych wraca do stałej.
def _get_app_version() -> str:
    try:
        out = subprocess.check_output(
            ["git", "describe", "--tags", "--always", "--dirty", "--abbrev=0"],
            cwd=Path(__file__).parent,
            stderr=subprocess.DEVNULL,
            timeout=1.0,
        )
        return out.decode().strip()
    except Exception:
        return "2026.07"

APP_VERSION = _get_app_version()

_ACTIVE_COLLECTION_FILE = Path(__file__).parent / ".active_collection"


# ============================================================
# UPDATE / SELF-UPDATE HELPERS (for local development)
# ============================================================

def _get_remote_latest_tag() -> str | None:
    """Pobiera najnowszy tag z origin (z GitHub)."""
    try:
        # Pobierz listę tagów z remote bez pełnego fetcha (szybkie)
        out = subprocess.check_output(
            ["git", "ls-remote", "--tags", "--sort=-v:refname", "origin"],
            cwd=Path(__file__).parent,
            stderr=subprocess.DEVNULL,
            timeout=8,
        )
        lines = out.decode().strip().splitlines()
        for line in lines:
            # Format: <sha>    refs/tags/v2026.08
            if "refs/tags/" in line:
                tag = line.split("refs/tags/")[-1].strip()
                # Pomijamy annotated tag pointers (kończą się na ^{})
                if not tag.endswith("^{}"):
                    return tag
        return None
    except Exception:
        return None


def _get_local_latest_tag() -> str:
    """Zwraca aktualny lokalny tag (lub commit jeśli nie ma tagu)."""
    try:
        out = subprocess.check_output(
            ["git", "describe", "--tags", "--abbrev=0"],
            cwd=Path(__file__).parent,
            stderr=subprocess.DEVNULL,
            timeout=2,
        )
        return out.decode().strip()
    except Exception:
        try:
            out = subprocess.check_output(
                ["git", "rev-parse", "--short", "HEAD"],
                cwd=Path(__file__).parent,
                stderr=subprocess.DEVNULL,
                timeout=2,
            )
            return out.decode().strip()
        except Exception:
            return "unknown"


def _git_pull() -> dict:
    """Pobiera najnowszy kod z origin (master, potem main jako fallback)."""
    repo = Path(__file__).parent
    last: dict = {"success": False, "stdout": "", "stderr": ""}
    for branch in ("master", "main"):
        try:
            out = subprocess.run(
                ["git", "pull", "origin", branch],
                cwd=repo,
                capture_output=True,
                text=True,
                timeout=60,
            )
            last = {
                "success": out.returncode == 0,
                "stdout": out.stdout.strip(),
                "stderr": out.stderr.strip(),
                "branch": branch,
            }
            if out.returncode == 0:
                return last
        except Exception as e:
            last = {"success": False, "stdout": "", "stderr": str(e), "branch": branch}
    return last


def _try_restart_service() -> dict:
    """
    Próbuje zrestartować aplikację.
    Najpierw próbuje user service (który używa waitress), potem zwraca instrukcję.
    """
    try:
        # Najpierw spróbuj user service (najczęstszy i najbardziej produkcyjny przypadek)
        result = subprocess.run(
            ["systemctl", "--user", "restart", "ai_analiza"],
            capture_output=True,
            text=True,
            timeout=15,
        )
        if result.returncode == 0:
            return {
                "success": True,
                "method": "user-service",
                "message": "Usługa użytkownika zrestartowana (używa Waitress)."
            }
    except Exception:
        pass

    # Fallback
    return {
        "success": False,
        "method": "manual",
        "message": "Kod zaktualizowany. Zrestartuj aplikację ręcznie najlepiej przez: ./restart-app.sh --user"
    }


# Proste cache dla GitHub Releases (żeby nie spamować API przy każdym sprawdzeniu)
_github_release_cache = {"data": None, "timestamp": 0}
GITHUB_RELEASE_CACHE_SECONDS = 300  # 5 minut


def _get_latest_github_release() -> dict | None:
    """Pobiera informacje o najnowszym releasie z GitHub (z cache)."""
    import time
    now = time.time()

    if _github_release_cache["data"] and (now - _github_release_cache["timestamp"] < GITHUB_RELEASE_CACHE_SECONDS):
        return _github_release_cache["data"]

    try:
        import urllib.request
        import json as json_module

        url = "https://api.github.com/repos/dragon15555000/AI_analiza_dokumentow/releases/latest"
        req = urllib.request.Request(url, headers={"Accept": "application/vnd.github+json", "User-Agent": "AI-Analiza-Dokumentow"})
        with urllib.request.urlopen(req, timeout=8) as response:
            data = json_module.loads(response.read().decode("utf-8"))

        release_info = {
            "tag_name": data.get("tag_name"),
            "name": data.get("name"),
            "body": data.get("body", ""),           # Markdown changelog
            "html_url": data.get("html_url"),
            "published_at": data.get("published_at"),
        }

        _github_release_cache["data"] = release_info
        _github_release_cache["timestamp"] = now
        return release_info

    except Exception:
        return None
ALLOWED_DOC_EXTENSIONS = frozenset(
    {"docx", "pdf", "xlsx", "xls", "csv", "md", "json", "txt",
     "jpg", "jpeg", "png", "tiff", "tif", "bmp"}
)


class EmbeddingError(RuntimeError):
    """Błąd generowania embeddingu — nie używaj wektora zerowego."""


def _load_persisted_collection() -> str | None:
    if _ACTIVE_COLLECTION_FILE.exists():
        name = _ACTIVE_COLLECTION_FILE.read_text(encoding="utf-8").strip()
        if name and re.match(r"^[\w\-]+$", name):
            return name
    return None


_persisted_col = _load_persisted_collection()
if _persisted_col:
    ACTIVE_COLLECTION = _persisted_col


def _persist_active_collection(name: str) -> None:
    try:
        _ACTIVE_COLLECTION_FILE.write_text(name, encoding="utf-8")
    except OSError as e:
        logger.warning("Nie zapisano aktywnej kolekcji: %s", e)


def _iter_mounted_windows_drives() -> list[Path]:
    """Zwraca zamontowane dyski Windows w WSL: /mnt/c, /mnt/d, ..."""
    mnt = Path("/mnt")
    try:
        if not mnt.is_dir():
            return []
        drives = [
            p.resolve()
            for p in mnt.iterdir()
            if p.is_dir() and re.fullmatch(r"[a-z]", p.name, re.IGNORECASE)
        ]
        return sorted(drives, key=lambda p: p.name.lower())
    except OSError as e:
        logger.debug("Nie można odczytać /mnt: %s", e)
        return []


def _resolve_allowed_roots() -> list[Path]:
    roots: list[Path] = []
    seen: set[str] = set()

    def add_root(path: Path) -> None:
        try:
            resolved = path.expanduser().resolve()
        except OSError:
            return
        if not resolved.exists():
            return
        key = str(resolved)
        if key not in seen:
            seen.add(key)
            roots.append(resolved)

    for raw in SEARCH_ROOTS:
        raw = raw.strip()
        if raw in {"/mnt/*", "/mnt/[a-z]", "/mnt/[A-Za-z]"}:
            for drive in _iter_mounted_windows_drives():
                add_root(drive)
            continue
        try:
            add_root(Path(raw))
        except OSError:
            continue
    if not roots:
        for fallback in [Path.home(), *_iter_mounted_windows_drives(), Path(__file__).parent.resolve()]:
            add_root(fallback)
    return roots


def _path_is_allowed(path: Path) -> bool:
    try:
        resolved = path.expanduser().resolve()
    except OSError:
        return False
    for root in _resolve_allowed_roots():
        try:
            resolved.relative_to(root)
            return True
        except ValueError:
            continue
    return False


def _path_is_browsable(path: Path) -> bool:
    """Folder można otworzyć tylko w obrębie SEARCH_ROOTS (bez /mnt, C: itd.)."""
    return _path_is_allowed(path)


def _default_browse_path() -> Path:
    gd = _google_drive_root()
    if gd:
        return gd
    roots = _resolve_allowed_roots()
    return roots[0] if roots else Path.home()


def _win_to_wsl(path: str) -> str | None:
    """Konwertuje ścieżkę Windows (G:\\...) na WSL przez wslpath."""
    path = (path or "").strip().strip('"').strip("'")
    if not path:
        return None
    if path.startswith("/"):
        return path
    try:
        proc = subprocess.run(
            ["wslpath", "-a", path],
            capture_output=True, text=True, timeout=15,
        )
        if proc.returncode == 0:
            out = proc.stdout.strip()
            if out:
                return out
    except Exception as e:
        logger.debug("wslpath failed: %s", e)
    return _normalize_sqlite_path_input(path) or None


def _normalize_browse_path(raw: str) -> str:
    """Akceptuje G:\\... lub /mnt/... — zwraca ścieżkę WSL."""
    raw = (raw or "").strip().strip('"').strip("'")
    if not raw:
        return str(_default_browse_path())
    if re.match(r"^[A-Za-z]:", raw):
        wsl = _win_to_wsl(raw)
        if wsl:
            return wsl
    return raw


def _google_drive_root() -> Path | None:
    """Google Drive w WSL: zwykle /mnt/g/Mój dysk (PL) lub /mnt/g/My Drive."""
    for folder in ("Mój dysk", "My Drive"):
        candidate = Path("/mnt/g") / folder
        if candidate.is_dir() and _path_is_allowed(candidate):
            try:
                return candidate.resolve()
            except OSError:
                return candidate
    g = Path("/mnt/g")
    if g.is_dir() and _path_is_allowed(g):
        try:
            return g.resolve()
        except OSError:
            return g
    return None


def _windows_user_profile() -> Path | None:
    """Profil użytkownika Windows w WSL: /mnt/c/Users/<login>."""
    c = Path("/mnt/c")
    if not c.is_dir():
        return None
    for name in (os.environ.get("USER"), os.environ.get("LOGNAME"), os.environ.get("USERNAME")):
        if not name:
            continue
        profile = c / "Users" / name
        if profile.is_dir() and _path_is_allowed(profile):
            try:
                return profile.resolve()
            except OSError:
                return profile
    return None


def _pick_windows_folder(initial_wsl: str = "") -> str | None:
    """Otwiera natywny dialog wyboru folderu Windows (działa z Google Drive G:)."""
    if platform.system() == "Windows":
        return None
    ps_lines = [
        "Add-Type -AssemblyName System.Windows.Forms",
        "$d = New-Object System.Windows.Forms.FolderBrowserDialog",
        "$d.Description = 'Wybierz folder z dokumentami (Google Drive, dysk G:, itp.)'",
        "$d.ShowNewFolderButton = $false",
    ]
    if initial_wsl:
        win_init = wsl_to_win(_normalize_browse_path(initial_wsl))
        if win_init:
            safe = win_init.replace("'", "''")
            ps_lines.append(f"$d.SelectedPath = '{safe}'")
    ps_lines += [
        "$r = $d.ShowDialog()",
        "if ($r -eq [System.Windows.Forms.DialogResult]::OK) { Write-Output $d.SelectedPath }",
    ]
    ps_script = "\n".join(ps_lines)
    import base64
    enc = base64.b64encode(ps_script.encode("utf-16-le")).decode("ascii")
    try:
        proc = subprocess.run(
            ["powershell.exe", "-NoProfile", "-STA", "-EncodedCommand", enc],
            capture_output=True, text=True, timeout=300,
        )
    except subprocess.TimeoutExpired:
        return None
    except FileNotFoundError:
        return None
    win_path = (proc.stdout or "").strip()
    if not win_path or proc.returncode != 0:
        return None
    return _win_to_wsl(win_path)


def _dir_is_readable(path: Path) -> bool:
    try:
        if not path.is_dir():
            return False
        with os.scandir(path) as it:
            next(it, None)
        return True
    except (PermissionError, OSError):
        return False


def _friendly_root_label(p: Path) -> str:
    s = str(p)
    m = re.match(r"/mnt/([a-z])$", s, re.IGNORECASE)
    if m:
        letter = m.group(1).upper()
        if letter == "G":
            return "Google Drive (G:)"
        return f"Dysk {letter}:"
    if p == Path.home():
        return "Home"
    name = p.name
    return name if name else s


def _root_kind_and_icon(path: Path) -> tuple[str, str]:
    s = str(path)
    if re.match(r"/mnt/[a-z]$", s, re.IGNORECASE):
        return "drive", "💾"
    return "root", "📂"


def _discover_browse_roots() -> list[dict]:
    """Punkty startowe importu: Google Drive, profil Windows i dozwolone dyski /mnt/<litera>."""
    drives: list[dict] = []
    seen: set[str] = set()

    gd = _google_drive_root()
    if gd:
        sp = str(gd)
        seen.add(sp)
        drives.append({
            "path": sp,
            "label": "Google Drive (G:)",
            "kind": "gdrive",
            "icon": "☁️",
            "accessible": _dir_is_readable(gd),
            "importable": True,
            "priority": 0,
        })

    profile = _windows_user_profile()
    if profile:
        sp = str(profile)
        if sp not in seen:
            seen.add(sp)
            drives.append({
                "path": sp,
                "label": "Komputer lokalny (C:)",
                "kind": "local",
                "icon": "💻",
                "accessible": _dir_is_readable(profile),
                "importable": True,
                "priority": 1,
            })

    for drive in _iter_mounted_windows_drives():
        if not _path_is_allowed(drive):
            continue
        try:
            sp = str(drive.resolve())
        except OSError:
            sp = str(drive)
        if sp in seen or (gd and sp == "/mnt/g"):
            continue
        seen.add(sp)
        drives.append({
            "path": sp,
            "label": _friendly_root_label(drive),
            "kind": "drive",
            "icon": "💾",
            "accessible": _dir_is_readable(drive),
            "importable": True,
            "priority": 2,
        })

    for i, root in enumerate(_resolve_allowed_roots()):
        try:
            sp = str(root.resolve())
        except OSError:
            sp = str(root)
        if sp in seen:
            continue
        if gd and sp == "/mnt/g":
            continue
        if profile and sp == str(profile):
            continue
        seen.add(sp)
        kind, icon = _root_kind_and_icon(root)
        drives.append({
            "path": sp,
            "label": _friendly_root_label(root),
            "kind": kind,
            "icon": icon,
            "accessible": _dir_is_readable(root),
            "importable": True,
            "priority": i + 3,
        })

    drives.sort(key=lambda d: (d["priority"], not d["accessible"]))
    return drives


def _validate_extensions(exts: list) -> list[str]:
    out: list[str] = []
    for e in exts:
        e = str(e).strip().lower().lstrip(".")
        if re.match(r"^[a-z0-9]{1,10}$", e) and e in ALLOWED_DOC_EXTENSIONS:
            out.append(e)
    if not out:
        return ["docx", "pdf", "xlsx", "xls", "csv", "md", "json"]
    return out


def _find_files_safe(folder: Path, exts: list) -> list[Path]:
    """Wyszukiwanie plików bez shella (find jako lista argv)."""
    validated = _validate_extensions(exts)
    seen: set[str] = set()
    files: list[Path] = []
    for ext in validated:
        proc = subprocess.run(
            ["find", str(folder), "-type", "f", "-name", f"*.{ext}"],
            capture_output=True,
            text=True,
            timeout=600,
        )
        if proc.returncode not in (0, 1):
            logger.warning("find zakończone kodem %s: %s", proc.returncode, proc.stderr[:200])
            continue
        for line in proc.stdout.splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                p = Path(line).resolve()
            except OSError:
                continue
            key = str(p)
            if key not in seen and p.is_file() and _path_is_allowed(p):
                seen.add(key)
                files.append(p)
    return files


def _find_file_by_name_safe(fname: str) -> Path | None:
    if not fname or not re.match(r"^[^\\/\0]+$", fname):
        return None
    for root in _resolve_allowed_roots():
        if not root.exists():
            continue
        proc = subprocess.run(
            ["find", str(root), "-type", "f", "-name", fname],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if proc.returncode not in (0, 1):
            continue
        for line in proc.stdout.splitlines()[:1]:
            line = line.strip()
            if not line:
                continue
            try:
                p = Path(line).resolve()
            except OSError:
                continue
            if p.is_file() and _path_is_allowed(p):
                return p
    return None


def _redact_sql_config(cfg: dict) -> dict:
    if not cfg:
        return {}
    out = dict(cfg)
    if out.get("password"):
        out["password"] = "********"
    return out


@app.before_request
def _require_api_key():
    # Pełne usunięcie wsparcia dla ?api_key= (zgodnie z zaleceniami audytu bezpieczeństwa).
    # Klucz API może być przekazywany TYLKO przez nagłówek HTTP "X-API-Key".
    # Powody: ochrona przed wyciekiem do logów, historii przeglądarki, Referer itd.
    # Szczególnie ważne dla endpointów SSE (/import/stream, /search/stream, /hybrid/stream).
    if not APP_API_KEY:
        return None
    if request.endpoint == "index":
        return None
    provided = request.headers.get("X-API-Key", "")
    if provided != APP_API_KEY:
        return jsonify(
            {"success": False, "error": "Brak lub nieprawidłowy klucz API (wymagany nagłówek X-API-Key)"}
        ), 401
    return None


def _is_local_request() -> bool:
    """Sprawdza czy żądanie pochodzi z localhost.
    Nagłówki X-Forwarded-For są ufane tylko gdy TRUST_PROXY=true w .env
    (używaj wyłącznie za zaufanym reverse proxy).
    """
    remote = request.remote_addr or ""
    if remote in ("127.0.0.1", "::1", "localhost"):
        return True
    if TRUST_PROXY:
        forwarded = request.headers.get("X-Forwarded-For", "") or request.headers.get("X-Real-IP", "")
        if forwarded:
            first = forwarded.split(",")[0].strip()
            if first in ("127.0.0.1", "::1", "localhost"):
                return True
    return False


def _require_config_access():
    """
    Ochrona tras konfiguracyjnych (nawet gdy APP_API_KEY jest puste).
    Wymaga: localhost LUB poprawny klucz API (jeśli jest ustawiony).
    """
    if _is_local_request():
        return None
    if not APP_API_KEY:
        # Na produkcji z 0.0.0.0 bez klucza — blokujemy dostęp do konfiguracji
        return jsonify({
            "success": False,
            "error": "Dostęp do konfiguracji tylko z localhost lub po podaniu X-API-Key"
        }), 403
    provided = request.headers.get("X-API-Key", "")
    if provided != APP_API_KEY:
        return jsonify({"success": False, "error": "Brak lub nieprawidłowy klucz API (wymagany nagłówek X-API-Key)"}), 401
    return None


# Konfiguracja LLM zapisywana przez UI (nadpisuje zmienne powyżej)
LLM_CONFIG_PATH = Path(__file__).parent / ".llm_config.json"


def _load_llm_config() -> dict:
    if LLM_CONFIG_PATH.exists():
        try:
            return json.loads(LLM_CONFIG_PATH.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def _save_llm_config(cfg: dict):
    try:
        LLM_CONFIG_PATH.write_text(json.dumps(cfg, indent=2), encoding="utf-8")
    except Exception as e:
        logger.warning("Błąd zapisu llm config: %s", e)


def _effective_app_api_key(cfg: dict | None = None) -> str:
    """Klucz ochrony API: UI (.llm_config.json) ma pierwszeństwo, pusty → fallback z .env."""
    data = cfg if cfg is not None else _load_llm_config()
    ui_key = (data.get("app_api_key") or "").strip()
    if ui_key:
        return ui_key
    return APP_API_KEY_ENV


def _apply_llm_config(cfg: dict):
    """Nadpisuje globalne zmienne LLM konfiguracją z pliku."""
    global DEFAULT_LLM_PROVIDER, OPENROUTER_API_KEY, OPENROUTER_MODEL
    global OPENROUTER_MODEL_VERIFY, OPENROUTER_FALLBACK_TO_OLLAMA
    global OLLAMA_URL, LLM_MODEL, APP_API_KEY, GEMINI_MODEL, GEMINI_API_KEY

    if cfg.get("provider"):
        DEFAULT_LLM_PROVIDER = cfg["provider"]
    if cfg.get("openrouter_key"):
        OPENROUTER_API_KEY = cfg["openrouter_key"]
    if cfg.get("openrouter_model"):
        OPENROUTER_MODEL = cfg["openrouter_model"]
    if cfg.get("openrouter_model_verify"):
        OPENROUTER_MODEL_VERIFY = cfg["openrouter_model_verify"]
    if "openrouter_fallback" in cfg:
        OPENROUTER_FALLBACK_TO_OLLAMA = bool(cfg["openrouter_fallback"])
    if cfg.get("ollama_url"):
        OLLAMA_URL = cfg["ollama_url"]
    if cfg.get("llm_model"):
        LLM_MODEL = cfg["llm_model"]
    if cfg.get("gemini_model"):
        GEMINI_MODEL = cfg["gemini_model"]
    if cfg.get("gemini_api_key"):
        GEMINI_API_KEY = cfg["gemini_api_key"]

    APP_API_KEY = _effective_app_api_key(cfg)

    # Synchronizuj konfigurację z llm_client
    llm_client._sync_llm_client_config(
        DEFAULT_LLM_PROVIDER=DEFAULT_LLM_PROVIDER,
        OPENROUTER_API_KEY=OPENROUTER_API_KEY,
        OPENROUTER_MODEL=OPENROUTER_MODEL,
        OPENROUTER_MODEL_VERIFY=OPENROUTER_MODEL_VERIFY,
        LLM_MODEL=LLM_MODEL,
        GEMINI_MODEL=GEMINI_MODEL,
        GEMINI_API_KEY=GEMINI_API_KEY,
    )


_apply_llm_config(_load_llm_config())


# ============================================================
# CHRONIONE TRASY KONFIGURACJI LLM (nawet bez APP_API_KEY)
# ============================================================

@app.route('/api/config/llm', methods=['GET'])
def get_llm_config():
    auth = _require_config_access()
    if auth:
        return auth

    cfg = _load_llm_config()
    env_gemini_key = os.environ.get("GEMINI_API_KEY", "").strip()

    # Zawsze zwracamy tylko flagi "czy klucz jest ustawiony" — nigdy preview kluczy
    return jsonify({
        "success": True,
        "provider": cfg.get("provider", DEFAULT_LLM_PROVIDER),
        "ollama_url": cfg.get("ollama_url", OLLAMA_URL),
        "llm_model": cfg.get("llm_model", LLM_MODEL),
        "gemini_model": cfg.get("gemini_model", GEMINI_MODEL),
        "openrouter_model": cfg.get("openrouter_model", OPENROUTER_MODEL),
        "openrouter_model_verify": cfg.get("openrouter_model_verify", OPENROUTER_MODEL_VERIFY),
        "openrouter_fallback": cfg.get("openrouter_fallback", OPENROUTER_FALLBACK_TO_OLLAMA),

        # Bezpieczne flagi zamiast podglądów kluczy
        "openrouter_key_set": bool(cfg.get("openrouter_key")),
        "gemini_key_set": bool(cfg.get("gemini_api_key") or env_gemini_key or GEMINI_API_KEY),
        "app_api_key_set": bool(_effective_app_api_key(cfg)),
        "app_api_key_source": (
            "ui" if (cfg.get("app_api_key") or "").strip()
            else ("env" if APP_API_KEY_ENV else None)
        ),
    })


@app.route('/api/config/llm', methods=['POST'])
def save_llm_config():
    auth = _require_config_access()
    if auth:
        return auth

    data = request.get_json() or {}

    current = _load_llm_config()

    # Aktualizuj tylko dozwolone pola
    allowed = ["provider", "ollama_url", "llm_model", "gemini_model", "openrouter_model",
               "openrouter_model_verify", "openrouter_fallback"]

    for k in allowed:
        if k in data:
            current[k] = data[k]

    # Klucze — zapisujemy tylko jeśli użytkownik je podał (nie nadpisujemy pustymi)
    if data.get("openrouter_key"):
        current["openrouter_key"] = data["openrouter_key"]
    if data.get("gemini_api_key"):
        current["gemini_api_key"] = data["gemini_api_key"]
    if "app_api_key" in data:
        ui_app_key = (data.get("app_api_key") or "").strip()
        if ui_app_key:
            current["app_api_key"] = ui_app_key
        else:
            # Pusty = usuń override UI → fallback do APP_API_KEY z .env
            current.pop("app_api_key", None)

    _save_llm_config(current)
    _apply_llm_config(current)

    return jsonify({"success": True, "message": "Konfiguracja LLM zapisana"})


# ============================================================
# LLM PROVIDER ABSTRACTION (Ollama <-> OpenRouter)
# ============================================================

def get_llm_provider(request_provider: str | None = None) -> str:
    """Zwraca aktywny provider: 'ollama', 'openrouter', 'gemini', 'ollama:<id>' lub ID custom endpointu."""
    if request_provider:
        p = request_provider.strip()
        pl = p.lower()
        if pl == "openrouter":
            return "openrouter"
        if pl == "gemini":
            return "gemini"
        if pl == "ollama" or pl.startswith("ollama:"):
            return "ollama"
        pool = _load_provider_pool()
        for endpoint in pool.get("custom_endpoints", []):
            if endpoint.get("id") == p and endpoint.get("active", True):
                return p
    return DEFAULT_LLM_PROVIDER


def _effective_llm_model(provider: str | None = None) -> str:
    """Model pokazywany w diagnostyce dla aktywnego providera."""
    prov = (provider or DEFAULT_LLM_PROVIDER or "").strip().lower()
    if prov == "openrouter":
        return OPENROUTER_MODEL
    if prov == "gemini":
        return GEMINI_MODEL
    if prov == "groq":
        return GROQ_MODEL
    return LLM_MODEL


def _is_gemini_base_url(url: str, provider_name: str = "") -> bool:
    """Sprawdza czy URL/provider wskazuje na natywne API Gemini."""
    return _custom_endpoint_is_gemini(url, provider_name)


def _redact_sensitive_log_text(text: str) -> str:
    """Redaguje oczywiste sekrety przed wysłaniem logów do chmury."""
    if not text:
        return ""
    patterns = [
        (r"(?i)(password|passwd|secret|token|api[_-]?key|authorization)\s*[:=]\s*\S+", r"\1=***REDACTED***"),
        (r"(?i)Bearer\s+\S+", "Bearer ***REDACTED***"),
    ]
    for pattern, repl in patterns:
        text = re.sub(pattern, repl, text)
    return text


def _resolve_audit_file_path(raw: str) -> Path | None:
    """Rozwiązuje ścieżkę pliku audytu (Windows/WSL) w obrębie SEARCH_ROOTS."""
    if not raw:
        return None
    try:
        path = Path(_normalize_browse_path(raw)).expanduser().resolve()
    except OSError:
        return None
    if not path.is_file() or not _path_is_allowed(path):
        return None
    if path.suffix.lower() not in GEMINI_AUDIT_EXTENSIONS:
        return None
    return path


def _stream_gemini_tokens(prompt: str, system: str = "", model: str | None = None,
                          max_tokens: int = 8192, temperature: float = 0.2):
    """Generator tokenów z Google Gemini API (streamGenerateContent)."""
    if not GEMINI_API_KEY:
        yield "Błąd: Brak GEMINI_API_KEY w .env"
        return
    effective_model = _normalize_gemini_model(model)
    url = _gemini_generate_url(effective_model, stream=True)
    payload: dict = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": temperature,
            "maxOutputTokens": max_tokens,
        },
    }
    if system:
        payload["systemInstruction"] = {"parts": [{"text": system}]}
    try:
        with _gemini_request_with_retry("POST", url, json=payload, stream=True, timeout=600) as r:
            if r.status_code >= 400:
                if r.status_code == 404:
                    yield f"[Błąd Gemini: {_gemini_model_hint(GEMINI_API_KEY, effective_model)}]"
                elif r.status_code == 429:
                    yield f"[RATE_LIMIT] {_gemini_error_message(429)}"
                else:
                    yield f"[Błąd Gemini streaming: {_gemini_error_message(r.status_code, r.text[:400])}]"
                return
            r.raise_for_status()
            for line in r.iter_lines(decode_unicode=True):
                if not line:
                    continue
                if isinstance(line, bytes):
                    line = line.decode("utf-8", errors="replace")
                if not line.startswith("data: "):
                    continue
                data_str = line[6:].strip()
                if not data_str or data_str == "[DONE]":
                    continue
                try:
                    chunk = json.loads(data_str)
                    for cand in chunk.get("candidates", []):
                        for part in cand.get("content", {}).get("parts", []):
                            token = part.get("text", "")
                            if token:
                                yield token
                except Exception:
                    continue
    except Exception as e:
        yield f"[Błąd Gemini streaming: {str(e)}]"


def _call_gemini(prompt: str, system: str, stream: bool, model: str,
                 max_tokens: int = 8192, temperature: float = 0.2):
    if not GEMINI_API_KEY:
        raise RuntimeError("Brak GEMINI_API_KEY w .env")
    effective_model = _normalize_gemini_model(model)
    if stream:
        return _gemini_request_with_retry(
            "POST",
            _gemini_generate_url(effective_model, stream=True),
            json={
                "contents": [{"role": "user", "parts": [{"text": prompt}]}],
                **({"systemInstruction": {"parts": [{"text": system}]}} if system else {}),
                "generationConfig": {"temperature": temperature, "maxOutputTokens": max_tokens},
            },
            stream=True,
            timeout=600,
            max_retries=3,
        )
    url = _gemini_generate_url(effective_model, stream=False)
    payload: dict = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": temperature, "maxOutputTokens": max_tokens},
    }
    if system:
        payload["systemInstruction"] = {"parts": [{"text": system}]}
    r = _gemini_request_with_retry("POST", url, json=payload, timeout=300, max_retries=3)
    data = r.json()
    text_parts = []
    for cand in data.get("candidates", []):
        for part in cand.get("content", {}).get("parts", []):
            if part.get("text"):
                text_parts.append(part["text"])
    return {"response": "".join(text_parts), "raw": data}


# ---- Rate limit helpers (OpenRouter) ----

# ============================================================
# HEALTH / STATUS DASHBOARD
# ============================================================

def _check_custom_endpoint_health(url: str, key: str = "", name: str = "") -> dict:
    """Sprawdza czy custom endpoint działa (Gemini, OpenAI-compatible lub Ollama /api/tags)."""
    if not url:
        return {"ok": False, "error": "Brak URL"}
    try:
        if _custom_endpoint_is_gemini(url, name):
            api_key = _gemini_key_from_url_or_value(url, key)
            t0 = time.time()
            ok, message, ms = _test_gemini_api(api_key)
            if ok:
                return {"ok": True, "url": url, "ms": ms, "error": None}
            return {
                "ok": False,
                "url": url,
                "ms": ms or int((time.time() - t0) * 1000),
                "error": message,
            }

        if _custom_endpoint_is_openai(url):
            t0 = time.time()
            headers = {"Content-Type": "application/json"}
            if key:
                headers["Authorization"] = f"Bearer {key}"
            r = requests.post(
                _openai_chat_completions_url(url),
                headers=headers,
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [{"role": "user", "content": "Say: OK"}],
                    "max_tokens": 5,
                },
                timeout=20,
            )
            ms = int((time.time() - t0) * 1000)
            return {
                "ok": r.status_code == 200,
                "url": url,
                "ms": ms,
                "status": r.status_code,
                "error": None if r.status_code == 200 else (r.text or "")[:200],
            }

        headers = {}
        if key:
            headers["Authorization"] = f"Bearer {key}"
        url_clean = url.rstrip("/")
        req = urllib.request.Request(
            f"{url_clean}/api/tags",
            headers=headers,
            method="GET"
        )
        with urllib.request.urlopen(req, timeout=5) as r:
            data = json.loads(r.read().decode("utf-8"))
            models = [m["name"] for m in data.get("models", [])]
            return {
                "ok": True,
                "url": url,
                "models_available": len(models),
                "error": None
            }
    except Exception as e:
        return {"ok": False, "url": url, "error": str(e)[:120]}


def _sanitize_for_prompt(text: str, max_len: int = 1400) -> str:
    """
    Podstawowa ochrona przed Prompt Injection.
    Usuwa/redaguje próby przejęcia kontroli nad LLM przez złośliwą treść w dokumentach.
    """
    if not text:
        return ""

    text = str(text)

    # 1. Usuwamy / redagujemy typowe frazy jailbreak / prompt injection
    jailbreak_patterns = [
        r"(?i)(ignore|disregard|forget|override|disobey).*?(previous|above|all|earlier|prior).*?instructions?",
        r"(?i)(you are now|act as|pretend to be|roleplay as).*?(different|new|another).*?(assistant|ai|model|character)",
        r"(?i)from now on.*?(you must|you will|always|never)",
        r"(?i)system prompt|initial instructions|hidden instructions",
    ]

    for pattern in jailbreak_patterns:
        text = re.sub(pattern, "[INSTRUCTION REDACTED]", text)

    # 2. Neutralizujemy bloki kodu markdown (często używane do ukrycia instrukcji)
    text = re.sub(r"```[\s\S]*?```", "[CODE BLOCK REDACTED]", text)

    # 3. Usuwamy nadmierne sekwencje backticków lub specjalnych znaków
    text = re.sub(r"`{3,}", "``", text)

    # 4. Ograniczamy długość (kontekst per chunk)
    if len(text) > max_len:
        text = text[:max_len] + " [TRUNCATED]"

    return text


def _check_qdrant_health() -> dict:
    """Sprawdza połączenie z Qdrant i aktywną kolekcję."""
    try:
        client = get_qdrant_client()
        collections = client.get_collections().collections
        collection_names = [c.name for c in collections]

        active_exists = ACTIVE_COLLECTION in collection_names

        points = 0
        if active_exists:
            info = client.get_collection(ACTIVE_COLLECTION)
            points = info.points_count or 0

        return {
            "ok": True,
            "collections_count": len(collection_names),
            "active_collection": ACTIVE_COLLECTION,
            "active_collection_exists": active_exists,
            "points_in_active": points,
            "error": None
        }
    except Exception as e:
        return {"ok": False, "error": str(e)[:150]}


def _ocr_health_status() -> dict:
    """Status OCR (Tesseract) — opcjonalny fallback dla skanów PDF i obrazów."""
    import shutil
    tesseract_bin = shutil.which("tesseract") is not None
    py_ok = pytesseract is not None
    pdf_ok = convert_from_path is not None
    available = py_ok and tesseract_bin
    hints: list[str] = []
    if not py_ok:
        hints.append("./venv/bin/pip install pytesseract pdf2image Pillow")
    if not pdf_ok:
        hints.append("pip install pdf2image (wymagane dla skanów PDF)")
    if not tesseract_bin:
        hints.append("sudo apt install tesseract-ocr tesseract-ocr-pol poppler-utils")
    return {
        "available": available,
        "python_module": py_ok,
        "tesseract_binary": tesseract_bin,
        "pdf2image": pdf_ok,
        "install_hint": " · ".join(hints) if hints else None,
        "lang": OCR_LANG,
        "note": "Uruchamiane automatycznie gdy PDF ma <20 znaków tekstu lub dla JPG/PNG/TIFF",
    }


def _file_parsers_health() -> dict:
    """Dostępność parserów plików (Excel/PDF/DOCX) — bez OCR."""
    return {
        "pdf_text": pdfplumber is not None,
        "excel": openpyxl is not None,
        "docx": docx is not None,
        "ocr_for_scans": _ocr_health_status()["available"],
        "note": "Excel: wszystkie arkusze + forensyka; PDF: tekst cyfrowy, potem OCR jeśli skan",
    }


@app.route('/qdrant/mode', methods=['GET'])
def qdrant_mode():
    is_cloud = bool(QDRANT_CLOUD_URL) and QDRANT_URL == QDRANT_CLOUD_URL
    return jsonify({
        "success": True,
        "mode": "cloud" if is_cloud else "local",
        "url": QDRANT_URL,
        "cloud_available": bool(QDRANT_CLOUD_URL),
    })


@app.route('/qdrant/switch', methods=['POST'])
def qdrant_switch():
    global QDRANT_URL, QDRANT_KEY, _qdrant_client
    data = request.get_json(force=True) or {}
    mode = data.get('mode')
    if mode == 'local':
        new_url, new_key = QDRANT_LOCAL_URL, QDRANT_LOCAL_KEY
    elif mode == 'cloud':
        if not QDRANT_CLOUD_URL:
            return jsonify({"success": False, "error": "Brak konfiguracji Qdrant Cloud (QDRANT_CLOUD_URL)"}), 400
        new_url, new_key = QDRANT_CLOUD_URL, QDRANT_CLOUD_KEY
    else:
        return jsonify({"success": False, "error": "Nieprawidłowy tryb. Użyj 'local' lub 'cloud'"}), 400
    with _qdrant_lock:
        QDRANT_URL = new_url
        QDRANT_KEY = new_key
        _qdrant_client = None
    logger.info(f"Przełączono Qdrant → {mode} ({new_url})")
    return jsonify({"success": True, "mode": mode, "url": new_url})


def call_llm(prompt: str, system: str = "", stream: bool = False,
             provider: str | None = None, model: str | None = None,
             max_tokens: int = 2000, temperature: float = 0.2) -> dict | requests.Response:
    """
    Uniwersalna funkcja do wywoływania LLM.
    Zwraca dict dla non-stream lub Response dla stream.
    provider moze byc: 'openrouter', 'ollama', lub custom endpoint ID.
    """
    pool = _load_provider_pool()
    custom_endpoint = None
    if provider:
        for endpoint in pool.get("custom_endpoints", []):
            if endpoint.get("id") == provider and endpoint.get("active", True):
                custom_endpoint = endpoint
                break

    if custom_endpoint:
        return _call_custom_endpoint(
            custom_endpoint.get("url", ""),
            custom_endpoint.get("key", ""),
            prompt, system, stream, model or LLM_MODEL,
            max_tokens, temperature,
        )

    prov = get_llm_provider(provider)

    if prov == "gemini":
        return _call_gemini(
            prompt, system, stream, model or GEMINI_MODEL, max_tokens, temperature
        )

    if prov == "openrouter":
        if not OPENROUTER_API_KEY:
            raise RuntimeError("Brak OPENROUTER_API_KEY w .env")
        return _call_openrouter(prompt, system, stream, model or OPENROUTER_MODEL, max_tokens, temperature)
    else:
        return _call_ollama(prompt, system, stream, model or LLM_MODEL, provider=provider)


def _llm_usage_from_result(result) -> dict:
    """Wyciąga usage z wyniku call_llm(); zwraca pusty dict, jeśli provider nie podał metryk."""
    if not isinstance(result, dict):
        return {}

    usage_src = result.get("usage")
    if not isinstance(usage_src, dict):
        raw = result.get("raw")
        if isinstance(raw, dict):
            usage_src = raw.get("usage") if isinstance(raw.get("usage"), dict) else raw

    if not isinstance(usage_src, dict):
        return {}

    def _as_int(value):
        try:
            if value is None:
                return None
            if isinstance(value, bool):
                return int(value)
            if isinstance(value, (int, float)):
                return max(0, int(value))
            text = str(value).strip()
            if not text:
                return None
            return max(0, int(float(text)))
        except Exception:
            return None

    prompt_tokens = _as_int(
        usage_src.get("prompt_tokens")
        or usage_src.get("input_tokens")
        or usage_src.get("prompt_eval_count")
    )
    completion_tokens = _as_int(
        usage_src.get("completion_tokens")
        or usage_src.get("output_tokens")
        or usage_src.get("eval_count")
    )
    total_tokens = _as_int(usage_src.get("total_tokens"))

    if total_tokens is None and prompt_tokens is not None and completion_tokens is not None:
        total_tokens = prompt_tokens + completion_tokens

    usage = {}
    if prompt_tokens is not None:
        usage["prompt_tokens"] = prompt_tokens
    if completion_tokens is not None:
        usage["completion_tokens"] = completion_tokens
    if total_tokens is not None:
        usage["total_tokens"] = total_tokens
    return usage


def stream_llm_tokens(prompt: str, system: str = "",
                      provider: str | None = None, model: str | None = None,
                      max_tokens: int = 2000, temperature: float = 0.2,
                      meta: dict | None = None):
    """
    Generator zwracający kolejne tokeny tekstu z LLM (Ollama, OpenRouter, custom endpoint).
    Używany w streamingowych endpointach.
    provider moze byc: 'openrouter', 'ollama', lub custom endpoint ID.
    """
    pool = _load_provider_pool()
    custom_endpoint = None
    if provider:
        for endpoint in pool.get("custom_endpoints", []):
            if endpoint.get("id") == provider and endpoint.get("active", True):
                custom_endpoint = endpoint
                break

    if custom_endpoint:
        url_base = custom_endpoint.get("url", "")
        if _custom_endpoint_is_gemini(url_base, custom_endpoint.get("name", "")):
            yield from _stream_gemini_custom_endpoint(
                url_base,
                custom_endpoint.get("key", ""),
                prompt, system,
                model or GEMINI_MODEL,
                max_tokens, temperature,
            )
            return
        if _custom_endpoint_is_openai(url_base):
            yield from _stream_openai_compatible(
                url_base,
                custom_endpoint.get("key", ""),
                prompt, system,
                model or LLM_MODEL,
                max_tokens, temperature,
            )
            return
        url = url_base.rstrip("/") + "/api/generate"
        key = custom_endpoint.get("key", "")
        headers = {}
        if key:
            headers["Authorization"] = f"Bearer {key}"
        payload = {
            "model": model or LLM_MODEL,
            "prompt": prompt,
            "system": system,
            "stream": True,
            "options": {"temperature": temperature}
        }
        try:
            with requests.post(url, json=payload, headers=headers, stream=True, timeout=300) as r:
                r.raise_for_status()
                for line in r.iter_lines():
                    if not line:
                        continue
                    try:
                        data = json.loads(line)
                        token = data.get("response", "")
                        if token:
                            yield token
                        if data.get("done", False):
                            break
                    except Exception:
                        continue
        except Exception as e:
            yield f"[Blad custom endpoint: {str(e)}]"
        return

    prov = get_llm_provider(provider)
    effective_model = model or (OPENROUTER_MODEL if prov == "openrouter" else LLM_MODEL)

    if prov == "gemini":
        yield from _stream_gemini_tokens(
            prompt, system, model or GEMINI_MODEL, max_tokens, temperature
        )
        return

    if prov == "openrouter":
        if not OPENROUTER_API_KEY:
            yield "Błąd: Brak klucza OPENROUTER_API_KEY"
            return
        last_error = None
        for candidate in _openrouter_model_candidates(effective_model):
            try:
                yield from _stream_openrouter_once(
                    prompt, system, candidate, max_tokens, temperature, meta=meta
                )
                if isinstance(meta, dict) and meta.get("provider_used") == "openrouter":
                    meta["model_used"] = meta.get("model_used") or candidate
                    meta["fallback_used"] = bool(candidate != effective_model)
                    if candidate != effective_model and meta.get("fallback_reason") is None:
                        meta["fallback_reason"] = "openrouter_model_unavailable"
                return
            except Exception as e:
                last_error = e
                if _is_openrouter_no_endpoints_error(e):
                    logger.warning(f"OpenRouter model {candidate} nie ma endpointów — próbuję model zapasowy")
                    continue
                raise

        if _is_openrouter_no_endpoints_error(last_error) and OPENROUTER_FALLBACK_TO_OLLAMA:
            logger.info("OpenRouter no-endpoints → fallback do Ollama")
            if isinstance(meta, dict):
                meta["provider_used"] = "ollama"
                meta["model_used"] = LLM_MODEL
                meta["fallback_used"] = True
                meta["fallback_reason"] = "openrouter_no_endpoints"
            ollama_url = _pick_ollama_url().rstrip("/") + "/api/generate"
            ollama_payload = {
                "model": LLM_MODEL,
                "prompt": prompt,
                "system": system,
                "stream": True,
                "options": {"temperature": temperature if temperature is not None else 0.2}
            }
            try:
                with requests.post(ollama_url, json=ollama_payload, stream=True, timeout=300) as r:
                    r.raise_for_status()
                    for line in r.iter_lines():
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            tok = data.get("response", "")
                            if tok:
                                yield tok
                            if data.get("done", False):
                                return
                        except Exception:
                            continue
                return
            except Exception as fb_err:
                yield f"[Błąd fallback na Ollama: {fb_err}]"
                return

        if _is_rate_limit_error(last_error) and OPENROUTER_FALLBACK_TO_OLLAMA:
            logger.info("OpenRouter rate limit → fallback do Ollama (OPENROUTER_FALLBACK_TO_OLLAMA=true)")
            ollama_url = _pick_ollama_url().rstrip("/") + "/api/generate"
            ollama_payload = {
                "model": LLM_MODEL,
                "prompt": prompt,
                "system": system,
                "stream": True,
                "options": {"temperature": temperature if temperature is not None else 0.2}
            }
            try:
                with requests.post(ollama_url, json=ollama_payload, stream=True, timeout=300) as r:
                    r.raise_for_status()
                    for line in r.iter_lines():
                        if not line:
                            continue
                        try:
                            data = json.loads(line)
                            tok = data.get("response", "")
                            if tok:
                                yield tok
                            if data.get("done", False):
                                return
                        except Exception:
                            continue
                return
            except Exception as fb_err:
                yield f"[Błąd fallback na Ollama: {fb_err}]"
                return

        if _is_rate_limit_error(last_error):
            friendly = "[RATE_LIMIT] Przekroczono limit zapytań OpenRouter. Poczekaj chwilę lub przełącz na Ollama w ustawieniach."
            yield friendly
        elif last_error is not None:
            yield f"[Błąd OpenRouter streaming: {last_error}]"

    else:
        # Ollama streaming
        url = _ollama_url_for_provider(provider).rstrip("/") + "/api/generate"
        payload = {
            "model": effective_model,
            "prompt": prompt,
            "system": system,
            "stream": True,
            "options": {"temperature": temperature}
        }
        try:
            with requests.post(url, json=payload, stream=True, timeout=300) as r:
                r.raise_for_status()
                for line in r.iter_lines():
                    if not line:
                        continue
                    try:
                        data = json.loads(line)
                        token = data.get("response", "")
                        if token:
                            yield token
                        if data.get("done", False):
                            break
                    except Exception:
                        continue
        except Exception as e:
            yield f"[Błąd Ollama streaming: {str(e)}]"


def _custom_endpoint_is_gemini(url: str, name: str = "") -> bool:
    """Czy URL/nazwa wskazuje na natywne API Google Gemini."""
    u = (url or "").lower()
    n = (name or "").lower().strip()
    if "generativelanguage.googleapis.com" in u:
        return True
    if n in ("g1.5f", "gemini") or n.startswith("gemini-"):
        return True
    return False


def _call_gemini_custom_endpoint(url: str, key: str, prompt: str, system: str,
                                 stream: bool, model: str = "",
                                 max_tokens: int = 2000, temperature: float = 0.2):
    """Natywny Gemini custom endpoint przez POST generateContent."""
    api_key = _gemini_key_from_url_or_value(url, key)
    if not api_key:
        raise RuntimeError("Brak klucza API Gemini")
    req_url = _gemini_custom_generate_url(url, api_key, model or GEMINI_MODEL, stream=stream)
    payload = _gemini_generate_payload(prompt, system, max_tokens, temperature)
    r = _gemini_request_with_retry(
        "POST",
        req_url,
        api_key=api_key,
        json=payload,
        stream=stream,
        timeout=300 if stream else 180,
    )
    if stream:
        return r
    data = r.json()
    return {"response": _gemini_response_text(data), "raw": data}


def _stream_gemini_custom_endpoint(url: str, key: str, prompt: str, system: str,
                                   model: str = "", max_tokens: int = 2000,
                                   temperature: float = 0.2):
    """Generator tokenów z natywnego Gemini streamGenerateContent."""
    try:
        with _call_gemini_custom_endpoint(
            url, key, prompt, system, True, model, max_tokens, temperature
        ) as r:
            if r.status_code == 429:
                yield f"[RATE_LIMIT] {_gemini_error_message(429)}"
                return
            if r.status_code >= 400:
                yield f"[Błąd Gemini custom streaming: {_gemini_error_message(r.status_code, (r.text or '')[:400])}]"
                return
            r.raise_for_status()
            for line in r.iter_lines():
                if not line:
                    continue
                decoded = line.decode("utf-8", errors="replace")
                if not decoded.startswith("data: "):
                    continue
                try:
                    chunk = json.loads(decoded[6:].strip())
                    text = _gemini_response_text(chunk)
                    if text:
                        yield text
                except Exception:
                    continue
    except Exception as exc:
        msg = str(exc)
        if "429" in msg or "Limit zapytań Gemini" in msg:
            yield f"[RATE_LIMIT] {_gemini_error_message(429)}"
        else:
            yield f"[Błąd Gemini custom streaming: {_redact_api_keys(msg)}]"


def _test_gemini_api(api_key: str, model: str | None = None) -> tuple[bool, str, int]:
    """Test Gemini API — weryfikuje model z listy generateContent."""
    if not api_key:
        return False, "Brak klucza API", 0
    effective_model = _normalize_gemini_model(model)
    t0 = time.time()
    try:
        ok, models, list_err = _list_gemini_generate_models(api_key)
        ms = int((time.time() - t0) * 1000)
        if not ok:
            return False, list_err or "Nie udało się pobrać listy modeli", ms
        if effective_model in models:
            return True, "OK", ms
        return False, _gemini_model_hint(api_key, effective_model), ms
    except Exception as exc:
        ms = int((time.time() - t0) * 1000)
        err = str(exc)
        if "404" in err:
            return False, _gemini_model_hint(api_key, effective_model), ms
        return False, err[:200], ms


def _custom_endpoint_is_openai(url: str) -> bool:
    """Czy URL wskazuje na OpenAI-compatible API (Groq, LM Studio, itp.)."""
    if _custom_endpoint_is_gemini(url):
        return False
    u = (url or "").lower().rstrip("/")
    return "openai/v1" in u or u.endswith("/v1")


def _openai_chat_completions_url(base_url: str) -> str:
    base = (base_url or "").rstrip("/")
    if base.endswith("/chat/completions"):
        return base
    if base.endswith("/v1"):
        return base + "/chat/completions"
    return base + "/chat/completions"


def _stream_openai_compatible(url: str, key: str, prompt: str, system: str,
                              model: str, max_tokens: int = 2000, temperature: float = 0.2):
    """Generator tokenów z OpenAI-compatible chat/completions (SSE)."""
    chat_url = _openai_chat_completions_url(url)
    headers = {"Content-Type": "application/json"}
    if key:
        headers["Authorization"] = f"Bearer {key}"
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})
    payload = {
        "model": model or LLM_MODEL,
        "messages": messages,
        "stream": True,
        "max_tokens": max_tokens,
        "temperature": temperature,
    }
    try:
        with requests.post(chat_url, headers=headers, json=payload, stream=True, timeout=300) as r:
            r.raise_for_status()
            for line in r.iter_lines():
                if not line:
                    continue
                line = line.decode("utf-8", errors="replace")
                if not line.startswith("data: "):
                    continue
                data_str = line[6:].strip()
                if data_str == "[DONE]":
                    return
                try:
                    chunk = json.loads(data_str)
                    choice = chunk.get("choices", [{}])[0]
                    delta = choice.get("delta", {})
                    token = delta.get("content", "")
                    if token:
                        yield token
                    if choice.get("finish_reason"):
                        return
                except Exception:
                    continue
    except Exception as e:
        yield f"[Błąd OpenAI-compatible streaming: {str(e)}]"


def _call_openai_compatible(url: str, key: str, prompt: str, system: str, stream: bool,
                            model: str, max_tokens: int = 2000, temperature: float = 0.2):
    chat_url = _openai_chat_completions_url(url)
    headers = {"Content-Type": "application/json"}
    if key:
        headers["Authorization"] = f"Bearer {key}"
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})
    payload = {
        "model": model or LLM_MODEL,
        "messages": messages,
        "stream": stream,
        "max_tokens": max_tokens,
        "temperature": temperature,
    }
    if stream:
        return requests.post(chat_url, headers=headers, json=payload, stream=True, timeout=300)
    r = requests.post(chat_url, headers=headers, json=payload, timeout=180)
    r.raise_for_status()
    return r.json()


def _call_custom_endpoint(url: str, key: str, prompt: str, system: str, stream: bool,
                          model: str = "", max_tokens: int = 2000, temperature: float = 0.2):
    """Custom endpoint: Gemini, OpenAI-compatible albo Ollama."""
    if _custom_endpoint_is_gemini(url):
        return _call_gemini_custom_endpoint(
            url, key, prompt, system, stream, model or GEMINI_MODEL, max_tokens, temperature
        )
    if _custom_endpoint_is_openai(url):
        return _call_openai_compatible(
            url, key, prompt, system, stream, model or LLM_MODEL, max_tokens, temperature
        )
    url_clean = url.rstrip("/") + "/api/generate"
    headers = {}
    if key:
        headers["Authorization"] = f"Bearer {key}"
    payload = {
        "model": model or LLM_MODEL,
        "prompt": prompt,
        "system": system,
        "stream": stream,
        "options": {"temperature": 0.2}
    }
    if stream:
        return requests.post(url_clean, json=payload, headers=headers, stream=True, timeout=300)
    else:
        r = requests.post(url_clean, json=payload, headers=headers, timeout=180)
        r.raise_for_status()
        return r.json()


    if _is_openrouter_no_endpoints_error(last_error) and OPENROUTER_FALLBACK_TO_OLLAMA:
        logger.info("OpenRouter no-endpoints (non-stream) → fallback do Ollama")
        result = _call_ollama(prompt, system, stream=False, model=LLM_MODEL)
        if isinstance(result, dict):
            return result
        return {"response": str(result)}

    raise last_error if last_error else RuntimeError("OpenRouter call failed")


# ---- Qdrant Client (reuse connection) ----
_qdrant_client = None
_qdrant_lock = threading.Lock()

def get_qdrant_client() -> QdrantClient:
    global _qdrant_client
    with _qdrant_lock:
        if _qdrant_client is None:
            from httpx import Limits
            _qdrant_client = QdrantClient(
                url=QDRANT_URL,
                api_key=QDRANT_KEY,
                timeout=30.0,
                limits=Limits(
                    max_keepalive_connections=5,
                    max_connections=20,
                    keepalive_expiry=30
                )
            )
        return _qdrant_client

# ---- Cache (dokumenty i sugestie) ----
_docs_cache = {"data": None, "ts": 0}
DOCS_CACHE_TTL = 300  # 5 minut
_suggestions_cache = {"data": None, "ts": 0}
SUGGESTIONS_TTL = 1800  # 30 minut

# ---- Multi-agent Swarm ----
FREE_MODELS_LIST = [
    ("meta-llama/llama-3.3-70b-instruct:free",  "Llama 3.3 70B ★"),
    ("google/gemini-2.0-flash-exp:free",         "Gemini 2.0 Flash"),
    ("mistralai/mistral-7b-instruct:free",       "Mistral 7B"),
    ("qwen/qwen-2.5-7b-instruct:free",           "Qwen 2.5 7B"),
    ("meta-llama/llama-3.1-8b-instruct:free",    "Llama 3.1 8B"),
]

SWARM_MODES = {
    "privacy": {
        "label": "🔒 Incognito",
        "desc": "Każdy worker dostaje losowy model i potasowane fragmenty. Żaden LLM nie widzi >20% dokumentu.",
        "workers": 5,
        "shuffle_chunks": True,
        "use_random_models": True,
        "max_chunks_per_worker": 4,
        "synthesis_model": "meta-llama/llama-3.3-70b-instruct:free",
        "hierarchical": False,
    },
    "speed": {
        "label": "⚡ Szybkość",
        "desc": "Maksymalny paralelizm z najszybszym modelem. Wyniki streamowane na bieżąco.",
        "workers": 4,
        "shuffle_chunks": False,
        "use_random_models": False,
        "preferred_model": ("google/gemini-2.0-flash-exp:free", "Gemini 2.0 Flash"),
        "max_chunks_per_worker": 5,
        "synthesis_model": "google/gemini-2.0-flash-exp:free",
        "hierarchical": False,
    },
    "quality": {
        "label": "🎯 Jakość",
        "desc": "Hierarchiczna synteza: workery → meta-analiza → finalny wynik z najlepszym modelem.",
        "workers": 3,
        "shuffle_chunks": False,
        "use_random_models": False,
        "preferred_model": ("meta-llama/llama-3.3-70b-instruct:free", "Llama 3.3 70B ★"),
        "max_chunks_per_worker": 6,
        "synthesis_model": "meta-llama/llama-3.3-70b-instruct:free",
        "hierarchical": True,
    },
}

# ---- Cache embeddingów (SQLite) ----
_CACHE_DB_PATH = Path(__file__).parent / "embedding_cache.db"

def _init_embed_cache():
    """Inicjalizuje tabelę w bazie SQLite, jeśli nie istnieje."""
    try:
        with sqlite3.connect(_CACHE_DB_PATH, timeout=10) as conn:
            conn.execute("CREATE TABLE IF NOT EXISTS embeddings (hash TEXT PRIMARY KEY, vector TEXT)")
        # Policz istniejące wpisy
        with sqlite3.connect(_CACHE_DB_PATH, timeout=10) as conn:
            cursor = conn.execute("SELECT COUNT(*) FROM embeddings")
            count = cursor.fetchone()[0]
        logger.info(f"Załadowano cache embeddingów (SQLite): {count} wpisów")
    except Exception as e:
        logger.warning(f"Błąd inicjalizacji cache embeddingów: {e}")

_init_embed_cache()

def _ensure_collection_exists():
    """Tworzy ACTIVE_COLLECTION jeśli nie istnieje (np. świeży lokalny Qdrant)."""
    try:
        from qdrant_client.models import VectorParams, Distance
        client = get_qdrant_client()
        if not client.collection_exists(ACTIVE_COLLECTION):
            client.create_collection(
                ACTIVE_COLLECTION,
                vectors_config=VectorParams(size=768, distance=Distance.COSINE)
            )
            logger.info(f"Kolekcja '{ACTIVE_COLLECTION}' utworzona automatycznie")
    except Exception as e:
        logger.warning(f"Nie udało się sprawdzić/utworzyć kolekcji '{ACTIVE_COLLECTION}': {e}")

_ensure_collection_exists()

def get_embedding(text: str) -> list:
    import hashlib as _hl
    key = _hl.sha256(text[:1500].encode('utf-8', errors='replace')).hexdigest()

    # 1. Sprawdzaj, czy wektor jest w bazie
    try:
        with sqlite3.connect(_CACHE_DB_PATH, timeout=10) as conn:
            cursor = conn.execute("SELECT vector FROM embeddings WHERE hash = ?", (key,))
            row = cursor.fetchone()
            if row:
                return json.loads(row[0])
    except Exception as e:
        logger.warning(f"Błąd odczytu cache embeddingów: {e}")

    # 2. Jeśli nie ma w cache, odpytaj Ollamę (z retry przy Connection reset)
    url = OLLAMA_URL + "/api/embeddings"
    payload = {"model": EMBED_MODEL, "prompt": text[:1500]}

    for attempt in range(4):  # max 4 próby
        try:
            req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"),
                                         headers={"Content-Type": "application/json"}, method="POST")
            with urllib.request.urlopen(req, timeout=45) as r:
                vec = json.loads(r.read().decode("utf-8"))["embedding"]

            # 3. Zapisz nowy wektor natychmiast do bazy SQLite
            try:
                with sqlite3.connect(_CACHE_DB_PATH, timeout=10) as conn:
                    conn.execute("INSERT OR REPLACE INTO embeddings (hash, vector) VALUES (?, ?)",
                                 (key, json.dumps(vec)))
            except Exception as e:
                logger.warning(f"Błąd zapisu cache embeddingów: {e}")

            return vec

        except Exception as e:
            if "Connection reset by peer" in str(e) or "104" in str(e):
                wait = (2 ** attempt) * 0.4   # 0.4s, 0.8s, 1.6s, 3.2s
                logger.warning(f"Ollama embedding reset (próba {attempt+1}/4) — czekam {wait:.1f}s...")
                time.sleep(wait)
                continue
            else:
                logger.error(f"Ollama Embedding Error: {e}")
                raise EmbeddingError(f"Błąd embeddingu Ollama: {e}") from e

    logger.error("Ollama Embedding Error: wszystkie próby nieudane (Connection reset)")
    raise EmbeddingError("Embedding niedostępny — Ollama nie odpowiada (connection reset)")

def get_embeddings_batch(texts: list, batch_size: int = 6) -> list:
    """Batch embeddings z mniejszą równoległością (domyślnie 6 zamiast 8), żeby mniej obciążać Ollamę."""
    results = [None] * len(texts)

    def embed_one(idx_text):
        idx, text = idx_text
        return idx, get_embedding(text)

    with ThreadPoolExecutor(max_workers=batch_size) as ex:
        futures = {ex.submit(embed_one, (i, t)): i for i, t in enumerate(texts)}
        for fut in as_completed(futures):
            try:
                idx, vec = fut.result()
                results[idx] = vec
            except EmbeddingError:
                raise
            except Exception as e:
                logger.error(f"Batch embedding error: {e}")
                raise EmbeddingError(f"Błąd batch embeddingu: {e}") from e

    return results
# Statystyki pingów modeli floty (utrwalane między restartami).
_persisted_model_live = _load_model_live_stats()
_model_live: dict = {}
for _mid in MODEL_REGISTRY:
    _defaults = {
        "ping_ok": None,
        "ping_ms": None,
        "ping_ts": None,
        "err_count": 0,
        "ok_count": 0,
        "avg_ms": None,
    }
    _saved = _persisted_model_live.get(_mid, {})
    if isinstance(_saved, dict):
        _defaults.update({k: _saved.get(k) for k in _defaults if k in _saved})
    _model_live[_mid] = _defaults

# Inicjalizuj models_fleet z globalnym stanem
_set_model_live(_model_live)
_set_MODEL_REGISTRY(MODEL_REGISTRY)


# Wrapper dla models_fleet._get_best_custom_endpoints — dostarcz provider_pool
def _get_best_custom_endpoints_wrapper(limit: int = 3) -> list[tuple]:
    """Wrapper dostarcza provider_pool do models_fleet._get_best_custom_endpoints."""
    from models_fleet import _get_best_custom_endpoints as _get_best_from_fleet
    pool = _load_provider_pool()
    return _get_best_from_fleet(limit=limit, provider_pool=pool)


def _embedding_query_for_mode(query_text: str, mode: str) -> str:
    """Zapytanie do embeddingu — bez historii rozmowy; detective lekko poszerza semantykę."""
    q = (query_text or "").strip()
    if mode != "detective" or not q:
        return q
    hint = (
        " niespójność anomalia rozbieżność kwota data umowa przetarg "
        "sprzeczność brak dokumentu podejrzany wzorzec"
    )
    return (q + hint)[:2400]


def _effective_search_limit(mode: str, limit: int) -> int:
    limit = min(max(int(limit or 5), 1), 20)
    if mode == "detective":
        floor = int(SEARCH_MODES["detective"].get("min_limit", 12))
        return max(limit, floor)
    return limit


def _diversify_contexts(raw: list, limit: int, max_per_file: int = 2) -> list:
    """Wybiera chunki z różnych plików — lepsze porównania między dokumentami."""
    if not raw or limit <= 0:
        return []
    per_file: dict[str, int] = {}
    out = []
    for item in raw:
        fname = item.get("file") or "Nieznany"
        if per_file.get(fname, 0) >= max_per_file:
            continue
        per_file[fname] = per_file.get(fname, 0) + 1
        out.append(item)
        if len(out) >= limit:
            break
    if len(out) < limit:
        seen = {(c.get("file"), (c.get("text") or "")[:80]) for c in out}
        for item in raw:
            key = (item.get("file"), (item.get("text") or "")[:80])
            if key in seen:
                continue
            out.append(item)
            seen.add(key)
            if len(out) >= limit:
                break
    return out


def _retrieve_search_contexts(
    client,
    query_text: str,
    limit: int,
    file_filter: str | None,
    mode: str,
) -> tuple[list, list]:
    """Wyszukiwanie wektorowe + opcjonalna dywersyfikacja plików (detective)."""
    from qdrant_client.models import Filter, FieldCondition, MatchValue

    embed_q = _embedding_query_for_mode(query_text, mode)
    vector = get_embedding(embed_q)
    eff_limit = _effective_search_limit(mode, limit)
    fetch_limit = eff_limit
    max_per_file = 3
    if mode == "detective":
        max_per_file = int(SEARCH_MODES["detective"].get("max_per_file", 2))
        if not file_filter:
            fetch_limit = min(eff_limit * 3, 45)

    if file_filter:
        qfilter = Filter(must=[FieldCondition(key="file", match=MatchValue(value=file_filter))])
        res = client.query_points(
            collection_name=ACTIVE_COLLECTION,
            query=vector,
            limit=fetch_limit,
            query_filter=qfilter,
        )
    else:
        res = client.query_points(
            collection_name=ACTIVE_COLLECTION,
            query=vector,
            limit=fetch_limit,
        )

    raw_contexts = []
    for point in res.points:
        p = point.payload
        raw_contexts.append({
            "file": p.get("file", ""),
            "text": p.get("text", ""),
            "score": float(point.score),
            "full_path": p.get("full_path", ""),
            "point_id": str(point.id),
        })

    if mode == "detective" and len(raw_contexts) > eff_limit:
        raw_contexts = _diversify_contexts(raw_contexts, eff_limit, max_per_file=max_per_file)

    llm_contexts = [{"file": c["file"], "text": c["text"]} for c in raw_contexts]
    results = [
        {
            "file": c.get("file") or "Nieznany",
            "score": f"{c.get('score', 0):.4f}",
            "point_id": c.get("point_id", ""),
            "snippet": _search_result_snippet(c.get("text", ""), query_text),
            "full_path": c.get("full_path", ""),
            "win_path": wsl_to_win(c.get("full_path", "")),
        }
        for c in raw_contexts
    ]
    return llm_contexts, results


def _build_search_prompt(
    query_text: str,
    contexts: list,
    mode: str,
    chat_context: str = "",
) -> tuple[str, str]:
    """Buduje prompt użytkownika i zwraca (prompt, system) — chat_context tylko do LLM."""
    cfg = SEARCH_MODES.get(mode, SEARCH_MODES["normal"])
    max_chunk = 1600 if mode == "detective" else 1400
    sanitized = [
        {"file": c.get("file", ""), "text": _sanitize_for_prompt(c.get("text", ""), max_chunk)}
        for c in contexts
    ]
    context_str = "\n\n".join(
        [f"[Dokument: {c['file']}]: {c['text']}" for c in sanitized]
    )
    chat_block = ""
    if chat_context and chat_context.strip():
        chat_block = CHAT_CONTEXT_BLOCK_TEMPLATE.format(
            chat_context=_sanitize_for_prompt(chat_context.strip(), 2500)
        )

    if mode == "detective":
        prompt = DETECTIVE_PROMPT_TEMPLATE.format(
            contexts=context_str,
            chat_block=chat_block,
            query=query_text,
            prompt_suffix=cfg['prompt_suffix'],
            format=cfg.get('format', '')
        )
    else:
        prompt = DEFAULT_PROMPT_TEMPLATE.format(
            contexts=context_str,
            chat_block=chat_block,
            query=query_text,
            prompt_suffix=cfg['prompt_suffix']
        )
    return prompt, cfg["system"]


def generate_answer(
    query: str,
    contexts: list,
    mode: str = "normal",
    provider: str | None = None,
    model: str | None = None,
    chat_context: str = "",
) -> str:
    """Generuje odpowiedź używając wybranego providera (ollama lub openrouter)."""
    prompt, system = _build_search_prompt(query, contexts, mode, chat_context=chat_context)
    cfg = SEARCH_MODES.get(mode, SEARCH_MODES["normal"])

    try:
        result = call_llm(
            prompt=prompt,
            system=system,
            stream=False,
            provider=provider,
            model=model
        )
        if isinstance(result, dict):
            return result.get("response", str(result))
        return str(result)
    except Exception as e:
        prov = get_llm_provider(provider)
        if _is_rate_limit_error(e) and OPENROUTER_FALLBACK_TO_OLLAMA and prov == "openrouter":
            logger.warning("generate_answer: OpenRouter 429 → fallback do Ollama")
            try:
                result = _call_ollama(prompt, system, stream=False, model=LLM_MODEL)
                if isinstance(result, dict):
                    return result.get("response", str(result))
                return str(result)
            except Exception as fb_e:
                return f"[RATE_LIMIT + Błąd fallback] {fb_e}"
        logger.error(f"LLM error in generate_answer ({prov}): {e}")
        if _is_rate_limit_error(e):
            return "[RATE_LIMIT] Przekroczono limit OpenRouter. Spróbuj później lub użyj Ollama."
        return f"Błąd syntezy LLM ({prov}): {e}"

def verify_answer(answer: str, contexts: list, query: str, provider: str | None = None, model: str | None = None) -> dict:
    """Krytyk — używa wybranego providera (call_llm)."""
    # Ochrona przed Prompt Injection w weryfikacji
    sanitized_contexts = [
        {"file": c.get("file", ""), "text": _sanitize_for_prompt(c.get("text", ""), 1000)}
        for c in contexts
    ]
    context_str = "\n\n".join([f"[{c['file']}]: {c['text']}" for c in sanitized_contexts])
    system = VERIFY_SYSTEM_PROMPT
    prompt = VERIFY_PROMPT_TEMPLATE.format(
        contexts=context_str,
        query=query,
        answer=answer
    )

    try:
        # Używamy dedykowanego modelu do weryfikacji (jeśli ustawiony), żeby nie dobić limitu tego samego modelu co główna odpowiedź
        effective_verify_model = model
        if not effective_verify_model and get_llm_provider(provider) == "openrouter":
            effective_verify_model = OPENROUTER_MODEL_VERIFY
        result = call_llm(prompt=prompt, system=system, stream=False, provider=provider, model=effective_verify_model)
        raw = result.get("response", str(result)) if isinstance(result, dict) else str(result)

        # Wyciągnij werdykt
        verdict = "NIEOKREŚLONY"
        justification = ""
        for line in raw.splitlines():
            l = line.strip()
            if l.startswith("WERDYKT:"):
                verdict = l.replace("WERDYKT:", "").strip()
            if l.startswith("UZASADNIENIE:"):
                justification = l.replace("UZASADNIENIE:", "").strip()

        confirmed = raw.count("✓")
        partial    = raw.count("⚠")
        hallucin   = raw.count("✗")
        total = confirmed + partial + hallucin
        confidence_pct = round(confirmed / total * 100) if total > 0 else None

        return {
            "success": True,
            "raw": raw,
            "verdict": verdict,
            "justification": justification,
            "confirmed": confirmed,
            "partial": partial,
            "hallucinations": hallucin,
            "confidence_pct": confidence_pct
        }
    except urllib.error.URLError as e:
        logger.error(f"LLM connection error (verify_answer): {e}")
        return {"success": False, "error": "Błąd połączenia z weryfikatorem (Ollama)."}
    except Exception as e:
        logger.error(f"LLM error in verify_answer: {e}")
        return {"success": False, "error": str(e)}

@app.route('/verify', methods=['POST'])
def verify_endpoint():
    data = request.get_json()
    answer   = data.get('answer', '').strip()
    query    = data.get('query', '').strip()
    contexts = data.get('contexts', [])
    llm_provider = data.get('llm_provider')
    openrouter_model = data.get('openrouter_model')
    if not answer or not query or not contexts:
        return jsonify({"success": False, "error": "Brak danych do weryfikacji"})
    # Przekazujemy provider, żeby weryfikacja też szła przez wybrany model (w tym osobny model_verify)
    result = verify_answer(answer, contexts, query, provider=llm_provider, model=openrouter_model)
    return jsonify(result)

def _safe_int_clamp(val, lo: int, hi: int, default: int = 1) -> int:
    try:
        return max(lo, min(hi, int(float(val))))
    except (TypeError, ValueError):
        return default


def highlight_backend(text: str, query: str) -> str:
    escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    if not query:
        return escaped
    words = [w.strip() for w in query.split() if len(w.strip()) > 2]
    if not words: return escaped
    for word in words:
        clean_word = re.sub(r'[.,\/#!$%\^&\*;:{}=\-_`~()]', '', word)
        root = clean_word[:-2] if len(clean_word) > 4 else clean_word
        if not root: continue
        try:
            pattern = re.compile(rf"({root}[a-zA-ZąćęłńóśźżĄĆĘŁŃÓŚŹŻ]*)", re.IGNORECASE)
            escaped = pattern.sub(r"<mark>\1</mark>", escaped)
        except re.error:
            continue
    return escaped


def _search_result_snippet(text: str, query: str, max_len: int = 220) -> str:
    """Krótki podgląd do listy wyników — pełny tekst przez /api/get_context."""
    plain = re.sub(r"<[^>]+>", "", text or "")
    plain = plain.replace("\n", " ").strip()
    if len(plain) <= max_len:
        return highlight_backend(plain, query)
    return highlight_backend(plain[:max_len].rsplit(" ", 1)[0] + "…", query)


def _mtime_in_range(path: Path, modified_after: str | None, modified_before: str | None) -> bool:
    if not modified_after and not modified_before:
        return True
    try:
        mtime = path.stat().st_mtime
    except OSError:
        return False
    if modified_after:
        try:
            ts = time.mktime(time.strptime(modified_after[:10], "%Y-%m-%d"))
            if mtime < ts:
                return False
        except ValueError:
            pass
    if modified_before:
        try:
            ts = time.mktime(time.strptime(modified_before[:10], "%Y-%m-%d")) + 86399
            if mtime > ts:
                return False
        except ValueError:
            pass
    return True


def _doc_modified_iso(doc: dict) -> str:
    meta = doc.get("metadata") or {}
    return (meta.get("modified") or "")[:10]


def _filter_documents_list(
    docs: list,
    ext_filter: list[str] | None = None,
    modified_after: str = "",
    modified_before: str = "",
) -> list:
    out = docs
    if ext_filter:
        allowed = {e.lower().lstrip(".") for e in ext_filter}
        out = [d for d in out if d.get("file", "").rsplit(".", 1)[-1].lower() in allowed]
    if modified_after:
        out = [d for d in out if _doc_modified_iso(d) >= modified_after[:10]]
    if modified_before:
        out = [d for d in out if _doc_modified_iso(d) <= modified_before[:10] or not _doc_modified_iso(d)]
    return out

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

def make_chunks(text: str) -> list:
    chunks = []
    step = CHUNK_SIZE - CHUNK_OVERLAP
    start = 0
    while start < len(text):
        end = start + CHUNK_SIZE
        chunk = text[start:end]
        # Zakończ na granicy zdania jeśli możliwe (ostatnia kropka/newline w ogonie)
        if end < len(text):
            for sep in ('. ', '.\n', '\n\n', '\n', ' '):
                cut = chunk.rfind(sep, CHUNK_SIZE // 2)
                if cut != -1:
                    chunk = chunk[:cut + len(sep)]
                    break
        if chunk.strip():
            chunks.append(chunk)
        start += max(len(chunk) - CHUNK_OVERLAP, step)
    return chunks

def _extract_xls(file_path: Path) -> str:
    """Parser starego formatu .xls — czyta też ukryte wiersze i kolumny."""
    try:
        import xlrd
        wb = xlrd.open_workbook(str(file_path), formatting_info=True)
        parts = []
        for sheet in wb.sheets():
            if sheet.nrows == 0:
                continue

            # Wykryj ukryte wiersze i kolumny
            hidden_rows = set()
            hidden_cols = set()
            try:
                for ri in range(sheet.nrows):
                    ri_obj = sheet.rowinfo_map.get(ri)
                    if ri_obj and ri_obj.hidden:
                        hidden_rows.add(ri)
            except Exception:
                pass
            try:
                for ci in range(sheet.ncols):
                    ci_obj = sheet.colinfo_map.get(ci)
                    if ci_obj and ci_obj.hidden:
                        hidden_cols.add(ci)
            except Exception:
                pass

            hidden_info = []
            if hidden_rows: hidden_info.append(f"{len(hidden_rows)} ukrytych wierszy")
            if hidden_cols: hidden_info.append(f"{len(hidden_cols)} ukrytych kolumn")

            lines = [f"\n=== Arkusz: {sheet.name}"
                     + (f" [UWAGA: {', '.join(hidden_info)}]" if hidden_info else "") + " ==="]

            # Nagłówki
            headers = []
            header_row = -1
            for ri in range(min(5, sheet.nrows)):
                row = [str(sheet.cell_value(ri, ci)).strip() for ci in range(sheet.ncols)]
                if len([v for v in row if v and v != '0.0']) >= 2:
                    headers = row
                    lines.append("Nagłówki: " + " | ".join(h for h in headers if h))
                    header_row = ri
                    break

            for ri in range(sheet.nrows):
                if ri == header_row:
                    continue
                is_hidden_row = ri in hidden_rows
                cells = []
                for ci in range(sheet.ncols):
                    val = sheet.cell_value(ri, ci)
                    if val == '' or val is None:
                        continue
                    ctype = sheet.cell_type(ri, ci)
                    if ctype == xlrd.XL_CELL_DATE:
                        try:
                            val = xlrd.xldate_as_datetime(val, wb.datemode).strftime('%Y-%m-%d')
                        except Exception:
                            pass
                    elif ctype == xlrd.XL_CELL_NUMBER and float(val) == int(val) and abs(val) < 1e12:
                        val = int(val)
                    hdr = headers[ci] if ci < len(headers) and headers[ci] else f"Kol{ci+1}"
                    hidden_col_mark = " [UKRYTA_KOL]" if ci in hidden_cols else ""
                    cells.append(f"{hdr}{hidden_col_mark}: {val}")

                if cells:
                    prefix = "[UKRYTY_WIERSZ] " if is_hidden_row else ""
                    lines.append(prefix + " | ".join(cells))

            parts.append("\n".join(lines))
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning(f"Błąd parsowania .xls {file_path.name}: {e}")
        return ""

def _extract_excel(file_path: Path) -> str:
    # Stary format — użyj xlrd
    if file_path.suffix.lower() == '.xls':
        return _extract_xls(file_path)
    parts = []
    # Dwa odczyty: raz z formułami, raz z wartościami (cache Excel)
    try:
        wb_vals = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
    except Exception:
        wb_vals = None
    try:
        wb_form = openpyxl.load_workbook(file_path, data_only=False, read_only=True)
    except Exception:
        wb_form = None

    if not wb_vals and not wb_form:
        return ""

    sheet_names = (wb_vals or wb_form).sheetnames

    for sheet_name in sheet_names:
        ws_v = wb_vals[sheet_name]  if wb_vals and sheet_name in (wb_vals.sheetnames)  else None
        ws_f = wb_form[sheet_name]  if wb_form and sheet_name in (wb_form.sheetnames)  else None
        if not ws_v and not ws_f:
            continue

        rows_v = list(ws_v.iter_rows(values_only=True))        if ws_v else []
        rows_f = list(ws_f.iter_rows(values_only=False))       if ws_f else []
        if not rows_v and not rows_f:
            continue

        # Wykryj ukryte wiersze i kolumny (tylko bez read_only)
        hidden_rows = set()
        hidden_cols = set()
        try:
            ws_meta = openpyxl.load_workbook(file_path, data_only=True)[sheet_name]
            for ri_h, rd in ws_meta.row_dimensions.items():
                if rd.hidden:
                    hidden_rows.add(ri_h - 1)  # 0-based
            for col_letter, cd in ws_meta.column_dimensions.items():
                if cd.hidden:
                    from openpyxl.utils import column_index_from_string
                    hidden_cols.add(column_index_from_string(col_letter) - 1)
        except Exception:
            pass

        hidden_info = []
        if hidden_rows: hidden_info.append(f"{len(hidden_rows)} ukrytych wierszy")
        if hidden_cols: hidden_info.append(f"{len(hidden_cols)} ukrytych kolumn")

        # Nagłówki kolumn — pierwsza niepusta linia z wartościami tekstowymi
        headers = []
        header_row_idx = -1
        source_rows = rows_v or [list(r) for r in rows_f]
        for ri_h, row in enumerate(source_rows[:5]):
            if sum(1 for c in row if c is not None and str(c).strip()) >= 2:
                headers = [str(c).strip() if c is not None else f"Kol{i+1}"
                           for i, c in enumerate(row)]
                header_row_idx = ri_h
                break

        lines = [f"\n=== Arkusz: {sheet_name}"
                 + (f" [UWAGA: {', '.join(hidden_info)}]" if hidden_info else "") + " ==="]
        if headers:
            lines.append("Nagłówki: " + " | ".join(headers))

        for ri, row_v in enumerate(rows_v):
            if ri == header_row_idx:
                continue
            vals = [v for v in row_v if v is not None and str(v).strip()]
            if not vals:
                continue

            is_hidden = ri in hidden_rows
            cells = []
            for ci, val in enumerate(row_v):
                if val is None:
                    continue
                hdr = headers[ci] if ci < len(headers) else f"Kol{ci+1}"
                hidden_col_mark = " [UKRYTA_KOL]" if ci in hidden_cols else ""

                formula = None
                if rows_f and ri < len(rows_f):
                    rf_row = rows_f[ri]
                    if ci < len(rf_row):
                        fval = rf_row[ci].value if hasattr(rf_row[ci], 'value') else None
                        if isinstance(fval, str) and fval.startswith('='):
                            formula = fval

                if formula:
                    cells.append(f"{hdr}{hidden_col_mark}: {val} [wzór: {formula}]")
                else:
                    cells.append(f"{hdr}{hidden_col_mark}: {val}")

            if cells:
                prefix = "[UKRYTY_WIERSZ] " if is_hidden else ""
                lines.append(prefix + " | ".join(cells))

        # Jeśli data_only zwrócił same None (plik niezapisany przez Excel),
        # fallback: odczytaj formuły bezpośrednio
        data_lines = len([l for l in lines if l.startswith("Kol") or ": " in l])
        if data_lines == 0 and rows_f:
            for rf_row in rows_f:
                cells = []
                for ci, cell in enumerate(rf_row):
                    v = cell.value if hasattr(cell, 'value') else None
                    if v is None:
                        continue
                    hdr = headers[ci] if ci < len(headers) else f"Kol{ci+1}"
                    cells.append(f"{hdr}: {v}")
                if cells:
                    lines.append(" | ".join(cells))

        parts.append("\n".join(lines))

    if wb_vals:
        try: wb_vals.close()
        except Exception: pass
    if wb_form:
        try: wb_form.close()
        except Exception: pass

    return "\n\n".join(parts)


def extract_file_metadata(f_path: Path) -> dict:
    """Zbiera bogate metadane pliku (systemowe + formatowe) do celów śledczych."""
    meta = {
        "file_name": f_path.name,
        "full_path": str(f_path),
        "size_bytes": None,
        "size_human": None,
        "created": None,
        "modified": None,
        "accessed": None,
        "file_type": f_path.suffix.lower().lstrip('.'),
        "embedded": {}
    }

    try:
        stat = f_path.stat()
        meta["size_bytes"] = stat.st_size
        meta["size_human"] = f"{stat.st_size / 1024:.1f} KB" if stat.st_size < 10*1024*1024 else f"{stat.st_size / (1024*1024):.2f} MB"

        # Czas modyfikacji jest wiarygodny wszędzie
        meta["modified"] = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stat.st_mtime))

        # Czas dostępu
        if hasattr(stat, "st_atime"):
            meta["accessed"] = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stat.st_atime))

        # Czas utworzenia — Windows ma st_ctime jako creation time, Unix ma jako change time
        if os.name == 'nt' and hasattr(stat, "st_ctime"):
            meta["created"] = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stat.st_ctime))
        else:
            # Na Linux/WSL często nie mamy wiarygodnego created time z filesystemu
            meta["created"] = None

    except Exception as e:
        logger.warning(f"Błąd pobierania stat dla {f_path.name}: {e}")

    # === Metadane wbudowane w pliki ===
    ext = f_path.suffix.lower()

    # PDF
    if ext == '.pdf' and pdfplumber:
        try:
            with pdfplumber.open(f_path) as pdf:
                if pdf.metadata:
                    meta["embedded"]["pdf"] = {k: str(v) for k, v in pdf.metadata.items() if v}
        except Exception:
            pass

    # DOCX
    if ext == '.docx' and docx:
        try:
            d = docx.Document(f_path)
            core_props = d.core_properties
            embedded = {}
            if core_props.author: embedded["author"] = core_props.author
            if core_props.title: embedded["title"] = core_props.title
            if core_props.last_modified_by: embedded["last_modified_by"] = core_props.last_modified_by
            if core_props.created: embedded["created"] = str(core_props.created)
            if core_props.modified: embedded["modified"] = str(core_props.modified)
            if embedded:
                meta["embedded"]["docx"] = embedded
        except Exception:
            pass

    # XLSX / XLS
    if ext in ['.xlsx', '.xls'] and openpyxl:
        try:
            wb = openpyxl.load_workbook(f_path, data_only=True)
            props = wb.properties
            embedded = {}
            if props.creator: embedded["author"] = props.creator
            if props.title: embedded["title"] = props.title
            if props.lastModifiedBy: embedded["last_modified_by"] = props.lastModifiedBy
            if props.created: embedded["created"] = str(props.created)
            if props.modified: embedded["modified"] = str(props.modified)
            if embedded:
                meta["embedded"]["excel"] = embedded
            wb.close()
        except Exception:
            pass

    return meta


def _ocr_pdf_pages(file_path: Path) -> str:
    """OCR wszystkich stron PDF przez Tesseract."""
    if not (pytesseract and convert_from_path):
        return ""
    logger.info("[OCR] Analiza wizualna pliku: %s...", file_path.name)
    pages = convert_from_path(file_path, dpi=200)
    ocr_text = []
    for page_img in pages:
        page_text = pytesseract.image_to_string(page_img, lang=OCR_LANG)
        ocr_text.append(page_text)
    text = "\n\n".join(ocr_text)
    logger.info("[OCR] Zakończono dla: %s (Pobrano znaków: %d)", file_path.name, len(text))
    return text


def _ocr_image_file(file_path: Path) -> str:
    """OCR pojedynczego pliku graficznego."""
    if not pytesseract:
        return ""
    try:
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang=OCR_LANG)
        logger.info("[OCR] Przetworzono obraz: %s", file_path.name)
        return text
    except ImportError:
        pass
    if convert_from_path:
        try:
            pages = convert_from_path(file_path, dpi=200)
            ocr_text = []
            for page_img in pages:
                page_text = pytesseract.image_to_string(page_img, lang=OCR_LANG)
                ocr_text.append(page_text)
            text = "\n\n".join(ocr_text)
            logger.info("[OCR] Przetworzono obraz przez pdf2image: %s", file_path.name)
            return text
        except Exception as ocr_err:
            logger.warning("Błąd OCR obrazu %s: %s", file_path.name, ocr_err)
    return ""


def extract_text(file_path: Path, force_ocr: bool = False) -> str:
    ext = file_path.suffix.lower()
    try:
        if ext in ['.md', '.txt']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif ext == '.json':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return json.dumps(json.load(f), indent=2, ensure_ascii=False)
        elif ext == '.docx' and docx:
            return "\n".join([p.text for p in docx.Document(file_path).paragraphs])
        
        elif ext == '.pdf':
            if force_ocr:
                try:
                    return _ocr_pdf_pages(file_path)
                except Exception as ocr_err:
                    logger.warning("Błąd wymuszonego OCR dla %s: %s", file_path.name, ocr_err)
                    return ""

            text = ""
            # Próba 1: Zwykłe wyciągnięcie tekstu cyfrowego
            if pdfplumber:
                with pdfplumber.open(file_path) as pdf:
                    text = "\n".join([page.extract_text() or "" for page in pdf.pages])

            # Próba 2 (Fallback OCR): Jeżeli plik to skan (brak tekstu), uruchom Tesseract
            if len(text.strip()) < 20 and pytesseract and convert_from_path:
                try:
                    text = _ocr_pdf_pages(file_path)
                except Exception as ocr_err:
                    logger.warning("Błąd OCR dla pliku %s: %s", file_path.name, ocr_err)

            return text

        # Obsługa obrazów bezpośrednio (jpg, png, tiff itp.) przez OCR
        elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp']:
            return _ocr_image_file(file_path)

        elif ext in ['.xlsx', '.xls'] and openpyxl:
            return _extract_excel(file_path)
        elif ext == '.csv':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    except Exception as e:
        logger.warning(f"Błąd parsowania pliku {file_path.name}: {e}")
    return ""

@app.route('/')
def index():
    return render_template("index.html", api_key_required=bool(APP_API_KEY))

@app.route('/stats', methods=['GET'])
def get_stats():
    try:
        client = get_qdrant_client()
        info = client.get_collection(collection_name=ACTIVE_COLLECTION)
        return jsonify({
            "success": True,
            "vectors_count": getattr(info, 'vectors_count', None) or info.points_count,
            "points_count": info.points_count,
            "status": info.status,
            "active_collection": ACTIVE_COLLECTION
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/stats/storage', methods=['GET'])
def stats_storage():
    global ACTIVE_COLLECTION
    try:
        client = get_qdrant_client()
        cols = client.get_collections().collections
        result = []
        total_vectors = 0
        total_points  = 0

        for c in cols:
            info = client.get_collection(c.name)
            pts   = info.points_count or 0
            vsize = info.config.params.vectors.size if info.config.params.vectors else 768
            # Szacunek: wektory (float32) + payload (avg 900 B/chunk) + indeks HNSW (~20%)
            vec_bytes     = pts * vsize * 4
            payload_bytes = pts * 900
            index_bytes   = int(vec_bytes * 0.20)
            total_bytes   = vec_bytes + payload_bytes + index_bytes

            result.append({
                "name":    c.name,
                "points":  pts,
                "indexed": info.indexed_vectors_count or 0,
                "status":  str(info.status),
                "vector_size": vsize,
                "distance": str(info.config.params.vectors.distance) if info.config.params.vectors else "Cosine",
                "est_mb":  round(total_bytes / 1_048_576, 2),
                "active":  c.name == ACTIVE_COLLECTION
            })
            total_vectors += pts
            total_points  += pts

        # Free tier Qdrant Cloud: ~4 GB disk / 1 GB RAM
        FREE_DISK_MB = 4096
        used_mb = sum(r["est_mb"] for r in result)
        return jsonify({
            "success": True,
            "collections": result,
            "total_points": total_points,
            "used_mb": round(used_mb, 2),
            "free_disk_mb": FREE_DISK_MB,
            "used_pct": round(used_mb / FREE_DISK_MB * 100, 1),
            "active_collection": ACTIVE_COLLECTION,
            "max_collections": 1000
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/collections/create', methods=['POST'])
def create_collection():
    global ACTIVE_COLLECTION, _qdrant_client
    data = request.get_json()
    name     = data.get('name', '').strip().replace(' ', '_')
    vec_size = int(data.get('vector_size', 768))
    distance = data.get('distance', 'Cosine').capitalize()
    switch   = data.get('switch_to', False)

    if not name:
        return jsonify({"success": False, "error": "Brak nazwy kolekcji"})
    if not name.replace('_','').replace('-','').isalnum():
        return jsonify({"success": False, "error": "Nazwa może zawierać tylko litery, cyfry, _ i -"})

    try:
        from qdrant_client.models import VectorParams, Distance
        dist_map = {"Cosine": Distance.COSINE, "Euclid": Distance.EUCLID, "Dot": Distance.DOT}
        dist = dist_map.get(distance, Distance.COSINE)

        client = get_qdrant_client()
        if client.collection_exists(name):
            return jsonify({"success": False, "error": f"Kolekcja '{name}' już istnieje"})

        client.create_collection(name, vectors_config=VectorParams(size=vec_size, distance=dist))

        if switch:
            ACTIVE_COLLECTION = name
            _persist_active_collection(name)
            _qdrant_client = None  # Reset połączenia po zmianie kolekcji
            _suggestions_cache["data"] = None; _docs_cache["data"] = None

        return jsonify({"success": True, "name": name, "switched": switch, "active": ACTIVE_COLLECTION})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/collections/switch', methods=['POST'])
def switch_collection():
    global ACTIVE_COLLECTION, _qdrant_client
    data = request.get_json()
    name = data.get('name', '').strip()
    if not name:
        return jsonify({"success": False, "error": "Brak nazwy"})
    try:
        client = get_qdrant_client()
        if not client.collection_exists(name):
            return jsonify({"success": False, "error": f"Kolekcja '{name}' nie istnieje"})
        ACTIVE_COLLECTION = name
        _persist_active_collection(name)
        _qdrant_client = None  # Reset połączenia po zmianie kolekcji
        _suggestions_cache["data"] = None; _docs_cache["data"] = None
        return jsonify({"success": True, "active_collection": ACTIVE_COLLECTION})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/collections/delete', methods=['POST'])
def delete_collection():
    global ACTIVE_COLLECTION, _qdrant_client
    data = request.get_json()
    name = data.get('name', '').strip()
    if not name:
        return jsonify({"success": False, "error": "Brak nazwy"})
    if name == ACTIVE_COLLECTION:
        return jsonify({"success": False, "error": "Nie można usunąć aktywnej kolekcji. Najpierw przełącz się na inną."})
    try:
        client = get_qdrant_client()
        client.delete_collection(name)
        _qdrant_client = None  # Reset połączenia po usunięciu kolekcji
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/browse', methods=['GET'])
def browse():
    raw = request.args.get('path', '').strip()
    show_all = request.args.get('all', '0') == '1'
    ext_filter = [e.lower().lstrip('.') for e in request.args.getlist('ext') if e.strip()]
    modified_after = request.args.get('modified_after', '').strip()
    modified_before = request.args.get('modified_before', '').strip()
    default_doc_exts = ['docx', 'pdf', 'xlsx', 'xls', 'csv', 'md', 'json', 'txt']
    allowed_exts = ext_filter if ext_filter else default_doc_exts

    if raw:
        p = Path(_normalize_browse_path(raw)).expanduser()
    else:
        p = _default_browse_path()

    if not p.exists() or not p.is_dir():
        parent = p.parent
        if parent.exists() and parent.is_dir() and _path_is_browsable(parent):
            p = parent
        else:
            p = _default_browse_path()

    try:
        p = p.resolve()
    except OSError:
        p = _default_browse_path()

    if not _path_is_browsable(p):
        roots_hint = ", ".join(str(r) for r in _resolve_allowed_roots()[:5])
        return jsonify({
            "success": False,
            "error": f"Ścieżka poza dozwolonymi katalogami (SEARCH_ROOTS). Dozwolone: {roots_hint}",
            "allowed_roots": [str(r) for r in _resolve_allowed_roots()],
        })

    try:
        entries = []
        total_children = 0
        permission_denied = 0

        for child in sorted(p.iterdir(), key=lambda c: (not c.is_dir(), c.name.lower())):
            total_children += 1
            try:
                if child.is_dir():
                    if not _path_is_browsable(child):
                        continue
                    entries.append({
                        "name": child.name,
                        "path": str(child),
                        "type": "dir",
                        "importable": _path_is_allowed(child),
                    })
                else:
                    ext = child.suffix.lower().lstrip('.')
                    is_doc = ext in allowed_exts
                    if show_all or is_doc:
                        if not _mtime_in_range(child, modified_after or None, modified_before or None):
                            continue
                        try:
                            st = child.stat()
                            mtime = time.strftime("%Y-%m-%d %H:%M", time.localtime(st.st_mtime))
                            size_kb = round(st.st_size / 1024, 1)
                        except OSError:
                            mtime, size_kb = "", 0
                        entries.append({
                            "name": child.name,
                            "path": str(child),
                            "type": "file",
                            "ext": ext,
                            "modified": mtime,
                            "size_kb": size_kb,
                        })
            except PermissionError:
                permission_denied += 1

        is_empty = (total_children == 0) or (len(entries) == 0 and permission_denied == 0)
        has_hidden_or_other = (total_children > 0) and (len(entries) == 0) and (permission_denied == 0)

        return jsonify({
            "success": True,
            "current": str(p),
            "win_path": wsl_to_win(str(p)),
            "parent": str(p.parent) if p != p.parent and _path_is_browsable(p.parent) else None,
            "entries": entries,
            "is_empty": is_empty,
            "total_children": total_children,
            "permission_denied": permission_denied,
            "has_unlisted_items": has_hidden_or_other or permission_denied > 0,
            "show_all": show_all,
            "allowed_roots": [str(r) for r in _resolve_allowed_roots()],
            "filters": {
                "ext": allowed_exts if not show_all else [],
                "modified_after": modified_after or None,
                "modified_before": modified_before or None,
            },
        })
    except PermissionError:
        return jsonify({"success": False, "error": f"Brak uprawnień do odczytu: {p}"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/browse/pick', methods=['POST'])
def browse_pick_folder():
    """Natywny dialog Windows — zwraca wybrany folder jako ścieżkę WSL."""
    data = request.get_json() or {}
    initial = (data.get("initial") or "").strip()
    try:
        if platform.system() == "Windows":
            return jsonify({
                "success": False,
                "error": "Dialog Windows działa z WSL — uruchom aplikację w WSL i otwórz UI w przeglądarce Windows.",
            }), 400
        wsl_path = _pick_windows_folder(initial)
        if not wsl_path:
            return jsonify({"success": False, "error": "Anulowano lub nie wybrano folderu"})
        p = Path(wsl_path).expanduser()
        try:
            p = p.resolve()
        except OSError:
            pass
        if not p.is_dir():
            return jsonify({"success": False, "error": f"Folder nie istnieje: {wsl_path}"})
        if not _path_is_allowed(p):
            roots_hint = ", ".join(str(r) for r in _resolve_allowed_roots()[:5])
            return jsonify({
                "success": False,
                "error": f"Folder poza SEARCH_ROOTS. Dozwolone: {roots_hint}",
                "allowed_roots": [str(r) for r in _resolve_allowed_roots()],
            })
        return jsonify({
            "success": True,
            "path": str(p),
            "win_path": wsl_to_win(str(p)),
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


# ============================================================
# AUTOMATYCZNE MAPOWANIE DYSKÓW LOKALNYCH (dla zakładki Import)
# ============================================================

def _discover_local_drives():
    """Punkty startowe z SEARCH_ROOTS — tylko ścieżki, które da się realnie przeglądać."""
    return _discover_browse_roots()


@app.route('/api/drives')
def api_drives():
    try:
        drives = _discover_local_drives()
        default_path = str(_default_browse_path())
        for drive in drives:
            drive["win_path"] = wsl_to_win(drive.get("path", ""))
        return jsonify({
            "success": True,
            "platform": platform.system(),
            "drives": drives,
            "default_path": default_path,
            "default_win_path": wsl_to_win(default_path),
            "allowed_roots": [str(r) for r in _resolve_allowed_roots()],
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/tasks', methods=['GET'])
def tasks_status():
    current, history = _tasks_snapshot()
    return jsonify({
        "success": True,
        "busy": current is not None,
        "current": current,
        "history": history,
    })


@app.route('/import/stream')
def import_stream():
    folder = request.args.get('folder', '').strip()
    if folder:
        folder = _normalize_browse_path(folder)
    exts = request.args.getlist('ext') or ['docx','pdf','xlsx','xls','csv','md','json']
    force_ocr = request.args.get('force_ocr', '0').strip().lower() in ('1', 'true', 'yes', 'on')

    def generate():
        def sse(event, data):
            import json as _json
            return f"event: {event}\ndata: {_json.dumps(data, ensure_ascii=False)}\n\n"

        task_id = ""
        try:
            folder_path = Path(folder).expanduser()
            if not folder or not folder_path.exists():
                yield sse("error", {"msg": f"Ścieżka nie istnieje: {folder}"})
                return
            if not folder_path.is_dir() or not _path_is_allowed(folder_path):
                yield sse("error", {"msg": "Ścieżka niedozwolona lub poza SEARCH_ROOTS"})
                return

            files = _find_files_safe(folder_path.resolve(), exts)
            if not files:
                yield sse("done", {"count": 0, "chunks": 0, "skipped": 0, "msg": "Brak kompatybilnych plików."})
                return

            ok, task = _start_heavy_task("document_import", f"Import dokumentów: {folder_path}", {
                "folder": str(folder_path),
                "ext": list(exts),
                "force_ocr": force_ocr,
            })
            if not ok:
                yield sse("error", _heavy_task_busy_payload(task))
                return
            task_id = task["id"]
            _update_heavy_task(task_id, total=len(files))

            yield sse("start", {"total": len(files), "force_ocr": force_ocr, "task_id": task_id})

            qdrant = get_qdrant_client()
            imported = 0
            skipped = 0
            total_chunks = 0
            new_chunks = 0
            ocr_count = 0

            for i, f_path in enumerate(files):
                _update_heavy_task(
                    task_id,
                    done=i,
                    total=len(files),
                    progress_pct=round(i / len(files) * 100) if files else 0,
                    current_item=str(f_path),
                )
                try:
                    ext_l = f_path.suffix.lower()
                    used_ocr = (force_ocr and ext_l == '.pdf') or ext_l in {
                        '.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp'
                    }
                    text = extract_text(f_path, force_ocr=force_ocr)
                    if used_ocr and text and len(text.strip()) >= 10:
                        ocr_count += 1
                    if not text or len(text.strip()) < 10:
                        skipped += 1
                        yield sse("skip", {"file": f_path.name, "reason": "pusty", "i": i+1, "total": len(files)})
                        continue

                    chunks = make_chunks(text)
                    file_new = 0

                    # Metadane pobieramy TYLKO RAZ na plik (nie w pętli batchy — ogromna oszczędność IO)
                    file_meta = extract_file_metadata(f_path)

                    # Deduplikacja — sprawdź które chunki już są w bazie
                    chunk_ids = [hashlib.md5(c.encode('utf-8', errors='replace')).hexdigest() for c in chunks]
                    existing_ids = set()
                    for batch_start in range(0, len(chunk_ids), 100):
                        batch = chunk_ids[batch_start:batch_start+100]
                        found = qdrant.retrieve(collection_name=ACTIVE_COLLECTION, ids=batch,
                                               with_payload=False, with_vectors=False)
                        existing_ids.update(p.id for p in found)

                    new_chunks_data = [(cid, chunk) for cid, chunk in zip(chunk_ids, chunks)
                                       if cid not in existing_ids]
                    total_chunks += len(chunks)

                    # Batch embeddings — 6 równolegle (zgodne z get_embeddings_batch default, żeby nie obciążać Ollamy)
                    BATCH = 6
                    embed_failed = False
                    for b in range(0, len(new_chunks_data), BATCH):
                        batch_items = new_chunks_data[b:b+BATCH]
                        batch_texts  = [item[1] for item in batch_items]
                        batch_ids    = [item[0] for item in batch_items]
                        try:
                            vectors = get_embeddings_batch(batch_texts, batch_size=BATCH)
                        except EmbeddingError as e:
                            skipped += 1
                            embed_failed = True
                            yield sse("skip", {"file": f_path.name, "reason": str(e)[:80], "i": i+1, "total": len(files)})
                            break

                        points  = [
                            PointStruct(
                                id=cid,
                                vector=vec,
                                payload={
                                    "file": f_path.name,
                                    "text": txt,
                                    "full_path": str(f_path),
                                    "metadata": file_meta
                                }
                            )
                            for cid, vec, txt in zip(batch_ids, vectors, batch_texts)
                            if vec and any(v != 0.0 for v in vec)
                        ]
                        if points:
                            qdrant.upsert(collection_name=ACTIVE_COLLECTION, points=points)
                        new_chunks += len(points)
                        file_new  += len(points)

                    if embed_failed:
                        continue

                    imported += 1
                    _update_heavy_task(
                        task_id,
                        done=i + 1,
                        total=len(files),
                        progress_pct=round((i + 1) / len(files) * 100) if files else 100,
                        current_item=str(f_path),
                    )
                    yield sse("file", {
                        "file": f_path.name,
                        "chunks": len(chunks),
                        "new": file_new,
                        "i": i+1,
                        "total": len(files),
                        "ocr": used_ocr,
                    })
                    time.sleep(0.02)
                except Exception as e:
                    skipped += 1
                    yield sse("skip", {"file": f_path.name, "reason": str(e)[:80], "i": i+1, "total": len(files)})

            ocr_note = f" · OCR: {ocr_count} plików" if ocr_count else ""
            _update_heavy_task(task_id, done=len(files), total=len(files), progress_pct=100, current_item=None)
            yield sse("done", {
                "count": imported,
                "chunks": total_chunks,
                "new_chunks": new_chunks,
                "skipped": skipped,
                "ocr_count": ocr_count,
                "task_id": task_id,
                "msg": f"Przetworzono {imported} plików · {new_chunks} nowych chunków · {total_chunks - new_chunks} duplikatów pominiętych{ocr_note}"
            })
        except Exception as e:
            logger.exception("import_stream error")
            yield sse("error", {"msg": str(e), "task_id": task_id or None})
            if task_id:
                _finish_heavy_task(task_id, status="error", error=str(e)[:500])
                task_id = ""
        finally:
            if task_id:
                _finish_heavy_task(task_id)

    return Response(stream_with_context(generate()), mimetype='text/event-stream',
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

@app.route('/collection/clear', methods=['POST'])
def clear_collection():
    try:
        client = get_qdrant_client()
        info = client.get_collection(ACTIVE_COLLECTION)
        count_before = info.points_count
        # Pobierz aktualny wymiar wektora zanim usuniesz kolekcję
        vsize = info.config.params.vectors.size if info.config.params.vectors else 768
        from qdrant_client.models import VectorParams, Distance
        client.delete_collection(ACTIVE_COLLECTION)
        client.create_collection(
            ACTIVE_COLLECTION,
            vectors_config=VectorParams(size=vsize, distance=Distance.COSINE)
        )
        _suggestions_cache["data"] = None; _docs_cache["data"] = None
        return jsonify({"success": True, "deleted": count_before})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/import', methods=['POST'])
def import_folder():
    data = request.get_json()
    return jsonify({"success": False, "error": "Użyj /import/stream (SSE)"})

def wsl_to_win(path: str) -> str:
    """Konwertuje ścieżkę WSL /mnt/c/... na Windows C:\\..."""
    if not path:
        return ""
    m = re.match(r'^/mnt/([a-zA-Z])(?:/(.*))?$', path)
    if m:
        drive = m.group(1).upper()
        rest = (m.group(2) or "").replace('/', '\\')
        return f"{drive}:\\{rest}" if rest else f"{drive}:\\"
    return path


def open_file_safely(win_path: str) -> bool:
    """
    Bezpiecznie otwiera plik za pomocą domyślnej aplikacji systemowej.
    Unika shell injection (w przeciwieństwie do poprzedniego os.popen).
    """
    if not win_path:
        return False

    try:
        if platform.system() == "Windows":
            # Najbezpieczniejsza metoda na Windows
            os.startfile(win_path)
            logger.info(f"Otwarto plik: {win_path}")
            return True
        else:
            # Fallback dla Linux/macOS (przydatne przy testach z WSL)
            subprocess.run(["xdg-open", win_path], check=False)
            logger.info(f"Otwarto plik (xdg-open): {win_path}")
            return True
    except Exception as e:
        logger.warning(f"Nie udało się otworzyć pliku '{win_path}': {e}")
        return False

@app.route('/file/open', methods=['POST'])
def file_open():
    data = request.get_json()
    wsl_path = data.get('path', '').strip()
    if not wsl_path:
        return jsonify({"success": False, "error": "Brak ścieżki"})

    try:
        check_path = Path(wsl_path).expanduser().resolve()
    except OSError:
        return jsonify({"success": False, "error": "Nieprawidłowa ścieżka"})
    if not check_path.is_file() or not _path_is_allowed(check_path):
        return jsonify({"success": False, "error": "Plik niedostępny lub poza dozwolonymi katalogami"})

    win_path = wsl_to_win(wsl_path)
    if not win_path:
        return jsonify({"success": False, "error": "Nie można skonwertować ścieżki"})

    success = open_file_safely(win_path)

    if success:
        return jsonify({"success": True, "win_path": win_path})
    else:
        return jsonify({"success": False, "error": "Nie udało się otworzyć pliku"})

@app.route('/hybrid/stream', methods=['POST'])
def hybrid_stream():
    """Wyszukiwanie hybrydowe: RAG + SQL równolegle — SSE."""
    data       = request.get_json()
    query_text = data.get('query', '').strip()
    chat_context = (data.get('chat_context') or '').strip()
    conn_cfg   = data.get('conn', {})
    schema_str = data.get('schema', '')
    limit      = min(int(data.get('limit', 5)), 20)
    mode       = data.get('mode', 'normal')
    if mode not in SEARCH_MODES:
        mode = 'normal'
    file_filter= data.get('file_filter', None)
    llm_provider = data.get('llm_provider')
    openrouter_model = data.get('openrouter_model')

    if not query_text:
        return jsonify({"success": False, "error": "Zapytanie puste"}), 400

    def generate():
        def sse(event, d):
            return f"event: {event}\ndata: {json.dumps(d, ensure_ascii=False)}\n\n"

        try:
            # 1. RAG — Qdrant query
            client = get_qdrant_client()
            try:
                rag_contexts, rag_results = _retrieve_search_contexts(
                    client, query_text, limit, file_filter, mode
                )
            except EmbeddingError as e:
                yield sse("error", {"error": str(e)})
                return

            yield sse("rag_results", {"results": rag_results, "contexts": rag_contexts})

            # 2. SQL — jeśli konfiguracja dostępna (z pełnymi guardami bezpieczeństwa)
            sql_data = {"success": False, "table": "", "columns": [], "rows": [], "sql": ""}
            sql_query = ""

            if conn_cfg and _sql_conn_configured(conn_cfg):
                try:
                    _require_sql_backend(conn_cfg)
                except RuntimeError as e:
                    sql_data["error"] = str(e)
                else:
                    try:
                        result = _run_text_to_sql_pipeline(
                            conn_cfg, query_text, schema_str,
                            fetch_limit=200, with_chart=False,
                        )
                        if result.get("success"):
                            sql_query = result.get("sql", "")
                            sql_data = {
                                "success": True,
                                "table": "wynik SQL",
                                "columns": result.get("columns", []),
                                "rows": result.get("rows", []),
                                "sql": sql_query[:120],
                                "total": result.get("total", 0),
                                "sql_corrections": result.get("sql_corrections", 0),
                            }
                        else:
                            sql_data["error"] = (result.get("error") or "Błąd SQL")[:200]
                            if result.get("sql"):
                                sql_data["sql"] = result["sql"][:120]
                    except Exception as e:
                        sql_data["error"] = str(e)[:200]

            yield sse("sql_results", {"sql_results": sql_data})

            # 3. LLM synteza — łączy RAG + SQL
            if not rag_contexts:
                yield sse("done", {"ai_answer": "Brak dokumentów w bazie RAG."})
                return

            rag_cap = 10 if mode == "detective" else 5
            rag_slice = rag_contexts[:rag_cap]
            prompt, system = _build_search_prompt(
                query_text, rag_slice, mode, chat_context=chat_context
            )
            if sql_data.get("success") and sql_data.get("rows"):
                rows_str = "\n".join([
                    " | ".join([f"{k}={v}" for k, v in zip(sql_data["columns"],
                             [r.get(col, "") for col in sql_data["columns"]])])
                    for r in sql_data["rows"][:10]
                ])
                prompt += f"\n\nDANE Z BAZY SQL (uzupełnienie — porównaj z dokumentami):\n{rows_str}\n"
                if mode == "detective":
                    prompt += (
                        "\nUwzględnij w sekcji Analiza śledcza rozbieżności między SQL a dokumentami, "
                        "jeśli występują.\n"
                    )
                else:
                    prompt += "Podaj: (1) co wynika z bazy danych, (2) co potwierdzają dokumenty, (3) wnioski.\n"

            # Streaming odpowiedzi LLM (obsługuje zarówno Ollama jak i OpenRouter)
            full_answer = ""
            llm_meta = {}

            try:
                for token in stream_llm_tokens(
                    prompt=prompt,
                    system=system,
                    provider=llm_provider,
                    model=openrouter_model,
                    meta=llm_meta,
                ):
                    if isinstance(token, str) and (token.startswith("[Błąd") or token.startswith("[RATE_LIMIT]") or token.startswith("[FALLBACK]")):
                        logger.warning(f"LLM stream error (hybrid): {token[:120]}")
                        yield sse("error", {"error": token, "provider": get_llm_provider(llm_provider)})
                        return
                    full_answer += token
                    yield sse("token", {"token": token})
            except Exception as e:
                logger.warning(f"LLM stream error: {e}")
                yield sse("error", {"error": f"Błąd streamingu LLM: {str(e)}"})

            done_payload = {"ai_answer": full_answer}
            if llm_meta:
                done_payload["llm_provider_used"] = llm_meta.get("provider_used")
                done_payload["llm_model_used"] = llm_meta.get("model_used")
                done_payload["llm_fallback_used"] = llm_meta.get("fallback_used", False)
                if llm_meta.get("fallback_reason"):
                    done_payload["llm_fallback_reason"] = llm_meta.get("fallback_reason")
            yield sse("done", done_payload)

        except Exception as e:
            yield sse("error", {"error": str(e)})

    return Response(stream_with_context(generate()), mimetype='text/event-stream',
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

@app.route('/search/stream', methods=['POST'])
def search_stream():
    """Wyszukiwanie z streamingiem LLM — wyniki natychmiast, odpowiedź słowo po słowie."""
    data       = request.get_json()
    query_text = data.get('query', '').strip()
    if not query_text:
        return jsonify({"success": False, "error": "Zapytanie puste"})
    chat_context = (data.get('chat_context') or '').strip()
    limit      = min(int(data.get('limit', 5)), 20)
    file_filter = data.get('file_filter', None)
    mode        = data.get('mode', 'normal')
    llm_provider = data.get('llm_provider')
    if mode not in SEARCH_MODES:
        mode = 'normal'

    def generate():
        def sse(event, d):
            return f"event: {event}\ndata: {json.dumps(d, ensure_ascii=False)}\n\n"

        try:
            client = get_qdrant_client()
            try:
                raw_contexts, results = _retrieve_search_contexts(
                    client, query_text, limit, file_filter, mode
                )
            except EmbeddingError as e:
                yield sse("error", {"error": str(e)})
                return

            # Wyślij wyniki od razu
            yield sse("results", {"results": results, "contexts": raw_contexts,
                                  "mode": mode, "mode_label": SEARCH_MODES[mode]["label"]})

            if not raw_contexts:
                yield sse("done", {"ai_answer": "Brak dokumentów."})
                return

            prompt, system = _build_search_prompt(
                query_text, raw_contexts, mode, chat_context=chat_context
            )

            # Streaming LLM — obsługuje Ollama i OpenRouter
            full_answer = ""
            openrouter_model = data.get("openrouter_model")
            llm_meta = {}

            try:
                for token in stream_llm_tokens(
                    prompt=prompt,
                    system=system,
                    provider=llm_provider,
                    model=openrouter_model,
                    meta=llm_meta,
                ):
                    if isinstance(token, str) and (token.startswith("[Błąd") or token.startswith("[RATE_LIMIT]") or token.startswith("[FALLBACK]")):
                        # Czysty błąd zamiast zaśmiecania odpowiedzi
                        logger.warning(f"LLM stream error (search): {token[:120]}")
                        yield sse("error", {"error": token, "provider": get_llm_provider(llm_provider)})
                        return
                    full_answer += token
                    yield sse("token", {"token": token})
            except Exception as e:
                logger.warning(f"LLM stream error w search_stream: {e}")
                yield sse("error", {"error": "Błąd streamingu modelu językowego."})

            done_payload = {"ai_answer": full_answer}
            if llm_meta:
                done_payload["llm_provider_used"] = llm_meta.get("provider_used")
                done_payload["llm_model_used"] = llm_meta.get("model_used")
                done_payload["llm_fallback_used"] = llm_meta.get("fallback_used", False)
                if llm_meta.get("fallback_reason"):
                    done_payload["llm_fallback_reason"] = llm_meta.get("fallback_reason")
            yield sse("done", done_payload)

        except Exception as e:
            yield sse("error", {"error": str(e)})

    return Response(stream_with_context(generate()), mimetype='text/event-stream',
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

@app.route('/search', methods=['POST'])
def search():
    data = request.get_json()
    query_text = data.get('query', '').strip()
    if not query_text: return jsonify({"success": False, "error": "Zapytanie puste"})
    chat_context = (data.get('chat_context') or '').strip()
    limit = min(int(data.get('limit', 5)), 20)
    file_filter = data.get('file_filter', None)
    mode = data.get('mode', 'normal')
    llm_provider = data.get('llm_provider')  # 'ollama' lub 'openrouter' z UI
    if mode not in SEARCH_MODES:
        mode = 'normal'

    try:
        client = get_qdrant_client()
        try:
            raw_contexts, results = _retrieve_search_contexts(
                client, query_text, limit, file_filter, mode
            )
        except EmbeddingError as e:
            return jsonify({"success": False, "error": str(e)})

        if file_filter and not raw_contexts:
            return jsonify({"success": True, "results": [], "ai_answer": "Brak dokumentów dla wybranego pliku."})

        openrouter_model = data.get("openrouter_model")
        ai_answer = (
            generate_answer(
                query_text,
                raw_contexts,
                mode,
                provider=llm_provider,
                model=openrouter_model,
                chat_context=chat_context,
            )
            if raw_contexts
            else "Brak dokumentów."
        )
        return jsonify({
            "success": True,
            "results": results,
            "ai_answer": ai_answer,
            "mode": mode,
            "mode_label": SEARCH_MODES[mode]["label"],
            "contexts": raw_contexts  # potrzebne do weryfikacji po stronie klienta
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/analyze', methods=['POST'])
def analyze():
    """Gotowe raporty śledcze — uproszczony POST zwracający pole answer."""
    data = request.get_json() or {}
    query_text = (data.get('query') or '').strip()
    if not query_text:
        return jsonify({"success": False, "error": "Zapytanie puste"})
    chat_context = (data.get('chat_context') or '').strip()
    limit = min(int(data.get('limit', 10)), 20)
    file_filter = data.get('file_filter')
    mode = data.get('mode', 'normal')
    llm_provider = data.get('llm_provider')
    openrouter_model = data.get('openrouter_model')
    if mode not in SEARCH_MODES:
        mode = 'normal'

    try:
        client = get_qdrant_client()
        try:
            raw_contexts, results = _retrieve_search_contexts(
                client, query_text, limit, file_filter, mode
            )
        except EmbeddingError as e:
            return jsonify({"success": False, "error": str(e)})

        answer = (
            generate_answer(
                query_text,
                raw_contexts,
                mode,
                provider=llm_provider,
                model=openrouter_model,
                chat_context=chat_context,
            )
            if raw_contexts
            else "Brak dokumentów w bazie."
        )
        return jsonify({
            "success": True,
            "answer": answer,
            "ai_answer": answer,
            "results": results,
            "mode": mode,
            "mode_label": SEARCH_MODES[mode]["label"],
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


GEMINI_AUDIT_SYSTEM = (
    "Jesteś elitarnym ekspertem ds. cyberbezpieczeństwa, DevOps i analizy systemowej. "
    "Analizujesz logi, dumpy i duże pliki tekstowe. "
    "Znajdź anomalie, błędy (ERROR, CRITICAL, FATAL) oraz potencjalne zagrożenia. "
    "Przygotuj raport po polsku w Markdown: Podsumowanie, Znalezione błędy, Sugerowane akcje naprawcze. "
    "Opieraj się wyłącznie na danych z pliku — nie zmyślaj."
)


@app.route('/audit/info', methods=['GET'])
def audit_info():
    """Status konfiguracji audytu Gemini."""
    return jsonify({
        "success": True,
        "gemini_configured": bool(GEMINI_API_KEY),
        "default_model": GEMINI_MODEL,
        "max_bytes": GEMINI_AUDIT_MAX_BYTES,
        "allowed_extensions": sorted(GEMINI_AUDIT_EXTENSIONS),
    })


@app.route('/audit/file/stream', methods=['POST'])
def audit_file_stream():
    """SSE: audyt dużego pliku tekstowego przez Gemini (cały plik w kontekście)."""
    data = request.get_json() or {}
    file_raw = (data.get("file") or data.get("path") or "").strip()
    query = (data.get("query") or data.get("prompt") or "").strip()
    model = (data.get("model") or GEMINI_MODEL).strip()
    redact = data.get("redact_secrets", True)

    if not GEMINI_API_KEY:
        def _err_gen():
            yield f"event: error\ndata: {json.dumps({'error': 'Brak GEMINI_API_KEY w .env — dodaj klucz z Google AI Studio.'}, ensure_ascii=False)}\n\n"
        return Response(stream_with_context(_err_gen()), mimetype='text/event-stream',
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    if not file_raw:
        def _err_gen():
            yield f"event: error\ndata: {json.dumps({'error': 'Podaj ścieżkę pliku.'}, ensure_ascii=False)}\n\n"
        return Response(stream_with_context(_err_gen()), mimetype='text/event-stream',
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    if not query:
        def _err_gen():
            yield f"event: error\ndata: {json.dumps({'error': 'Podaj pytanie / instrukcję analizy.'}, ensure_ascii=False)}\n\n"
        return Response(stream_with_context(_err_gen()), mimetype='text/event-stream',
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    def generate():
        def sse(event, payload):
            return f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n"

        try:
            path = _resolve_audit_file_path(file_raw)
            if not path:
                yield sse("error", {
                    "error": (
                        "Plik niedostępny, poza SEARCH_ROOTS lub niedozwolone rozszerzenie "
                        f"({', '.join(sorted(GEMINI_AUDIT_EXTENSIONS))})."
                    ),
                })
                return

            size = path.stat().st_size
            if size > GEMINI_AUDIT_MAX_BYTES:
                yield sse("error", {
                    "error": (
                        f"Plik za duży ({size:,} B). Limit: {GEMINI_AUDIT_MAX_BYTES:,} B "
                        f"(GEMINI_AUDIT_MAX_BYTES)."
                    ),
                })
                return

            yield sse("progress", {
                "message": f"Wczytywanie {path.name}…",
                "bytes": size,
                "file": str(path),
            })

            content = path.read_text(encoding="utf-8", errors="ignore")
            if redact:
                content = _redact_sensitive_log_text(content)

            yield sse("progress", {
                "message": f"Analiza Gemini ({len(content):,} znaków)…",
                "chars": len(content),
                "model": model,
            })

            full_prompt = f"{query}\n\n--- ZAWARTOŚĆ PLIKU ({path.name}) ---\n{content}"
            full_answer = ""

            for token in _stream_gemini_tokens(
                full_prompt,
                system=GEMINI_AUDIT_SYSTEM,
                model=model,
                max_tokens=8192,
                temperature=0.2,
            ):
                if isinstance(token, str) and token.startswith("[Błąd"):
                    yield sse("error", {"error": token})
                    return
                full_answer += token
                yield sse("token", {"token": token})

            yield sse("done", {
                "ai_answer": full_answer,
                "file": str(path),
                "file_name": path.name,
                "model": model,
                "chars_analyzed": len(content),
                "bytes_analyzed": size,
            })
        except Exception as e:
            logger.warning("audit_file_stream error: %s", e)
            yield sse("error", {"error": str(e)[:300]})

    return Response(stream_with_context(generate()), mimetype='text/event-stream',
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route('/compare', methods=['POST'])
def compare_documents():
    """SSE: porównanie dwóch dokumentów — streaming wyników porównania przez LLM."""
    _require_api_key()
    data = request.get_json() or {}
    file_a = (data.get('file_a') or '').strip()
    file_b = (data.get('file_b') or '').strip()
    focus  = (data.get('focus') or 'general').strip()

    if not file_a or not file_b:
        def _err():
            yield sse('error', {'error': 'Wymagane file_a i file_b'})
        return Response(stream_with_context(_err()), mimetype='text/event-stream',
                        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})
    if file_a == file_b:
        def _err():
            yield sse('error', {'error': 'Wybierz dwa różne pliki'})
        return Response(stream_with_context(_err()), mimetype='text/event-stream',
                        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

    def generate():
        def sse(event, d):
            import json as _j
            return f"event: {event}\ndata: {_j.dumps(d, ensure_ascii=False)}\n\n"

        try:
            client = get_qdrant_client()

            def _fetch_chunks(fname: str) -> list[str]:
                from qdrant_client.models import Filter, FieldCondition, MatchValue
                qfilter = Filter(must=[FieldCondition(key="file", match=MatchValue(value=fname))])
                records, _ = client.scroll(
                    collection_name=ACTIVE_COLLECTION,
                    scroll_filter=qfilter,
                    limit=30,
                    with_payload=["text"],
                    with_vectors=False,
                )
                return [r.payload.get("text", "") for r in records if r.payload.get("text")]

            chunks_a = _fetch_chunks(file_a)
            chunks_b = _fetch_chunks(file_b)

            if not chunks_a:
                yield sse('error', {'error': f'Brak chunków dla pliku: {file_a}'})
                return
            if not chunks_b:
                yield sse('error', {'error': f'Brak chunków dla pliku: {file_b}'})
                return

            MAX_CHARS = 8000
            text_a = "\n\n".join(chunks_a)[:MAX_CHARS]
            text_b = "\n\n".join(chunks_b)[:MAX_CHARS]

            focus_labels = {
                'general':    'ogólne porównanie treści i struktury',
                'dates':      'daty, terminy i harmonogram',
                'parties':    'strony, sygnatury i upoważnienia',
                'financial':  'kwoty, wartości i warunki finansowe',
                'legal':      'zobowiązania prawne i klauzule umowne',
                'technical':  'wymagania techniczne i specyfikacje',
            }
            focus_desc = focus_labels.get(focus, focus_labels['general'])

            system = (
                "Jesteś ekspertem ds. analizy dokumentów. Porównaj dwa dokumenty skupiając się na: "
                f"{focus_desc}. "
                "Wskaż kluczowe podobieństwa, różnice i rozbieżności. "
                "Odpowiadaj po polsku, używaj punktów i tabel gdzie to pomocne."
            )
            prompt = (
                f"=== DOKUMENT A: {file_a} ===\n{text_a}\n\n"
                f"=== DOKUMENT B: {file_b} ===\n{text_b}\n\n"
                f"Przeprowadź szczegółowe porównanie tych dwóch dokumentów "
                f"z uwzględnieniem: {focus_desc}."
            )

            yield sse('start', {
                'file_a': file_a,
                'file_b': file_b,
                'chunks_a': len(chunks_a),
                'chunks_b': len(chunks_b),
                'focus': focus_desc,
            })

            comparison_text = ""
            for token in stream_llm_tokens(prompt=prompt, system=system, max_tokens=2000, temperature=0.2):
                comparison_text += token
                yield sse('chunk', {'token': token})

            yield sse('done', {
                'file_a': file_a,
                'file_b': file_b,
                'chunks_a': len(chunks_a),
                'chunks_b': len(chunks_b),
                'comparison': comparison_text,
            })
        except Exception as e:
            logger.exception("compare_documents error")
            yield sse('error', {'error': str(e)[:300]})

    return Response(stream_with_context(generate()), mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


@app.route('/api/collection/profile', methods=['GET'])
def collection_profile():
    """Profiluje kolekcję wg typów plików i sugeruje tryb wyszukiwania."""
    try:
        client = get_qdrant_client()
        records, _ = client.scroll(
            collection_name=ACTIVE_COLLECTION,
            limit=500,
            with_payload=["file"],
            with_vectors=False,
        )
        if not records:
            return jsonify({"success": True, "profile": "empty", "suggestion_text": "", "suggestion_mode": "normal"})

        from collections import Counter
        ext_counts: Counter = Counter()
        for r in records:
            fname = r.payload.get("file", "")
            if fname:
                ext = Path(fname).suffix.lower().lstrip(".")
                if ext:
                    ext_counts[ext] += 1

        total = sum(ext_counts.values()) or 1
        top_ext, top_count = ext_counts.most_common(1)[0] if ext_counts else ("", 0)
        top_ratio = top_count / total

        # Determine profile
        financial_exts = {"xlsx", "xls", "csv", "ods"}
        legal_exts     = {"pdf", "docx", "doc", "odt"}
        code_exts      = {"py", "js", "ts", "java", "cs", "cpp", "go", "rb"}

        if ext_counts.keys() & financial_exts and top_ratio > 0.5 and top_ext in financial_exts:
            profile = "financial"
            suggestion_mode = "normal"
            suggestion_text = f"Kolekcja zawiera głównie pliki finansowe ({top_ext.upper()}). Rozważ tryb Detektyw do głębokiej analizy."
            suggestion_mode = "detective"
        elif ext_counts.keys() & legal_exts and top_ratio > 0.4 and top_ext in legal_exts:
            profile = "legal"
            suggestion_mode = "detective"
            suggestion_text = f"Kolekcja zawiera głównie dokumenty prawne ({top_ext.upper()}). Zalecany tryb: Detektyw."
        elif ext_counts.keys() & code_exts:
            profile = "technical"
            suggestion_mode = "normal"
            suggestion_text = "Kolekcja zawiera kod źródłowy. Tryb normalny jest odpowiedni."
        elif top_ratio > 0.7:
            profile = "homogeneous"
            suggestion_mode = "normal"
            suggestion_text = f"Jednorodna kolekcja ({top_ext.upper()}). Tryb normalny jest odpowiedni."
        else:
            profile = "mixed"
            suggestion_mode = "detective"
            suggestion_text = "Różnorodna kolekcja dokumentów. Tryb Detektyw da lepsze wyniki."

        return jsonify({
            "success": True,
            "profile": profile,
            "suggestion_mode": suggestion_mode,
            "suggestion_text": suggestion_text,
            "ext_counts": dict(ext_counts.most_common(5)),
        })
    except Exception as e:
        logger.warning(f"collection_profile error: {e}")
        return jsonify({"success": False, "error": str(e)[:200]})


# ===== MULTI-AGENT SWARM =====

def _split_chunks_for_swarm(chunks: list, cfg: dict) -> list[list]:
    """Dzieli chunki na grupy workerów wg trybu."""
    import random as _rnd
    n_workers = cfg["workers"]
    max_per = cfg["max_chunks_per_worker"]
    data = list(chunks)
    if cfg.get("shuffle_chunks"):
        _rnd.shuffle(data)
    data = data[:n_workers * max_per]
    n = min(n_workers, len(data))
    groups: list[list] = [[] for _ in range(n)]
    for i, chunk in enumerate(data):
        groups[i % n].append(chunk)
    return [g for g in groups if g]


def _assign_swarm_models(n_workers: int, cfg: dict) -> list[tuple]:
    """Zwraca liste (model_id, model_name) dla kazdego workera.
    Inteligentnie wybiera custom endpointy jesli sa dostepne i maja wysoki score."""
    import random as _rnd
    pool = list(FREE_MODELS_LIST)

    custom_endpoints = _get_best_custom_endpoints_wrapper(limit=2)
    if custom_endpoints:
        pool.extend(custom_endpoints)

    if cfg.get("use_random_models"):
        _rnd.shuffle(pool)
        return [pool[i % len(pool)] for i in range(n_workers)]

    preferred = cfg.get("preferred_model", pool[0] if pool else FREE_MODELS_LIST[0])
    return [preferred] * n_workers


@app.route('/agents/swarm/stream', methods=['POST'])
def agents_swarm_stream():
    """SSE: równoległa analiza dokumentów przez N niezależnych agentów LLM."""
    _require_api_key()
    data = request.get_json() or {}
    query = (data.get('query') or '').strip()
    swarm_mode = data.get('swarm_mode', 'quality')
    file_filter = data.get('file_filter') or None

    def sse(event, d):
        import json as _j
        return f"event: {event}\ndata: {_j.dumps(d, ensure_ascii=False)}\n\n"

    if not query:
        def _err():
            yield sse('error', {'error': 'Zapytanie jest puste'})
        return Response(stream_with_context(_err()), mimetype='text/event-stream',
                        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

    cfg = SWARM_MODES.get(swarm_mode, SWARM_MODES['quality'])

    def generate():
        try:
            yield sse('progress', {'msg': 'Pobieranie fragmentów z bazy wektorowej…', 'step': 1})

            client = get_qdrant_client()
            try:
                raw_contexts, _ = _retrieve_search_contexts(client, query, 20, file_filter, 'detective')
            except EmbeddingError as e:
                yield sse('error', {'error': str(e)})
                return

            if not raw_contexts:
                yield sse('error', {'error': 'Brak dokumentów w bazie dla tego zapytania.'})
                return

            groups = _split_chunks_for_swarm(raw_contexts, cfg)
            n_workers = len(groups)
            models = _assign_swarm_models(n_workers, cfg)

            yield sse('init', {
                'n_workers': n_workers,
                'mode': swarm_mode,
                'mode_label': cfg['label'],
                'mode_desc': cfg['desc'],
                'total_chunks': len(raw_contexts),
            })

            for i, ((model_id, model_name), group) in enumerate(zip(models, groups)):
                yield sse('worker_start', {
                    'worker_id': i,
                    'model': model_name,
                    'chunks': len(group),
                    'files': list({c['file'] for c in group})[:3],
                })

            # Równoległe wywołania LLM (ThreadPoolExecutor)
            def _run_worker(worker_id: int, context_chunks: list, model_id: str, model_name: str):
                try:
                    ctx = "\n\n---\n\n".join(
                        f"[{c['file']}]\n{c['text']}" for c in context_chunks
                    )
                    prompt = (
                        f"Pytanie: {query}\n\n"
                        f"Przeanalizuj poniższe fragmenty dokumentów i odpowiedz konkretnie.\n\n"
                        f"{ctx}\n\n"
                        f"Jeśli w tych fragmentach brakuje odpowiedzi — napisz krótko: "
                        f"'Brak danych w tym fragmencie.'"
                    )
                    result = call_llm(
                        prompt=prompt,
                        system="Jesteś analitykiem dokumentów. Odpowiadaj po polsku, zwięźle i konkretnie.",
                        stream=False,
                        provider='openrouter',
                        model=model_id,
                    )
                    text = _llm_response_text(result)
                    return ('done', worker_id, text, model_name, None)
                except Exception as exc:
                    return ('error', worker_id, None, model_name, str(exc)[:200])

            partial: dict[int, str] = {}
            worker_timeout = 120
            with ThreadPoolExecutor(max_workers=max(1, n_workers)) as ex:
                futures = {
                    ex.submit(_run_worker, i, group, model_id, model_name): i
                    for i, ((model_id, model_name), group) in enumerate(zip(models, groups))
                }
                try:
                    completed_futures = as_completed(futures, timeout=worker_timeout)
                    for fut in completed_futures:
                        wid = futures[fut]
                        try:
                            kind, worker_id, text, model_name, err = fut.result()
                        except Exception as exc:
                            partial[wid] = ''
                            yield sse('worker_error', {
                                'worker_id': wid,
                                'model': models[wid][1] if wid < len(models) else '?',
                                'error': str(exc)[:200],
                            })
                            continue
                        if kind == 'done':
                            partial[worker_id] = text
                            yield sse('worker_done', {
                                'worker_id': worker_id,
                                'model': model_name,
                                'preview': text[:240] + ('…' if len(text) > 240 else ''),
                            })
                        else:
                            partial[worker_id] = ''
                            yield sse('worker_error', {'worker_id': worker_id, 'model': model_name, 'error': err})
                except FuturesTimeoutError:
                    for fut, wid in futures.items():
                        if fut.done():
                            continue
                        fut.cancel()
                        partial[wid] = ''
                        yield sse('worker_error', {
                            'worker_id': wid,
                            'model': models[wid][1] if wid < len(models) else '?',
                            'error': f'Timeout — brak odpowiedzi w {worker_timeout} s',
                        })

            if len(partial) < n_workers:
                yield sse('error', {'error': 'Timeout — workery nie odpowiedziały w czasie 120 s.'})
                return

            # Synteza
            yield sse('synthesis_start', {'msg': 'Synteza wyników przez orkiestratora…'})

            parts = "\n\n===\n\n".join(
                f"Agent {i + 1} ({models[i][1]}):\n{partial.get(i, '[brak]')}"
                for i in range(n_workers)
            )
            if cfg.get('hierarchical'):
                synthesis_prompt = (
                    f"Masz wyniki {n_workers} niezależnych agentów analizujących różne fragmenty dokumentów.\n\n"
                    f"Pytanie użytkownika: {query}\n\n"
                    f"Wyniki agentów:\n{parts}\n\n"
                    f"Zadanie: pomiń powtórzenia, połącz kluczowe ustalenia ze wszystkich agentów, "
                    f"wskaż ewentualne sprzeczności, i podaj finalną, wyczerpującą odpowiedź."
                )
            else:
                synthesis_prompt = (
                    f"Pytanie: {query}\n\n"
                    f"Odpowiedzi cząstkowe agentów:\n{parts}\n\n"
                    f"Podaj zwięzłą, kompletną odpowiedź łącząc wszystkie ustalenia."
                )

            synth_model = cfg['synthesis_model']
            final = call_llm(
                prompt=synthesis_prompt,
                system="Jesteś ekspertem od syntezy informacji z dokumentów. Odpowiadaj po polsku.",
                stream=False,
                provider='openrouter',
                model=synth_model,
            )
            final_text = _llm_response_text(final)
            final_usage = _llm_usage_from_result(final)
            if not final_usage and final_text:
                approx_tokens = max(1, len(final_text.split()))
                final_usage = {"completion_tokens": approx_tokens, "total_tokens": approx_tokens}

            yield sse('done', {
                'answer': final_text,
                'n_workers': n_workers,
                'mode': swarm_mode,
                'mode_label': cfg['label'],
                'usage': final_usage,
            })

        except Exception as e:
            logger.exception("agents_swarm_stream error")
            yield sse('error', {'error': str(e)[:300]})

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'},
    )


@app.route('/agents/swarm/modes', methods=['GET'])
def agents_swarm_modes():
    """Zwraca definicje trybów swarm do UI."""
    return jsonify({
        "success": True,
        "modes": {k: {"label": v["label"], "desc": v["desc"]} for k, v in SWARM_MODES.items()},
    })


# ===== MODEL FLEET =====

@app.route('/api/models/registry', methods=['GET'])
def models_registry_endpoint():
    """Zwraca rejestr modeli + live statystyki + custom endpointy."""
    out = {}
    for mid, info in MODEL_REGISTRY.items():
        live = _model_live.get(mid, {})
        out[mid] = {**info, "live": live, "score": _model_score(mid)}

    pool = _load_provider_pool()
    all_custom_stats = _load_custom_endpoint_stats()
    for endpoint in pool.get("custom_endpoints", []):
        if not endpoint.get("active"):
            continue
        eid = endpoint.get("id")
        stats = all_custom_stats.get(eid, {
            "ping_ok": None, "ping_ms": None, "ping_ts": None,
            "err_count": 0, "ok_count": 0, "avg_ms": None
        })
        out[eid] = {
            "kind": "custom_endpoint",
            "name": endpoint.get("name", endpoint.get("url")),
            "short": endpoint.get("name", "Custom")[:20],
            "provider": "Custom",
            "icon": "🔌",
            "context_k": 128,
            "speed_tier": 2,
            "quality_tier": 2,
            "free": True,
            "rate_rpm": 60,
            "rate_day": 1000,
            "tags": ["custom", "swarm"],
            "url": endpoint.get("url"),
            "live": stats,
            "score": _custom_endpoint_score(eid)
        }

    return jsonify({"success": True, "models": out})


@app.route('/api/llm/fleet', methods=['GET'])
def llm_fleet_endpoint():
    """Kompatybilny endpoint floty LLM używany przez testy i starszy frontend."""
    task = request.args.get("task", "general").strip() or "general"
    models_payload = models_registry_endpoint().get_json() or {}
    models = models_payload.get("models") or {}
    providers = []
    for model_id, info in models.items():
        providers.append({
            "id": model_id,
            "model_id": model_id,
            "name": info.get("name") or info.get("short") or model_id,
            "provider": info.get("provider", ""),
            "score": info.get("score", 0),
            "ping_ms": (info.get("live") or {}).get("avg_ms") or (info.get("live") or {}).get("ping_ms"),
            "ping_ok": (info.get("live") or {}).get("ping_ok"),
            "context_k": info.get("context_k"),
            "rate_rpm": info.get("rate_rpm"),
            "rate_day": info.get("rate_day"),
            "tags": info.get("tags", []),
        })
    providers.sort(key=lambda item: item.get("score") or 0, reverse=True)
    return jsonify({
        "success": True,
        "task": task,
        "auto_route": bool(_load_llm_config().get("fleet_auto_route", False)),
        "providers": providers,
        "models": models,
    })


@app.route('/api/models/ping/all', methods=['POST'])
def models_ping_all():
    """SSE: pinguje wszystkie modele + custom endpointy równolegle i streamuje wyniki."""
    _require_api_key()
    TEST = "Odpowiedz jednym słowem: OK"
    pool = _load_provider_pool()
    custom_endpoints = [e for e in pool.get("custom_endpoints", []) if e.get("active")]
    total_count = len(MODEL_REGISTRY) + len(custom_endpoints)

    def generate():
        import json as _json

        def sse(event, data):
            payload = {"event": event, "data": data}
            return f"data: {_json.dumps(payload, ensure_ascii=False)}\n\n"

        yield sse('start', {'count': total_count})

        def _ping_model(mid: str):
            t0 = time.time()
            try:
                r = requests.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {_pick_openrouter_key()}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "http://localhost",
                        "X-Title": "AI Analiza Dokumentów",
                    },
                    json={
                        "model": mid,
                        "messages": [{"role": "user", "content": TEST}],
                        "max_tokens": 5,
                        "temperature": 0.0,
                    },
                    timeout=30,
                )
                ms  = int((time.time() - t0) * 1000)
                if r.status_code != 200:
                    return ('model', mid, False, ms, (r.text or f"HTTP {r.status_code}")[:120])
                txt = _llm_response_text(r.json())
                ok  = bool(txt and not txt.startswith('['))
                return ('model', mid, ok, ms, None)
            except Exception as exc:
                return ('model', mid, False, int((time.time() - t0) * 1000), str(exc)[:120])

        def _ping_custom(endpoint_id: str, endpoint_url: str, endpoint_key: str, endpoint_name: str):
            t0 = time.time()
            try:
                health = _check_custom_endpoint_health(endpoint_url, endpoint_key, endpoint_name)
                ms = int(health.get("ms") or ((time.time() - t0) * 1000))
                ok = bool(health.get("ok"))
                return ('custom', endpoint_id, ok, ms, health.get("error"))
            except Exception as exc:
                ms = int((time.time() - t0) * 1000)
                return ('custom', endpoint_id, False, ms, str(exc)[:120])

        ping_jobs = [( _ping_model, (mid,)) for mid in MODEL_REGISTRY]
        ping_jobs.extend([
            (_ping_custom, (e['id'], e['url'], e.get('key', ''), e.get('name', '')))
            for e in custom_endpoints
        ])

        done = 0
        max_workers = min(12, max(1, total_count))
        with ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = [ex.submit(fn, *args) for fn, args in ping_jobs]
            try:
                completed_futures = as_completed(futures, timeout=90)
                for fut in completed_futures:
                    try:
                        kind, eid, ok, ms, err = fut.result()
                    except Exception as exc:
                        yield sse('result', {
                            'kind': 'model',
                            'model_id': '?',
                            'ok': False, 'ms': 0, 'err': str(exc)[:120],
                            'score': 0, 'avg_ms': None,
                        })
                        done += 1
                        continue

                    if kind == 'model':
                        live = _model_live[eid]
                        live['ping_ok'] = ok; live['ping_ms'] = ms
                        live['ping_ts'] = time.time()
                        if ok:
                            live['ok_count'] += 1
                            prev = live.get('avg_ms')
                            live['avg_ms'] = ms if prev is None else int(prev * 0.7 + ms * 0.3)
                        else:
                            live['err_count'] += 1
                        _save_model_live_stats(_model_live)
                        score = _model_score(eid)
                        live_copy = dict(live)
                        yield sse('result', {
                            'kind': 'model',
                            'model_id': eid,
                            'ok': ok, 'ms': ms, 'err': err,
                            'score': score,
                            'avg_ms': live_copy.get('avg_ms'),
                        })
                    elif kind == 'custom':
                        all_stats = _load_custom_endpoint_stats()
                        stats = all_stats.get(eid, {
                            "ping_ok": None, "ping_ms": None, "ping_ts": None,
                            "err_count": 0, "ok_count": 0, "avg_ms": None
                        })
                        stats['ping_ok'] = ok
                        stats['ping_ms'] = ms
                        stats['ping_ts'] = time.time()
                        if ok:
                            stats['ok_count'] += 1
                            prev = stats.get('avg_ms')
                            stats['avg_ms'] = ms if prev is None else int(prev * 0.7 + ms * 0.3)
                        else:
                            stats['err_count'] += 1
                        all_stats[eid] = stats
                        _save_custom_endpoint_stats(all_stats)
                        yield sse('result', {
                            'kind': 'custom',
                            'endpoint_id': eid,
                            'ok': ok, 'ms': ms, 'err': err,
                            'avg_ms': stats.get('avg_ms'),
                        })
                    done += 1
            except FuturesTimeoutError:
                for fut in futures:
                    fut.cancel()
                yield sse('timeout', {})

        yield sse('done', {'pinged': done})

    return Response(stream_with_context(generate()), mimetype='text/event-stream',
                    headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})


@app.route('/api/models/recommend', methods=['GET'])
def models_recommend():
    """Zwraca ranking modeli dla zadanego typu zadania i min. kontekstu."""
    task  = request.args.get('task', 'analysis')
    ctx_k = int(request.args.get('context_k', 8))
    task_tags = {
        'analysis': ['analiza', 'prawo', 'raporty', 'rozumowanie'],
        'speed':    ['szybkość', 'krótkie_pytania', 'klasyfikacja'],
        'data':     ['dane_strukturalne', 'tabele', 'ekstrakcja'],
        'long':     ['bardzo_długi_kontekst', 'długi_kontekst'],
    }.get(task, [])
    ranked = []
    for mid, info in MODEL_REGISTRY.items():
        if info.get('context_k', 0) < ctx_k:
            continue
        base  = _model_score(mid)
        bonus = sum(4 for t in info.get('tags', []) if t in task_tags)
        ranked.append({'model_id': mid, 'name': info['name'], 'icon': info['icon'],
                       'score': round(base + bonus, 1)})
    ranked.sort(key=lambda x: x['score'], reverse=True)
    return jsonify({"success": True, "ranked": ranked[:5], "task": task})


# ===== PROVIDER POOL API =====

@app.route('/api/providers', methods=['GET'])
def providers_list():
    """Lista wszystkich skonfigurowanych dostawców."""
    pool = _load_provider_pool()
    # Maskuj klucze API — pokaż tylko ostatnie 6 znaków
    safe = json.loads(json.dumps(pool))
    for entry in safe.get("openrouter_keys", []):
        k = entry.get("key", "")
        entry["key_masked"] = ("*" * max(0, len(k) - 6)) + k[-6:] if k else ""
        del entry["key"]
    for entry in safe.get("custom_endpoints", []):
        k = entry.get("key", "")
        entry["key_masked"] = ("*" * max(0, len(k) - 6)) + k[-6:] if k else ""
        entry["api_type"] = (
            "gemini" if _custom_endpoint_is_gemini(entry.get("url", ""), entry.get("name", ""))
            else "openai" if _custom_endpoint_is_openai(entry.get("url", ""))
            else "ollama"
        )
        if "key" in entry:
            del entry["key"]
    pool_or = [e for e in pool.get("openrouter_keys", []) if e.get("active")]
    pool_ol = [e for e in pool.get("ollama_urls", []) if e.get("active")]
    # effective = wpisy w .providers.json albo fallback z .env (gdy pula UI pusta)
    safe["defaults"] = {
        "openrouter_key_set": bool(OPENROUTER_API_KEY),
        "ollama_url": OLLAMA_URL,
        "gemini_key_set": bool(GEMINI_API_KEY),
        "gemini_model": GEMINI_MODEL,
        "pool_active_keys": len(pool_or),
        "pool_active_ollamas": len(pool_ol),
        "active_keys": len(pool_or) if pool_or else (1 if OPENROUTER_API_KEY else 0),
        "active_ollamas": len(pool_ol) if pool_ol else (1 if OLLAMA_URL else 0),
        "using_env_fallback_or": not pool_or and bool(OPENROUTER_API_KEY),
        "using_env_fallback_ollama": not pool_ol and bool(OLLAMA_URL),
    }
    return jsonify({"success": True, "pool": safe})


@app.route('/api/providers', methods=['POST'])
def providers_add():
    """Dodaje nowego dostawcę do puli."""
    if not _is_local_request():
        return jsonify({"success": False, "error": "Dostępne tylko z localhost"}), 403
    data = request.get_json() or {}
    ptype = data.get("type")  # openrouter_key | ollama_url | custom_endpoint
    if ptype not in ("openrouter_key", "ollama_url", "custom_endpoint"):
        return jsonify({"success": False, "error": "Nieznany typ dostawcy"})

    import uuid
    pool = _load_provider_pool()
    entry_id = str(uuid.uuid4())[:8]

    if ptype == "openrouter_key":
        key = (data.get("key") or "").strip()
        if not key:
            return jsonify({"success": False, "error": "Brak klucza API"})
        pool.setdefault("openrouter_keys", []).append({
            "id": entry_id, "name": data.get("name", f"Konto {entry_id}"),
            "key": key, "active": True,
        })
    elif ptype == "ollama_url":
        url = (data.get("url") or "").strip().rstrip("/")
        if not url:
            return jsonify({"success": False, "error": "Brak adresu URL"})
        pool.setdefault("ollama_urls", []).append({
            "id": entry_id, "name": data.get("name", url),
            "url": url, "active": True,
        })
    elif ptype == "custom_endpoint":
        url = (data.get("url") or "").strip().rstrip("/")
        if not url:
            return jsonify({"success": False, "error": "Brak adresu URL"})
        pool.setdefault("custom_endpoints", []).append({
            "id": entry_id, "name": data.get("name", url),
            "url": url, "key": data.get("key", ""), "active": True,
        })

    _save_provider_pool(pool)
    return jsonify({"success": True, "id": entry_id})


@app.route('/api/providers/<entry_id>', methods=['DELETE'])
def providers_delete(entry_id: str):
    """Usuwa dostawcę z puli."""
    if not _is_local_request():
        return jsonify({"success": False, "error": "Dostępne tylko z localhost"}), 403
    pool = _load_provider_pool()
    removed = False
    for section in ("openrouter_keys", "ollama_urls", "custom_endpoints"):
        before = len(pool.get(section, []))
        pool[section] = [e for e in pool.get(section, []) if e.get("id") != entry_id]
        if len(pool[section]) < before:
            removed = True
    if removed:
        _save_provider_pool(pool)
        return jsonify({"success": True})
    return jsonify({"success": False, "error": "Nie znaleziono"}), 404


@app.route('/api/providers/<entry_id>/toggle', methods=['POST'])
def providers_toggle(entry_id: str):
    """Włącza/wyłącza dostawcę."""
    if not _is_local_request():
        return jsonify({"success": False, "error": "Dostępne tylko z localhost"}), 403
    pool = _load_provider_pool()
    for section in ("openrouter_keys", "ollama_urls", "custom_endpoints"):
        for entry in pool.get(section, []):
            if entry.get("id") == entry_id:
                entry["active"] = not entry.get("active", True)
                _save_provider_pool(pool)
                return jsonify({"success": True, "active": entry["active"]})
    return jsonify({"success": False, "error": "Nie znaleziono"}), 404


@app.route('/api/providers/<entry_id>/test', methods=['POST'])
def providers_test(entry_id: str):
    """Testuje połączenie z dostawcą (krótki ping)."""
    pool = _load_provider_pool()
    entry = None
    ptype = None
    for section in ("openrouter_keys", "ollama_urls", "custom_endpoints"):
        for e in pool.get(section, []):
            if e.get("id") == entry_id:
                entry = e; ptype = section; break

    if not entry:
        # Może to być test domyślnego
        if entry_id == "default_openrouter":
            entry = {"key": OPENROUTER_API_KEY}; ptype = "openrouter_keys"
        elif entry_id == "default_ollama":
            entry = {"url": OLLAMA_URL}; ptype = "ollama_urls"
        elif entry_id == "default_gemini":
            ok, message, ms = _test_gemini_api(GEMINI_API_KEY)
            return jsonify({"success": ok, "ms": ms, "error": None if ok else message})
        else:
            return jsonify({"success": False, "error": "Nie znaleziono"}), 404

    t0 = time.time()
    try:
        if ptype == "openrouter_keys":
            key = entry.get("key", "")
            return jsonify(_test_openrouter_api_key(key))
        elif ptype == "ollama_urls":
            url = entry.get("url", "").rstrip("/")
            key = entry.get("key", "")
            headers = {}
            if key:
                headers["Authorization"] = f"Bearer {key}"
            r = requests.get(f"{url}/api/tags", headers=headers, timeout=8)
            ms = int((time.time() - t0) * 1000)
            return jsonify({"success": r.status_code == 200, "ms": ms, "status": r.status_code})
        elif ptype == "custom_endpoints":
            url = entry.get("url", "").rstrip("/")
            key = entry.get("key", "")
            name = entry.get("name", "") or entry.get("label", "")
            if _custom_endpoint_is_gemini(url, name):
                health = _check_custom_endpoint_health(url, key, name)
                return jsonify({
                    "success": bool(health.get("ok")),
                    "ms": health.get("ms", 0),
                    "error": health.get("error"),
                })
            elif _custom_endpoint_is_openai(url):
                headers = {"Content-Type": "application/json"}
                if key:
                    headers["Authorization"] = f"Bearer {key}"
                r = requests.post(
                    _openai_chat_completions_url(url),
                    headers=headers,
                    json={"model": "llama-3.3-70b-versatile",
                          "messages": [{"role": "user", "content": "Say: OK"}],
                          "max_tokens": 5},
                    timeout=20,
                )
                ms = int((time.time() - t0) * 1000)
                return jsonify({"success": r.status_code == 200, "ms": ms, "status": r.status_code,
                               "error": None if r.status_code == 200 else r.text[:200]})
            else:
                headers = {}
                if key:
                    headers["Authorization"] = f"Bearer {key}"
                r = requests.get(f"{url}/api/tags", headers=headers, timeout=8)
                ms = int((time.time() - t0) * 1000)
                return jsonify({"success": r.status_code == 200, "ms": ms, "status": r.status_code})
    except Exception as exc:
        ms = int((time.time() - t0) * 1000)
        return jsonify({"success": False, "ms": ms, "error": str(exc)[:200]})


# ===== SWARM REPORTS API =====

@app.route('/api/swarm/reports', methods=['POST'])
def swarm_reports_save():
    """Zapisuje wynik roju agentów jako raport JSON."""
    data = request.get_json() or {}
    query   = (data.get('query') or '').strip()
    mode    = (data.get('mode') or 'quality').strip()
    mode_label = (data.get('mode_label') or mode).strip()
    n_workers  = int(data.get('n_workers') or 0)
    answer  = (data.get('answer') or '').strip()

    if not query or not answer:
        return jsonify({"success": False, "error": "Wymagane query i answer"})

    _ensure_swarm_reports_dir()
    import uuid
    report_id = str(uuid.uuid4())[:8]
    title = (query[:57] + '…') if len(query) > 60 else query
    report = {
        "id": report_id,
        "title": title,
        "query": query,
        "mode": mode,
        "mode_label": mode_label,
        "n_workers": n_workers,
        "answer": answer,
        "saved_at": datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%S'),
    }
    (_SWARM_REPORTS_DIR / f"{report_id}.json").write_text(
        json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    return jsonify({"success": True, "id": report_id})


@app.route('/api/swarm/reports', methods=['GET'])
def swarm_reports_list():
    """Lista zapisanych raportów roju, posortowana od najnowszego."""
    _ensure_swarm_reports_dir()
    reports = []
    for p in _SWARM_REPORTS_DIR.glob("*.json"):
        try:
            r = json.loads(p.read_text(encoding="utf-8"))
            # Bezpieczny widok — answer może być długi, zwróć tylko preview
            reports.append({
                "id": r.get("id"),
                "title": r.get("title"),
                "query": r.get("query"),
                "mode": r.get("mode"),
                "mode_label": r.get("mode_label"),
                "n_workers": r.get("n_workers"),
                "saved_at": r.get("saved_at"),
                "answer_preview": (r.get("answer") or "")[:300],
            })
        except Exception:
            continue
    reports.sort(key=lambda x: x.get("saved_at", ""), reverse=True)
    return jsonify({"success": True, "reports": reports})


@app.route('/api/swarm/reports/<report_id>', methods=['GET'])
def swarm_reports_get(report_id: str):
    """Pełny raport roju (z pełną treścią answer)."""
    if not re.match(r'^[a-f0-9]{8}$', report_id):
        return jsonify({"success": False, "error": "Nieprawidłowe ID"}), 400
    p = _SWARM_REPORTS_DIR / f"{report_id}.json"
    if not p.exists():
        return jsonify({"success": False, "error": "Nie znaleziono"}), 404
    try:
        return jsonify({"success": True, "report": json.loads(p.read_text(encoding="utf-8"))})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/api/swarm/reports/<report_id>', methods=['DELETE'])
def swarm_reports_delete(report_id: str):
    """Usuwa zapisany raport roju."""
    if not _is_local_request():
        return jsonify({"success": False, "error": "Dostępne tylko z localhost"}), 403
    if not re.match(r'^[a-f0-9]{8}$', report_id):
        return jsonify({"success": False, "error": "Nieprawidłowe ID"}), 400
    p = _SWARM_REPORTS_DIR / f"{report_id}.json"
    if not p.exists():
        return jsonify({"success": False, "error": "Nie znaleziono"}), 404
    p.unlink()
    return jsonify({"success": True})


@app.route('/suggestions', methods=['GET'])
def get_suggestions():
    force = request.args.get('force', '0') == '1'
    now = time.time()
    if not force and _suggestions_cache["data"] and (now - _suggestions_cache["ts"]) < SUGGESTIONS_TTL:
        return jsonify({"success": True, "suggestions": _suggestions_cache["data"], "cached": True})
    try:
        client = get_qdrant_client()
        # Losowa próbka: pobierz ~300 chunków, wybierz 25 z różnych plików
        records, _ = client.scroll(
            collection_name=ACTIVE_COLLECTION,
            limit=300,
            offset=None,
            with_payload=["file", "text"],
            with_vectors=False
        )
        import random
        seen_files = set()
        sample = []
        random.shuffle(records)
        for r in records:
            fname = r.payload.get("file", "")
            if fname not in seen_files and len(r.payload.get("text", "")) > 100:
                seen_files.add(fname)
                sample.append(r.payload)
            if len(sample) >= 25:
                break

        context = "\n\n".join([
            f"[{p['file']}]: {p['text'][:600]}" for p in sample
        ])
        prompt = (
            "Na podstawie poniższych fragmentów dokumentów śledczych, wygeneruj dokładnie 8 konkretnych pytań analitycznych. "
            "Każde pytanie musi być zwięzłe (max 12 słów), dotyczyć konkretnych faktów, kwot, dat, nazwisk lub zdarzeń z tych dokumentów. "
            "Zwróć TYLKO listę 8 pytań, każde w osobnej linii, bez numeracji, bez komentarzy.\n\n"
            f"DOKUMENTY:\n{context}"
        )
        system_msg = "Jesteś analitykiem śledczym. Odpowiadasz wyłącznie po polsku. Zwracasz tylko listę pytań."
        result = _call_ollama(prompt, system_msg, stream=False, model=LLM_MODEL)
        raw = result.get("response", "") if isinstance(result, dict) else ""

        lines = [l.strip().lstrip("-•·1234567890.). ") for l in raw.strip().splitlines()]
        suggestions = [l for l in lines if len(l) > 10][:8]

        _suggestions_cache["data"] = suggestions
        _suggestions_cache["ts"] = now
        return jsonify({"success": True, "suggestions": suggestions, "cached": False})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

NOISE_PATTERNS = [
    # Licencje i pliki prawne bibliotek
    r'^(GPL|LGPL|MIT|Apache License|OFL|OFL-\d|SIL Open Font|Architects Daughter SIL|LICENSE|COPYING|NOTICE)(\.\w+)?$',
    # Logi i pliki tymczasowe
    r'.*\.log$', r'^cat-log.*', r'^SaveAs_log.*', r'^ChangeLog.*',
    # Pliki bibliotek/fontów
    r'^OFL.*\.txt$', r'^OFL-FAQ.*',
    # Zaszyfrowane/tokenizowane nazwy URL (długie stringi base64/JWT z = lub +)
    r'^[A-Za-z0-9+/=]{40,}.*',
    # Dokumentacja techniczna niezwiązana ze sprawą (można rozszerzyć)
    r'^sss\..*',
]

import re as _re

def _is_noise(fname: str) -> str:
    """Zwraca powód detekcji szumu lub pusty string jeśli plik jest OK."""
    for pat in NOISE_PATTERNS:
        if _re.match(pat, fname, _re.IGNORECASE):
            return f"wzorzec: {pat[:40]}"
    # Zaszyfrowana nazwa — długa, zawiera = lub + ale nie jest normalną nazwą pliku
    if len(fname) > 60 and ('=' in fname or (fname.count('+') > 2)):
        return "zaszyfrowana nazwa (token URL)"
    return ""

@app.route('/documents/scan-noise', methods=['GET'])
def scan_noise():
    try:
        client = get_qdrant_client()
        file_chunks = {}
        offset = None
        while True:
            records, offset = client.scroll(
                collection_name=ACTIVE_COLLECTION, limit=250, offset=offset,
                with_payload=["file"], with_vectors=False
            )
            for r in records:
                fname = r.payload.get("file", "")
                if fname:
                    file_chunks[fname] = file_chunks.get(fname, 0) + 1
            if offset is None:
                break

        noise = []
        for fname, count in file_chunks.items():
            reason = _is_noise(fname)
            if reason:
                noise.append({"file": fname, "chunks": count, "reason": reason})

        noise.sort(key=lambda x: -x["chunks"])
        return jsonify({"success": True, "noise": noise, "total_chunks": sum(n["chunks"] for n in noise)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/documents/delete', methods=['POST'])
def delete_documents():
    data = request.get_json()
    files_to_delete = data.get('files', [])
    if not files_to_delete:
        return jsonify({"success": False, "error": "Brak listy plików"})
    try:
        from qdrant_client.models import Filter, FieldCondition, MatchAny
        client = get_qdrant_client()
        # Skasuj wszystkie punkty gdzie payload.file znajduje się na liście do usunięcia
        client.delete(
            collection_name=ACTIVE_COLLECTION,
            points_selector=Filter(
                must=[FieldCondition(key="file", match=MatchAny(any=files_to_delete))]
            )
        )
        _suggestions_cache["data"] = None; _docs_cache["data"] = None
        return jsonify({"success": True, "files_count": len(files_to_delete)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})

@app.route('/documents', methods=['GET'])
def get_documents():
    force = request.args.get('force', '0') == '1'
    ext_filter = request.args.getlist('ext')
    modified_after = request.args.get('modified_after', '').strip()
    modified_before = request.args.get('modified_before', '').strip()
    now = time.time()
    if not force and not ext_filter and not modified_after and not modified_before \
            and _docs_cache["data"] and (now - _docs_cache["ts"]) < DOCS_CACHE_TTL:
        return jsonify({"success": True, "documents": _docs_cache["data"],
                        "total": len(_docs_cache["data"]), "cached": True})
    try:
        client = get_qdrant_client()
        file_chunks = {}
        file_paths  = {}
        file_meta   = {}
        offset = None
        while True:
            records, offset = client.scroll(
                collection_name=ACTIVE_COLLECTION,
                limit=250, offset=offset,
                with_payload=["file", "full_path", "metadata"],
                with_vectors=False
            )
            for r in records:
                fname = r.payload.get("file", "")
                if fname:
                    file_chunks[fname] = file_chunks.get(fname, 0) + 1
                    if fname not in file_paths and r.payload.get("full_path"):
                        file_paths[fname] = r.payload["full_path"]
                    if fname not in file_meta and r.payload.get("metadata"):
                        file_meta[fname] = r.payload["metadata"]
            if offset is None:
                break
        docs = sorted(
            [{
                "file": f,
                "chunks": c,
                "full_path": file_paths.get(f, ""),
                "metadata": file_meta.get(f)
            }
             for f, c in file_chunks.items()],
            key=lambda x: -x["chunks"]
        )
        _docs_cache["data"] = docs
        _docs_cache["ts"]   = now
        filtered = _filter_documents_list(docs, ext_filter or None, modified_after, modified_before)
        return jsonify({
            "success": True,
            "documents": filtered,
            "total": len(filtered),
            "total_unfiltered": len(docs),
            "cached": False,
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/api/get_context', methods=['GET'])
def get_context():
    """Lazy-load fragmentu dokumentu (chunk z Qdrant lub skrót z pliku źródłowego)."""
    point_id = (request.args.get('point_id') or '').strip()
    file_name = (request.args.get('file') or '').strip()
    query = (request.args.get('query') or '').strip()
    source = (request.args.get('source') or 'chunk').strip().lower()

    if not point_id and not file_name:
        return jsonify({"success": False, "error": "Wymagany point_id lub file"}), 400

    try:
        client = get_qdrant_client()

        if point_id and source != 'file':
            records = client.retrieve(
                collection_name=ACTIVE_COLLECTION,
                ids=[point_id],
                with_payload=True,
                with_vectors=False,
            )
            if records:
                p = records[0].payload or {}
                text = p.get("text", "")
                if query:
                    text = highlight_backend(text, query)
                full_path = p.get("full_path", "")
                return jsonify({
                    "success": True,
                    "point_id": point_id,
                    "file": p.get("file", file_name),
                    "text": text,
                    "full_path": full_path,
                    "win_path": wsl_to_win(full_path),
                    "metadata": p.get("metadata"),
                    "source": "qdrant",
                })

        if file_name:
            from qdrant_client.models import Filter, FieldCondition, MatchValue

            if query and source == 'chunk':
                try:
                    vector = get_embedding(query)
                    res = client.query_points(
                        collection_name=ACTIVE_COLLECTION,
                        query=vector,
                        limit=1,
                        query_filter=Filter(
                            must=[FieldCondition(key="file", match=MatchValue(value=file_name))]
                        ),
                    )
                    if res.points:
                        p = res.points[0].payload or {}
                        text = highlight_backend(p.get("text", ""), query)
                        full_path = p.get("full_path", "")
                        return jsonify({
                            "success": True,
                            "point_id": str(res.points[0].id),
                            "file": file_name,
                            "text": text,
                            "full_path": full_path,
                            "win_path": wsl_to_win(full_path),
                            "metadata": p.get("metadata"),
                            "source": "qdrant_search",
                        })
                except EmbeddingError as e:
                    return jsonify({"success": False, "error": str(e)})

            scroll_filter = Filter(
                must=[FieldCondition(key="file", match=MatchValue(value=file_name))]
            )
            records, _ = client.scroll(
                collection_name=ACTIVE_COLLECTION,
                scroll_filter=scroll_filter,
                limit=1,
                with_payload=True,
                with_vectors=False,
            )
            if records:
                p = records[0].payload or {}
                text = highlight_backend(p.get("text", ""), query)
                full_path = p.get("full_path", "")
                return jsonify({
                    "success": True,
                    "point_id": str(records[0].id),
                    "file": file_name,
                    "text": text,
                    "full_path": full_path,
                    "win_path": wsl_to_win(full_path),
                    "metadata": p.get("metadata"),
                    "source": "qdrant_scroll",
                })

            path = _find_file_by_name_safe(file_name)
            if path and path.is_file() and _path_is_allowed(path):
                excerpt = extract_text(path)[:12000]
                if query:
                    excerpt = highlight_backend(excerpt, query)
                return jsonify({
                    "success": True,
                    "file": file_name,
                    "text": excerpt,
                    "full_path": str(path),
                    "win_path": wsl_to_win(str(path)),
                    "metadata": extract_file_metadata(path),
                    "source": "file",
                    "truncated": len(excerpt) >= 12000,
                })

        return jsonify({"success": False, "error": "Nie znaleziono fragmentu dokumentu"}), 404
    except Exception as e:
        logger.exception("get_context")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/export/metadata_report', methods=['POST'])
def export_metadata_report_docx():
    """Generuje raport DOCX z metadanymi wybranych dokumentów (dla śledczych)."""
    data = request.get_json() or {}
    selected = data.get('files', [])  # lista obiektów z file, full_path, metadata

    if not selected:
        return jsonify({"success": False, "error": "Brak wybranych plików"}), 400

    try:
        import docx as _docx
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = _docx.Document()

        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Nagłówek
        h = doc.add_heading('Raport Metadanych Wybranych Dokumentów', 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Wygenerowano: {time.strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph(f'Liczba plików: {len(selected)}')
        doc.add_paragraph()

        # Podsumowanie
        doc.add_heading('Podsumowanie', level=1)
        doc.add_paragraph('Poniżej znajdują się metadane systemowe oraz wbudowane dla wybranych dokumentów. '
                          'Dane pochodzą z oryginalnych plików w momencie importu.')

        # Dla każdego pliku
        for idx, f in enumerate(selected, 1):
            doc.add_heading(f'{idx}. {f.get("file", "Nieznany plik")}', level=1)

            meta = f.get('metadata') or {}

            # Metadane systemowe
            doc.add_heading('Metadane systemowe', level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Pole'
            hdr_cells[1].text = 'Wartość'

            for key in ['size_human', 'created', 'modified', 'accessed']:
                if meta.get(key):
                    row_cells = table.add_row().cells
                    row_cells[0].text = key.replace('_', ' ').title()
                    row_cells[1].text = str(meta[key])

            # Metadane wbudowane
            if meta.get('embedded'):
                doc.add_heading('Metadane wbudowane w plik', level=2)
                for fmt, emb in meta['embedded'].items():
                    p = doc.add_paragraph()
                    p.add_run(f'{fmt.upper()}: ').bold = True
                    p.add_run(str(emb))

            # Pełna ścieżka
            if meta.get('full_path'):
                p = doc.add_paragraph()
                p.add_run('Pełna ścieżka: ').bold = True
                p.add_run(meta['full_path'])

            doc.add_paragraph()  # odstęp

        # Stopka
        doc.add_paragraph('─' * 70)
        p = doc.add_paragraph()
        run = p.add_run('Wygenerowano przez AI Analiza Dokumentów | Narzędzie wspomagające analizę śledczą')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        fname = f'raport_metadanych_{time.strftime("%Y%m%d_%H%M")}.docx'
        return send_file(buf, as_attachment=True, download_name=fname,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/export/docx', methods=['POST'])
def export_docx():
    data        = request.get_json()
    query       = data.get('query', '')
    ai_answer   = data.get('ai_answer', '')
    results     = data.get('results', [])
    mode_label  = data.get('mode_label', 'Standardowy')

    try:
        import docx as _docx
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = _docx.Document()

        # Styl dokumentu
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Nagłówek
        h = doc.add_heading('Raport Śledczy — Analiza RAG', 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Data: {time.strftime("%d.%m.%Y %H:%M")}   |   Tryb: {mode_label}')
        doc.add_paragraph(f'Kolekcja: {ACTIVE_COLLECTION}   |   Wyników: {len(results)}')
        doc.add_paragraph()

        # Zapytanie
        doc.add_heading('Zapytanie', level=1)
        p = doc.add_paragraph()
        p.add_run(query).bold = True

        # Odpowiedź LLM
        doc.add_heading('Synteza AI (Llama3)', level=1)
        for line in ai_answer.split('\n'):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(style='List Bullet' if line.startswith('-') else 'Normal')
            p.add_run(line.lstrip('- '))

        # Źródła
        doc.add_heading(f'Dokumenty źródłowe ({len(results)})', level=1)
        for i, r in enumerate(results, 1):
            doc.add_heading(f'{i}. {r.get("file","?")}', level=2)
            score_p = doc.add_paragraph()
            run = score_p.add_run(f'Dopasowanie: {float(r.get("score",0))*100:.1f}%')
            run.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)
            run.font.size = Pt(9)
            # Usuń tagi HTML z tekstu
            clean = re.sub(r'<[^>]+>', '', r.get('text', ''))
            doc.add_paragraph(clean[:1500])

            if r.get('win_path'):
                p = doc.add_paragraph()
                run = p.add_run(f'Ścieżka: {r["win_path"]}')
                run.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)
                run.font.size = Pt(9)

            doc.add_paragraph()

        # Stopka
        doc.add_paragraph('─' * 60)
        doc.add_paragraph('Wygenerowano przez AI Analiza Dokumentów | Llama3 + Qdrant Cloud').runs[0].font.size = Pt(8)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        fname = f'raport_{time.strftime("%Y%m%d_%H%M")}.docx'
        return send_file(buf, as_attachment=True, download_name=fname,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


_NETWORK_SYSTEM = (
    "Jesteś ekspertem analityki śledczej. Twoim zadaniem jest wyciągnięcie jak najbogatszej sieci powiązań z dokumentów.\n\n"
    "ZASADY BEZWZGLĘDNE:\n"
    "1. Używaj WYŁĄCZNIE prawdziwych nazw, kwot, dat i faktów z dokumentów.\n"
    "2. NIGDY nie używaj placeholderów typu 'nazwa', 'firma', 'X', 'osoba A'.\n"
    "3. Każdy label musi być rzeczywistą wartością z tekstu.\n\n"
    "FORMAT JSON (zwracaj tylko poprawny JSON):\n"
    '{\n'
    '  "nodes": [ {"id": "jan_kowalski", "type": "osoba", "label": "Jan Kowalski"}, ... ],\n'
    '  "edges": [ {"source": "jan_kowalski", "target": "abc_spolka", "label": "prezes zarządu", '
    '"doc": "umowa_2023.pdf", "evidence": "Podpisał umowę 12.03.2023", "date": "2023-03-12", "strength": 3}, ... ]\n'
    '}\n\n'
    "Typy węzłów (type): osoba, firma, kwota, dokument, przetarg, umowa, inne\n\n"
    "Bogate typy relacji (label krawędzi) — używaj precyzyjnych:\n"
    "- prezes / członek zarządu / dyrektor\n"
    "- podpisał umowę / aneks\n"
    "- zapłacił / wystawił fakturę / otrzymał płatność\n"
    "- wygrał przetarg / złożył ofertę\n"
    "- zlecił / wykonał usługę\n"
    "- jest właścicielem / beneficjentem\n"
    "- zatwierdził / skontrolował\n\n"
    "W polu 'evidence' wstaw krótki, dosłowny cytat z dokumentu.\n"
    "Jeśli w tekście pojawia się data — dodaj ją w polu 'date' (YYYY-MM-DD lub YYYY-MM).\n"
    "W polu 'strength' (1-5) podaj siłę powiązania na podstawie dowodów w dokumencie.\n"
    "id: snake_case, bez polskich znaków, max 40 znaków."
)


def _aggregate_network_edges(raw_edges: list, seen_nodes: dict) -> list:
    """Grupuje krawędzie, sumuje siłę i zbiera dowody."""
    edge_groups = defaultdict(list)
    for e in raw_edges:
        src = str(e.get("source", "")).strip()
        tgt = str(e.get("target", "")).strip()
        if not src or not tgt or src not in seen_nodes or tgt not in seen_nodes:
            continue
        edge_groups[(src, tgt)].append({
            "label": (e.get("label") or "")[:40],
            "doc": (e.get("doc") or "")[:90],
            "evidence": (e.get("evidence") or "")[:180],
            "date": (e.get("date") or "")[:20],
            "strength": _safe_int_clamp(e.get("strength", 1), 1, 5, default=1),
        })

    clean_edges = []
    for (src, tgt), items in edge_groups.items():
        labels = [it["label"] for it in items if it["label"]]
        main_label = max(set(labels), key=labels.count) if labels else items[0]["label"]

        seen_ev = set()
        evidences = []
        for it in items:
            ev_key = (it["doc"], it["evidence"][:80])
            if ev_key not in seen_ev:
                seen_ev.add(ev_key)
                evidences.append({
                    "doc": it["doc"],
                    "evidence": it["evidence"],
                    "date": it["date"],
                })
            if len(evidences) >= 6:
                break

        total_strength = sum(it["strength"] for it in items) + (len(items) - 1)
        clean_edges.append({
            "source": src,
            "target": tgt,
            "label": main_label,
            "doc": items[0]["doc"],
            "evidence": items[0]["evidence"],
            "date": items[0]["date"],
            "strength": max(1, min(12, total_strength)),
            "evidence_count": len(evidences),
            "evidences": evidences,
        })

    clean_edges.sort(key=lambda e: e.get("strength", 1), reverse=True)
    return clean_edges


def _network_graph_stats(nodes: list, edges: list) -> dict:
    """Statystyki grafu bez dodatkowego wywołania LLM."""
    degree = defaultdict(int)
    for e in edges:
        degree[e["source"]] += e.get("strength", 1)
        degree[e["target"]] += e.get("strength", 1)

    type_counts = defaultdict(int)
    for n in nodes:
        type_counts[n.get("type", "inne")] += 1

    rel_counts = defaultdict(int)
    for e in edges:
        rel_counts[e.get("label") or "(bez etykiety)"] += 1

    dates = sorted(d for e in edges if (d := (e.get("date") or "").strip()) and d[0].isdigit())

    hubs = sorted(
        [{"id": n["id"], "label": n.get("label", n["id"]), "type": n.get("type", "inne"),
          "score": degree.get(n["id"], 0)}
         for n in nodes if degree.get(n["id"], 0) > 0],
        key=lambda h: h["score"],
        reverse=True,
    )[:8]

    top_relations = sorted(rel_counts.items(), key=lambda x: x[1], reverse=True)[:8]

    return {
        "entity_counts": dict(type_counts),
        "top_hubs": hubs,
        "top_relations": [{"label": k, "count": v} for k, v in top_relations],
        "date_span": {"min": dates[0], "max": dates[-1]} if dates else None,
        "strong_edges": sum(1 for e in edges if e.get("strength", 1) >= 3),
    }


def _network_ai_briefing(query: str, nodes: list, edges: list, stats: dict,
                         provider=None, model=None) -> str:
    """Krótki briefing śledczy na podstawie zbudowanego grafu."""
    if not nodes or not edges:
        return ""

    hub_lines = ", ".join(f"{h['label']} ({h['score']})" for h in stats.get("top_hubs", [])[:5])
    rel_lines = ", ".join(f"{r['label']}×{r['count']}" for r in stats.get("top_relations", [])[:5])
    sample_edges = edges[:12]
    edge_txt = "\n".join(
        f"- {e.get('source')} → {e.get('target')}: {e.get('label')} "
        f"(siła {e.get('strength', 1)}, {e.get('doc', '')})"
        for e in sample_edges
    )

    prompt = (
        f"ZAPYTANIE UŻYTKOWNIKA: {query}\n\n"
        f"GRAF: {len(nodes)} węzłów, {len(edges)} relacji.\n"
        f"Kluczowe węzły (huby): {hub_lines or 'brak'}.\n"
        f"Najczęstsze relacje: {rel_lines or 'brak'}.\n\n"
        f"Przykładowe krawędzie:\n{edge_txt}\n\n"
        "Napisz krótki briefing śledczy (max 6 punktów):\n"
        "1) główne podmioty i ich role,\n"
        "2) najsilniejsze powiązania finansowe/prawne,\n"
        "3) potencjalne rozbieżności lub luki [WYMAGA SPRAWDZENIA],\n"
        "4) 2-3 konkretne pytania do dalszej analizy.\n"
        "Tylko fakty z grafu — bez wymyślania."
    )
    system = (
        "Jesteś analitykiem śledczym. Pisz po polsku, zwięźle, punktami. "
        "Oznaczaj niepewności tagiem [WYMAGA SPRAWDZENIA]."
    )
    try:
        result = call_llm(
            prompt=prompt, system=system, stream=False,
            provider=provider, model=model, max_tokens=900,
        )
        return (result.get("response", "") if isinstance(result, dict) else str(result)).strip()
    except Exception as e:
        logger.warning(f"Briefing sieci powiązań: {e}")
        return ""


@app.route('/network', methods=['POST'])
def build_network():
    """
    Buduje sieć powiązań partiami (po 5 dokumentów).
    Zwraca SSE: progress → done (graf + statystyki + opcjonalny briefing AI).
    """
    data = request.get_json() or {}
    query = data.get('query', '').strip()
    limit = min(int(data.get('limit', 10)), 20)
    llm_provider = data.get('llm_provider')
    openrouter_model = data.get("openrouter_model")
    include_briefing = data.get("include_briefing", True)
    BATCH_SIZE = 5

    if not query:
        return jsonify({"success": False, "error": "Brak zapytania"})

    def generate():
        def sse(event, payload):
            return f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n"

        try:
            client = get_qdrant_client()
            try:
                vector = get_embedding(query)
            except EmbeddingError as e:
                yield sse("error", {"error": str(e)})
                return

            res = client.query_points(collection_name=ACTIVE_COLLECTION, query=vector, limit=limit)
            all_contexts = [
                {"file": p.payload.get("file", ""), "text": p.payload.get("text", "")}
                for p in res.points
            ]

            if not all_contexts:
                yield sse("error", {"error": "Nie znaleziono dokumentów pasujących do zapytania"})
                return

            batches = [all_contexts[i:i + BATCH_SIZE] for i in range(0, len(all_contexts), BATCH_SIZE)]
            total = len(all_contexts)
            total_batches = len(batches)

            seen_nodes = {}
            raw_edges = []

            yield sse("progress", {
                "processed": 0,
                "total": total,
                "batch": 0,
                "total_batches": total_batches,
                "message": f"Znaleziono {total} fragmentów w {total_batches} partiach…",
            })

            for i, batch in enumerate(batches):
                yield sse("progress", {
                    "processed": min(i * BATCH_SIZE, total),
                    "total": total,
                    "batch": i + 1,
                    "total_batches": total_batches,
                    "message": f"Analiza partii {i + 1}/{total_batches} (LLM ekstrahuje encje)…",
                })

                batch_str = "\n\n".join(
                    f"[{c['file']}]: {_sanitize_for_prompt(c['text'], 800)}" for c in batch
                )
                prompt = (
                    f"DOKUMENTY (partia {i + 1}/{total_batches}):\n{batch_str}\n\n"
                    f"ZADANIE: {query}\n\n"
                    "Wyciągnij sieć powiązań używając PRAWDZIWYCH nazw, kwot i danych.\n"
                    "NIE używaj placeholderów. JSON:"
                )

                try:
                    result = call_llm(
                        prompt=prompt,
                        system=_NETWORK_SYSTEM,
                        stream=False,
                        provider=llm_provider,
                        model=openrouter_model,
                        max_tokens=2800,
                    )
                    raw = result.get("response", "") if isinstance(result, dict) else str(result)
                    json_match = re.search(r'\{[\s\S]*\}', raw)
                    if not json_match:
                        continue

                    graph = json.loads(json_match.group())

                    for n in graph.get("nodes", []):
                        nid = str(n.get("id", "")).strip()
                        if nid and nid not in seen_nodes:
                            seen_nodes[nid] = {
                                "id": nid,
                                "type": n.get("type", "inne"),
                                "label": (n.get("label") or nid)[:40],
                            }

                    for e in graph.get("edges", []):
                        raw_edges.append(e)

                except Exception as batch_err:
                    logger.warning(f"Błąd partii sieci {i + 1}: {batch_err}")
                    continue

            clean_nodes = list(seen_nodes.values())
            clean_edges = _aggregate_network_edges(raw_edges, seen_nodes)
            stats = _network_graph_stats(clean_nodes, clean_edges)

            briefing = ""
            if include_briefing and clean_nodes and clean_edges:
                yield sse("progress", {
                    "processed": total,
                    "total": total,
                    "batch": total_batches,
                    "total_batches": total_batches,
                    "message": "Generuję briefing śledczy AI…",
                })
                briefing = _network_ai_briefing(
                    query, clean_nodes, clean_edges, stats,
                    provider=llm_provider, model=openrouter_model,
                )

            yield sse("done", {
                "success": True,
                "nodes": clean_nodes,
                "edges": clean_edges,
                "query": query,
                "sources": len(all_contexts),
                "stats": stats,
                "briefing": briefing,
            })

        except Exception as e:
            logger.exception("Błąd build_network")
            yield sse("error", {"error": str(e)})

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'},
    )


@app.route('/timeline', methods=['POST'])
def timeline_endpoint():
    """Oś czasu — ekstrahuje daty z dokumentów i buduje chronologię (SSE)."""
    data = request.get_json() or {}
    query_text = (data.get('query') or '').strip()
    limit = min(int(data.get('limit', 20)), 50)
    llm_provider = data.get('llm_provider')
    openrouter_model = data.get('openrouter_model')

    def generate():
        def sse(event, payload):
            return f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n"

        try:
            client = get_qdrant_client()
            payloads = []

            # Pobieramy dokumenty — wg query lub ostatnie
            if query_text:
                yield sse("progress", {"message": "Wyszukiwanie dokumentów…"})
                try:
                    vector = get_embedding(query_text)
                    res = client.query_points(
                        collection_name=ACTIVE_COLLECTION,
                        query=vector,
                        limit=limit,
                        with_payload=True,
                        with_vectors=False
                    )
                    payloads = [p.payload for p in res.points]
                except Exception as e:
                    logger.warning(f"Błąd przy pobieraniu embeddingu: {e}")
                    payloads = []

            if not payloads:
                # Fallback — pobierz ostatnie dokumenty
                yield sse("progress", {"message": "Pobieranie ostatnich dokumentów…"})
                records, _ = client.scroll(
                    collection_name=ACTIVE_COLLECTION,
                    limit=limit,
                    with_payload=True,
                    with_vectors=False
                )
                payloads = [r.payload for r in records]

            if not payloads:
                yield sse("done", {
                    "success": True,
                    "timeline": [],
                    "message": "Brak dokumentów do analizy."
                })
                return

            # Pre-filtracja: szukamy dat za pomocą RegEx
            yield sse("progress", {"message": "Ekstrakcja dat (RegEx)…"})
            date_pattern = re.compile(
                r'\b(?:\d{1,4}[.\/-]\d{1,2}[.\/-]\d{1,4}|'
                r'\d{1,2}\s+(?:stycznia|lutego|marca|kwietnia|maja|czerwca|lipca|sierpnia|września|października|listopada|grudnia)\s+(?:19|20)\d{2})\b',
                re.IGNORECASE
            )

            date_chunks = []
            for p in payloads:
                text = p.get("text", "")
                if date_pattern.search(text):
                    date_chunks.append({
                        "file": p.get("file", "Nieznany"),
                        "text": text[:600]
                    })

            if not date_chunks:
                yield sse("done", {
                    "success": True,
                    "timeline": [],
                    "message": "Nie wykryto żadnych dat w badanych dokumentach."
                })
                return

            # Przygotowujemy kontekst — max 15 fragmentów
            context_str = "\n\n".join([
                f"[{c['file']}]: {c['text']}" for c in date_chunks[:15]
            ])

            system_msg = (
                "Jesteś śledczym analitykiem danych. Twoim zadaniem jest znalezienie wszystkich zdarzeń powiązanych z datami "
                "w podanym tekście i przypisanie do nich krótkich, obiektywnych opisów. "
                "Zwróć odpowiedź WYŁĄCZNIE jako poprawny JSON w formacie: "
                "[{\"date\": \"YYYY-MM-DD\", \"event\": \"opis zdarzenia\", \"entity\": \"powiązany podmiot (opcjonalnie)\", \"doc\": \"nazwa pliku\"}]. "
                "Jeśli brakuje dokładnego dnia, użyj YYYY-MM-01. Ważne: Zwróć sam JSON, bez żadnego tekstu pobocznego i bez bloku markdown ```json."
            )

            prompt = f"DOKUMENTY:\n{context_str}\n\nWyciągnij chronologię z dokumentów w formacie JSON:"

            yield sse("progress", {"message": "Analiza LLM…"})
            try:
                result = call_llm(
                    prompt=prompt,
                    system=system_msg,
                    stream=False,
                    provider=llm_provider,
                    model=openrouter_model,
                    max_tokens=1500
                )
                raw_resp = result.get("response", str(result)) if isinstance(result, dict) else str(result)
            except Exception as llm_err:
                logger.exception(f"Timeline LLM error: {llm_err}")
                yield sse("error", {"error": f"Błąd LLM: {str(llm_err)[:200]}"})
                return

            # Parsowanie JSON'a
            json_match = re.search(r'\[\s*\{.*\}\s*\]', raw_resp, re.DOTALL)
            if json_match:
                parsed_timeline = json.loads(json_match.group(0))
                # Sortujemy rosnąco po dacie
                parsed_timeline.sort(key=lambda x: x.get('date', '9999-99-99'))
                yield sse("done", {
                    "success": True,
                    "timeline": parsed_timeline
                })
            else:
                logger.warning(f"Timeline: brak JSON w odpowiedzi LLM: {raw_resp[:200]}")
                yield sse("error", {
                    "error": "LLM nie zwrócił poprawnego formatu JSON.",
                    "raw": raw_resp[:300]
                })

        except Exception as e:
            logger.exception("timeline_endpoint error")
            yield sse("error", {"error": str(e)})

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'},
    )


# ============================================================
# ANALIZA (raporty śledcze + forensyka Excel)
# ============================================================

_EXCEL_ERR_MARKERS = ("#REF!", "#DIV/0!", "#VALUE!", "#NAME?", "#NULL!", "#NUM!", "#N/A", "#NA")


def _fmt_excel_value(v) -> str:
    if v is None:
        return "—"
    if isinstance(v, float):
        if abs(v) >= 1e9 or (0 < abs(v) < 0.0001):
            return f"{v:.6g}"
        return f"{v:,.4f}".replace(",", " ")
    if isinstance(v, int):
        return f"{v:,}".replace(",", " ")
    return str(v)[:120]


def _compress_int_ranges(nums: list[int]) -> str:
    if not nums:
        return ""
    nums = sorted(set(nums))
    parts = []
    start = prev = nums[0]
    for n in nums[1:]:
        if n == prev + 1:
            prev = n
            continue
        parts.append(f"{start}" if start == prev else f"{start}–{prev}")
        start = prev = n
    parts.append(f"{start}" if start == prev else f"{start}–{prev}")
    return ", ".join(parts)


def _excel_goal_seek_hint(stored: float) -> tuple[bool, str, int]:
    """Heurystyka śladu Goal Seek / Solver (nadmiarowa precyzja)."""
    if not isinstance(stored, float) or stored != stored:
        return False, "", 0
    text = f"{stored:.14f}".rstrip("0")
    if "." not in text:
        return False, "", 0
    decimals = len(text.split(".")[1])
    if decimals < 7:
        return False, "", decimals
    if abs(stored) < 1e-12:
        return False, "", decimals
    return (
        True,
        f"Wartość {_fmt_excel_value(stored)} ma {decimals} miejsc po przecinku — typowe dla Goal Seek / Solver.",
        decimals,
    )


def _excel_sum_range(ws_v, range_spec: str) -> tuple[float, int] | None:
    from openpyxl.utils import range_boundaries
    spec = range_spec.strip().replace("$", "")
    if ":" not in spec:
        return None
    try:
        min_col, min_row, max_col, max_row = range_boundaries(spec)
        total = 0.0
        count = 0
        for r in ws_v.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col):
            for c in r:
                if isinstance(c.value, (int, float)):
                    total += float(c.value)
                    count += 1
        return (total, count) if count else None
    except Exception:
        return None


def _excel_eval_simple_formula(formula: str, sheet_values: dict) -> float | None:
    """Proste formuły binarne: =A1+B1, =C3-D3, =E1*E2."""
    m = re.match(r"^=([A-Z]+\d+)\s*([+\-*/])\s*([A-Z]+\d+)$", formula.strip(), re.IGNORECASE)
    if not m:
        return None
    vals = []
    for coord in (m.group(1).upper(), m.group(3).upper()):
        v = sheet_values.get(coord)
        if not isinstance(v, (int, float)):
            return None
        vals.append(float(v))
    op = m.group(2)
    if op == "+":
        return vals[0] + vals[1]
    if op == "-":
        return vals[0] - vals[1]
    if op == "*":
        return vals[0] * vals[1]
    if op == "/" and vals[1] != 0:
        return vals[0] / vals[1]
    return None


def _excel_add_finding(findings: list, summary: dict, **kw) -> None:
    findings.append(kw)
    sev = kw.get("severity", "warning")
    cat = kw.get("category", "")
    if sev == "critical":
        summary["errors"] = summary.get("errors", 0) + 1
    if cat == "override":
        summary["overrides"] = summary.get("overrides", 0) + 1
    if cat == "goal_seek":
        summary["goal_seek"] = summary.get("goal_seek", 0) + 1
    if sev == "warning":
        summary["warnings"] = summary.get("warnings", 0) + 1
    if cat == "excel_error":
        summary["excel_errors"] = summary.get("excel_errors", 0) + 1
    if cat == "external_ref":
        summary["external_refs"] = summary.get("external_refs", 0) + 1


def _excel_forensics(file_path: Path) -> dict:
    """Analiza forensyczna Excel — formuły, cache, ukrycia, odwołania zewnętrzne."""
    findings: list[dict] = []
    summary: dict = {
        "total_formulas": 0,
        "errors": 0,
        "goal_seek": 0,
        "overrides": 0,
        "warnings": 0,
        "hidden_rows": 0,
        "hidden_cols": 0,
        "hidden_sheets": 0,
        "excel_errors": 0,
        "external_refs": 0,
        "critical": 0,
        "high": 0,
    }
    by_sheet: dict[str, dict] = {}

    try:
        wb_f = openpyxl.load_workbook(file_path, data_only=False)
        wb_v = openpyxl.load_workbook(file_path, data_only=True)
    except Exception as e:
        return {"success": False, "error": str(e)}

    props = wb_f.properties
    workbook_meta = {
        "creator": getattr(props, "creator", None) or "—",
        "last_modified_by": getattr(props, "lastModifiedBy", None) or "—",
        "created": str(props.created) if getattr(props, "created", None) else "—",
        "modified": str(props.modified) if getattr(props, "modified", None) else "—",
        "sheet_count": len(wb_f.sheetnames),
        "sheet_names": list(wb_f.sheetnames),
    }

    sheet_values: dict[str, dict] = {}
    for sn in wb_v.sheetnames:
        ws = wb_v[sn]
        sheet_values[sn] = {}
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    sheet_values[sn][cell.coordinate] = cell.value

    calc_values: dict[str, float] = {}
    calc_engine = "none"
    try:
        from xlcalculator import ModelCompiler, Evaluator
        compiler = ModelCompiler()
        model = compiler.read_and_parse_archive(str(file_path))
        evaluator = Evaluator(model)
        for sn in wb_f.sheetnames:
            ws_f = wb_f[sn]
            for row in ws_f.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and cell.value.startswith("="):
                        key = f"{sn}!{cell.coordinate}"
                        try:
                            val = evaluator.evaluate(key)
                            if isinstance(val, (int, float)):
                                calc_values[key] = float(val)
                        except Exception:
                            pass
        if calc_values:
            calc_engine = "xlcalculator"
    except Exception:
        calc_engine = "none"

    for sn in wb_f.sheetnames:
        by_sheet[sn] = {"formulas": 0, "findings": 0, "hidden_rows": 0, "hidden_cols": 0}
        ws_f = wb_f[sn]
        state = getattr(ws_f, "sheet_state", "visible") or "visible"
        if state in ("hidden", "veryHidden"):
            summary["hidden_sheets"] += 1
            _excel_add_finding(
                findings, summary,
                severity="high",
                category="hidden_sheet",
                sheet=sn,
                cell="(arkusz)",
                formula="",
                stored=None,
                calculated=None,
                issue=f"Ukryty arkusz: {sn}",
                detail=f"Stan arkusza: {state}. Dane mogą być niewidoczne w standardowym widoku.",
                recommendation="Odsłoń arkusz w Excelu lub uwzględnij go w analizie.",
            )

    for sn in wb_v.sheetnames:
        ws = wb_v[sn]
        hidden_rows = sorted(ri for ri, rd in ws.row_dimensions.items() if rd.hidden)
        hidden_cols = sorted(
            cl for cl, cd in ws.column_dimensions.items() if cd.hidden
        )

        if hidden_rows:
            summary["hidden_rows"] += len(hidden_rows)
            by_sheet.setdefault(sn, {})["hidden_rows"] = len(hidden_rows)
            samples = []
            for ri in hidden_rows[:12]:
                cells = []
                for cell in ws[ri]:
                    if cell.value is not None and str(cell.value).strip():
                        cells.append(f"{cell.column_letter}={_fmt_excel_value(cell.value)}")
                if cells:
                    samples.append(f"w.{ri}: " + ", ".join(cells[:6]))
            _excel_add_finding(
                findings, summary,
                severity="high",
                category="hidden_rows",
                sheet=sn,
                cell=f"wiersze {_compress_int_ranges(hidden_rows)}",
                formula="",
                stored=None,
                calculated=None,
                issue=f"{len(hidden_rows)} ukrytych wierszy",
                detail=(
                    f"Zakres wierszy: {_compress_int_ranges(hidden_rows)}. "
                    + ("Próbka danych: " + "; ".join(samples[:4]) if samples else "Wiersze bez danych liczbowych/tekstowych.")
                ),
                recommendation="Sprawdź, czy ukryte wiersze nie zawierają kwot pominiętych w sumach widocznych.",
            )

        if hidden_cols:
            summary["hidden_cols"] += len(hidden_cols)
            by_sheet.setdefault(sn, {})["hidden_cols"] = len(hidden_cols)
            _excel_add_finding(
                findings, summary,
                severity="high",
                category="hidden_cols",
                sheet=sn,
                cell=f"kolumny {', '.join(hidden_cols[:15])}{'…' if len(hidden_cols) > 15 else ''}",
                formula="",
                stored=None,
                calculated=None,
                issue=f"{len(hidden_cols)} ukrytych kolumn",
                detail=f"Kolumny: {', '.join(hidden_cols[:20])}. Mogą zawierać dane wyłączone z widoku.",
                recommendation="Zweryfikuj sumy i tabele przestawne pod kątem ukrytych kolumn.",
            )

    for sn in wb_f.sheetnames:
        if sn not in wb_v.sheetnames:
            continue
        ws_f = wb_f[sn]
        ws_v = wb_v[sn]
        flat_values = sheet_values.get(sn, {})

        for row in ws_v.iter_rows():
            for cell in row:
                val = cell.value
                if isinstance(val, str) and any(m in val.upper() for m in _EXCEL_ERR_MARKERS):
                    _excel_add_finding(
                        findings, summary,
                        severity="critical",
                        category="excel_error",
                        sheet=sn,
                        cell=cell.coordinate,
                        formula="",
                        stored=val,
                        calculated=None,
                        issue=f"Błąd Excel w komórce: {val}",
                        detail="Komórka zawiera kod błędu — formuła lub odwołanie jest nieprawidłowe.",
                        recommendation="Przejdź do komórki i napraw odwołanie lub dzielenie przez zero.",
                    )

        for row in ws_f.iter_rows():
            for cell_f in row:
                formula = cell_f.value
                if not isinstance(formula, str) or not formula.startswith("="):
                    continue

                summary["total_formulas"] += 1
                by_sheet.setdefault(sn, {})["formulas"] = by_sheet[sn].get("formulas", 0) + 1
                coord = cell_f.coordinate
                stored = flat_values.get(coord)
                calc_key = f"{sn}!{coord}"
                formula_short = formula if len(formula) <= 120 else formula[:117] + "…"

                if re.search(r"\[[^\]]+\]", formula):
                    _excel_add_finding(
                        findings, summary,
                        severity="warning",
                        category="external_ref",
                        sheet=sn,
                        cell=coord,
                        formula=formula_short,
                        stored=stored,
                        calculated=None,
                        issue="Odwołanie do zewnętrznego pliku",
                        detail="Formuła odwołuje się do innego skoroszytu — wynik zależy od pliku, którego tu nie weryfikujemy.",
                        recommendation="Otwórz powiązany plik i sprawdź, czy wartości są aktualne.",
                    )

                if stored is None:
                    _excel_add_finding(
                        findings, summary,
                        severity="warning",
                        category="no_cache",
                        sheet=sn,
                        cell=coord,
                        formula=formula_short,
                        stored=None,
                        calculated=None,
                        issue="Brak zapisanej wartości (cache)",
                        detail="Excel nie zapisał wyniku formuły (np. LibreOffice, import bez przeliczenia).",
                        recommendation="Otwórz plik w Excelu, Ctrl+Alt+F9 (przelicz) i zapisz ponownie.",
                    )
                    continue

                if isinstance(stored, float):
                    is_gs, gs_detail, dec = _excel_goal_seek_hint(stored)
                    if is_gs:
                        _excel_add_finding(
                            findings, summary,
                            severity="high",
                            category="goal_seek",
                            sheet=sn,
                            cell=coord,
                            formula=formula_short,
                            stored=stored,
                            calculated=None,
                            issue=f"Podejrzenie Goal Seek ({dec} miejsc dziesiętnych)",
                            detail=gs_detail,
                            recommendation="Sprawdź historię zmian komórki i zgodność z założeniami budżetu.",
                        )

                expected: float | None = None
                if calc_key in calc_values:
                    expected = calc_values[calc_key]
                else:
                    simple = _excel_eval_simple_formula(formula, flat_values)
                    if simple is not None:
                        expected = simple
                        calc_engine = calc_engine or "simple"

                if expected is not None and isinstance(stored, (int, float)):
                    try:
                        diff = abs(float(stored) - float(expected))
                        base = max(abs(float(expected)), 1e-9)
                        rel = diff / base
                        if diff > 0.01 and rel > 0.001:
                            _excel_add_finding(
                                findings, summary,
                                severity="critical",
                                category="override",
                                sheet=sn,
                                cell=coord,
                                formula=formula_short,
                                stored=stored,
                                calculated=round(float(expected), 6),
                                diff=round(diff, 4),
                                delta_pct=round(rel * 100, 2),
                                issue="Nadpisanie / rozbieżność cache vs formuła",
                                detail=(
                                    f"W komórce zapisano {_fmt_excel_value(stored)}, "
                                    f"a z formuły wynika {_fmt_excel_value(expected)} "
                                    f"(Δ={_fmt_excel_value(diff)}, {rel*100:.2f}%)."
                                ),
                                recommendation="Porównaj z wersją archiwalną; możliwe ręczne wklejenie wartości.",
                            )
                    except (TypeError, ValueError):
                        pass

                inner = formula.split("(", 1)
                if len(inner) == 2 and inner[0].upper() == "=SUM":
                    args = inner[1].rstrip(")").replace(";", ",")
                    if ":" in args and "!" not in args.split(":")[0]:
                        range_spec = args.split(",")[0].strip()
                        sum_result = _excel_sum_range(ws_v, range_spec)
                        if sum_result is not None and isinstance(stored, (int, float)):
                            total, _ = sum_result
                            diff = abs(float(stored) - total)
                            if diff > 0.02:
                                _excel_add_finding(
                                    findings, summary,
                                    severity="critical",
                                    category="sum_mismatch",
                                    sheet=sn,
                                    cell=coord,
                                    formula=formula_short,
                                    stored=stored,
                                    calculated=round(total, 4),
                                    diff=round(diff, 4),
                                    issue="Błąd sumowania SUM",
                                    detail=(
                                        f"Suma zakresu {range_spec} = {_fmt_excel_value(total)}, "
                                        f"w komórce {_fmt_excel_value(stored)} (Δ={_fmt_excel_value(diff)})."
                                    ),
                                    recommendation="Przelicz zakres ręcznie lub sprawdź ukryte wiersze w SUM.",
                                )

                m_avg = re.match(r"^=AVERAGE\(([^)]+)\)$", formula, re.IGNORECASE)
                if m_avg and isinstance(stored, (int, float)):
                    spec = m_avg.group(1).replace(";", ",").strip()
                    if ":" in spec and "!" not in spec:
                        sum_result = _excel_sum_range(ws_v, spec)
                        if sum_result is not None:
                            total, cnt = sum_result
                            avg = total / cnt
                            diff = abs(float(stored) - avg)
                            if diff > 0.05:
                                _excel_add_finding(
                                    findings, summary,
                                    severity="high",
                                    category="average_mismatch",
                                    sheet=sn,
                                    cell=coord,
                                    formula=formula_short,
                                    stored=stored,
                                    calculated=round(avg, 4),
                                    diff=round(diff, 4),
                                    issue="Błąd średniej AVERAGE",
                                    detail=(
                                        f"Średnia z zakresu ≈ {_fmt_excel_value(avg)}, "
                                        f"zapisano {_fmt_excel_value(stored)}."
                                    ),
                                    recommendation="Sprawdź puste komórki w zakresie — AVERAGE je pomija, SUM nie.",
                                )

    wb_f.close()
    wb_v.close()

    sev_order = {"critical": 0, "high": 1, "warning": 2}
    findings.sort(key=lambda x: (sev_order.get(x.get("severity"), 3), x.get("sheet", ""), x.get("cell", "")))

    for f in findings:
        if f.get("severity") == "critical":
            summary["critical"] += 1
        elif f.get("severity") == "high":
            summary["high"] += 1
        by_sheet.setdefault(f.get("sheet", "?"), {})["findings"] = (
            by_sheet.get(f.get("sheet", "?"), {}).get("findings", 0) + 1
        )

    total_findings = len(findings)
    max_show = 100
    truncated = total_findings > max_show
    if truncated:
        findings = findings[:max_show]

    penalty = summary["critical"] * 12 + summary["high"] * 5 + summary["warnings"] * 2
    confidence_pct = max(8, min(98, 95 - penalty)) if summary["total_formulas"] else 90
    if not findings:
        confidence_pct = 92

    return {
        "success": True,
        "file": file_path.name,
        "summary": summary,
        "findings": findings,
        "findings_total": total_findings,
        "findings_truncated": truncated,
        "workbook_meta": workbook_meta,
        "by_sheet": by_sheet,
        "calc_engine": calc_engine,
        "confidence_pct": confidence_pct,
    }


@app.route('/analyze/excel', methods=['POST'])
def analyze_excel():
    data      = request.get_json()
    win_path  = data.get('win_path', '')
    wsl_path  = data.get('wsl_path', '')
    fname     = data.get('file', '')

    # Ustal ścieżkę WSL
    path = None

    # 1. Bezpośrednia ścieżka WSL
    if wsl_path and Path(wsl_path).exists():
        path = Path(wsl_path)

    # 2. Ścieżka Windows → WSL
    if not path and win_path:
        m = re.match(r'^([A-Za-z]):\\(.*)', win_path)
        if m:
            p = Path(f"/mnt/{m.group(1).lower()}/{m.group(2).replace(chr(92),'/')}")
            if p.exists():
                path = p

    # 3. Szukaj po nazwie rekurencyjnie (bez shella)
    if not path and fname:
        path = _find_file_by_name_safe(fname)

    if not path:
        return jsonify({"success": False, "error": f"Nie znaleziono pliku '{fname}' — zaimportuj go ponownie z opcją śledzenia ścieżek (full_path)."})

    if openpyxl is None:
        return jsonify({"success": False, "error": "Brak biblioteki openpyxl — zainstaluj zależności Python."})

    try:
        path = path.expanduser().resolve()
    except OSError:
        return jsonify({"success": False, "error": "Nieprawidłowa ścieżka"})
    if not path.is_file() or not _path_is_allowed(path):
        return jsonify({"success": False, "error": "Plik niedostępny lub poza dozwolonymi katalogami (SEARCH_ROOTS)"})

    result = _excel_forensics(path)

    # Komentarz LLM przy istotnych znaleziskach
    notable = [
        f for f in result.get("findings", [])
        if f.get("severity") in ("critical", "high")
    ]
    if notable and result.get("success"):
        s = result.get("summary", {})
        meta = result.get("workbook_meta", {})
        ctx_lines = []
        for f in notable[:12]:
            extra = ""
            if f.get("calculated") is not None:
                extra = f" | oczekiwane: {f.get('calculated')}"
            if f.get("delta_pct") is not None:
                extra += f" | Δ%={f.get('delta_pct')}"
            ctx_lines.append(
                f"[{f.get('category','?')}] {f['sheet']}!{f['cell']}: {f['issue']}{extra} — {f.get('detail','')}"
            )
        prompt = (
            f"Plik: {path.name}\n"
            f"Autor: {meta.get('creator')} | Modyfikował: {meta.get('last_modified_by')} | "
            f"Data modyfikacji: {meta.get('modified')}\n"
            f"Formuły: {s.get('total_formulas', 0)} | Krytyczne: {s.get('critical', 0)} | "
            f"Wysokie: {s.get('high', 0)} | Goal Seek: {s.get('goal_seek', 0)} | "
            f"Silnik weryfikacji: {result.get('calc_engine', 'none')}\n\n"
            f"Znaleziska ({len(notable)} istotnych, pierwsze {min(12, len(notable))}):\n"
            + "\n".join(ctx_lines)
            + "\n\nOdpowiedz po polsku (max 12 zdań):\n"
            "1) Najpoważniejsze ryzyko śledcze/finansowe\n"
            "2) Manipulacja vs błąd techniczny\n"
            "3) Co sprawdzić dalej (konkretnie)"
        )
        system = (
            "Jesteś biegłym rewidentem śledczym. Analizujesz anomalie w Excelu. "
            "Pisz konkretnie, bez ogólników."
        )
        try:
            out = call_llm(prompt, system=system, stream=False, max_tokens=900, temperature=0.2)
            if isinstance(out, dict):
                result["llm_comment"] = out.get("response", str(out))
            else:
                result["llm_comment"] = str(out)[:4000]
        except Exception as e:
            result["llm_comment"] = f"(Błąd LLM: {e})"

    return jsonify(result)


# ============================================================

# ============================================================
# KOŃCOWA RECENZJA KODU CAŁEJ APLIKACJI
# Data: 01.06.2026
# Wykonane na prośbę użytkownika
# ============================================================
#
# Przegląd i poprawki wykonane w tej sesji:
#
# 1. Dodano nową zakładkę "Raporty śledcze" z 8 gotowymi szablonami
#    demonstracyjnymi możliwości analitycznych programu (umowy, przepływy,
#    osoby decyzyjne, przetargi, ryzyka, własność itp.).
#
# 2. Wprowadzono podstawowe przetwarzanie partiami w funkcji build_network()
#    (partie po 5 dokumentów) + informacja o liczbie przetworzonych partii.
#
# 3. Ulepszono komunikaty błędów w Sieci powiązań.
#
# 4. Dodano/ulepszono:
#    - Globalny licznik tokenów w topbarze
#    - Lepsze liczniki tokenów przy odpowiedziach LLM
#    - Bardziej widoczne paski postępu przy imporcie/wektoryzacji
#    - Toast po udanej aktualizacji z GitHub
#
# 5. Poprawiono skrypty pomocnicze (start, stop, restart-app.sh, install-user-service.sh)
#    z kolorowymi komunikatami i opcjonalnym git pull przy starcie.
#
# 6. Dodano zarządzanie App API Key w ustawieniach UI.
#
# 7. Zmieniono domyślny provider LLM na OpenRouter (oprócz importu/embeddings).
#
# 8. Dodano i udoskonalono system self-update w modalu diagnostycznym.
#
# 9. Zaktualizowano dokumentację (README.md, AGENTS.md, IMPROVEMENT_PLAN.md).
#
# 10. Przeprowadzono końcową pełną recenzję kodu i naprawiono pozostałości
#     po wcześniejszych agresywnych refaktoryzacjach.
#
# Stan na koniec sesji: aplikacja jest znacznie bardziej przyjazna,
# stabilna przy ciężkich analizach i lepiej przygotowana do codziennego
# użytku w trybie deweloperskim / near-production.
#
# Data zakończenia recenzji: 01 czerwca 2026
# ================================================================

# 2026-06-01 — Opcja A dla problemu z SQL config JSON parse errors (na żądanie użytkownika "A"):

# 2026-06-02 — Security remediation (z raportu przegląd PR #9–#18):
#   - Przywrócono _is_sql_safe, _validate_sql_table_refs, _get_sql_conn i powiązane
#     helpery SQL (usunięte w 3837d5a/ea77923). Bez nich hybrid_stream() z
#     LLM-generated SQL był narażony na injection (teraz fail-closed + walidacja).
#   - Dodano guard PYMSSQL_AVAILABLE w ścieżce hybrydowej.
#   - Drobne poprawki XSS w network error/evidence panel (escapeHtml).
# Najwyższy risk z audytu (latent SQL injection) — zaadresowany.
# - Całkowicie odłączono sqlLoadSavedConn() z window.onload oraz showTab w index.html
# - Zakładka SQL (#navItemSql) jest domyślnie ukryta (display:none)
# - ensureSqlTabVisibility() pokazuje ją automatycznie TYLKO gdy /sql/config ma has_config:true
# - Dodano przycisk w modalu diagnostycznym do ręcznego włączenia (pierwsza konfiguracja)
# - Dodano pełne stub endpointy dla wszystkich /sql/* wywoływanych z UI
# - Zaimplementowano brakujący /health (używa _check_* helpers + sql_available/configured)
#   → kolorowe kropki, staty topbara, bogaty modal diagnostyczny działają poprawnie
# - Dodano PYMSSQL_AVAILABLE + _load_sql_config
# Efekt: błąd "Unexpected token '<'" nigdy nie pojawia się dla użytkowników bez SQL.
# ================================================================

# ============================================================
# MODUŁ SQL — MS SQL Server + SQLite (Text-to-SQL + wektoryzacja do Qdrant)
# ============================================================

@app.route('/sql/config', methods=['GET'])
def sql_config_get():
    """Zwraca zapisaną konfigurację SQL lub informację, że jej nie ma."""
    try:
        cfg = _load_sql_config()
        if cfg and _sql_conn_configured(cfg):
            safe = _redact_sql_config(cfg)
            return jsonify({
                "success": True,
                "has_config": True,
                "config": {
                    "type": safe.get("type", "mssql"),
                    "server": safe.get("server", ""),
                    "port": safe.get("port", 1433),
                    "database": safe.get("database", ""),
                    "user": safe.get("user", ""),
                    "db_path": safe.get("db_path", ""),
                },
            })
        return jsonify({"success": True, "has_config": False})
    except Exception as e:
        return jsonify({"success": False, "has_config": False, "error": str(e)})


@app.route('/sql/config', methods=['POST'])
def sql_config_post():
    """Zapisuje konfigurację SQL (dla UI)."""
    try:
        data = request.get_json() or {}
        existing = _load_sql_config()
        for k in ["type", "server", "port", "database", "user", "db_path"]:
            if k in data:
                existing[k] = data[k]
        if data.get("password"):
            existing["password"] = data["password"]
        _save_sql_config(existing)
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


def _sql_text_columns(conn, table: str, dialect: str) -> list[str]:
    """Kolumny tekstowe/numeryczne — zwykły kursor (pymssql as_dict nie obsługuje metadanych)."""
    cur = conn.cursor()
    try:
        if dialect == "sqlite":
            cur.execute(f'PRAGMA table_info("{table}")')
            text_types = ("TEXT", "CHAR", "CLOB", "VARCHAR", "NVARCHAR", "INT", "INTEGER", "REAL", "FLOAT", "NUMERIC", "DECIMAL", "DATE", "DATETIME")
            return [r[1] for r in cur.fetchall() if any(t in (r[2] or "").upper() for t in text_types)]
        cur.execute("""
            SELECT COLUMN_NAME FROM INFORMATION_SCHEMA.COLUMNS
            WHERE TABLE_NAME=%s AND DATA_TYPE IN
            ('varchar','nvarchar','text','ntext','char','nchar','int','decimal','numeric','float','date','datetime')
            ORDER BY ORDINAL_POSITION
        """, (table,))
        return [r[0] for r in cur.fetchall()]
    finally:
        cur.close()


def _sql_list_tables(conn, dialect: str) -> list[dict]:
    cur = conn.cursor()
    try:
        if dialect == "sqlite":
            cur.execute("""
                SELECT name, type FROM sqlite_master
                WHERE type IN ('table','view') AND name NOT LIKE 'sqlite_%'
                ORDER BY name
            """)
            return [{"name": r[0], "type": r[1].upper()} for r in cur.fetchall()]
        cur.execute("""
            SELECT TABLE_NAME, TABLE_TYPE
            FROM INFORMATION_SCHEMA.TABLES
            WHERE TABLE_TYPE IN ('BASE TABLE','VIEW')
            ORDER BY TABLE_NAME
        """)
        return [{"name": r[0], "type": r[1]} for r in cur.fetchall()]
    finally:
        cur.close()


def _sql_table_count(conn, table: str, dialect: str) -> int:
    qtable = _sql_quote_ident(table, dialect)
    cur = conn.cursor()
    try:
        cur.execute(f"SELECT COUNT(*) AS cnt FROM {qtable}")
        row = cur.fetchone()
        return int(row[0])
    finally:
        cur.close()


def _sql_fetch_table_rows(cur, table: str, columns: list[str], dialect: str, batch: int):
    col_list = ", ".join(_sql_quote_ident(c, dialect) for c in columns)
    qtable = _sql_quote_ident(table, dialect)
    cur.execute(f"SELECT {col_list} FROM {qtable}")
    while True:
        batch_rows = cur.fetchmany(batch)
        if not batch_rows:
            break
        for row in batch_rows:
            if isinstance(row, dict):
                yield row
            elif hasattr(row, "keys"):
                yield {k: row[k] for k in row.keys()}
            else:
                yield {columns[i]: row[i] for i in range(len(columns))}


def _sql_row_to_text(table: str, row: dict, label_col: str = "") -> str:
    label = str(row.get(label_col, "")) if label_col else ""
    parts = [f"{k}: {v}" for k, v in row.items() if v is not None and str(v).strip()]
    prefix = f"[{table}] {label}\n" if label else f"[{table}]\n"
    return prefix + " | ".join(parts)


@app.route('/sql/test', methods=['POST'])
def sql_test():
    """Test połączenia z bazą SQL (MS SQL Server lub SQLite)."""
    cfg = request.get_json() or {}
    try:
        if _sql_dialect(cfg) == "sqlite" and not _sql_conn_configured(cfg):
            return jsonify({
                "success": False,
                "error": "Podaj ścieżkę do pliku .db (np. /mnt/g/dane/eksport.db)",
            })
        _require_sql_backend(cfg)
        dialect = _sql_dialect(cfg)
        conn = _get_sql_conn(cfg, readonly=True)
        cur = conn.cursor()
        if dialect == "sqlite":
            cur.execute("SELECT sqlite_version()")
            version = f"SQLite {cur.fetchone()[0]}"
        else:
            cur.execute("SELECT @@VERSION")
            version = cur.fetchone()[0]
        tables = _sql_list_tables(conn, dialect)
        conn.close()
        return jsonify({"success": True, "version": version[:120], "tables": tables, "dialect": dialect})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/sql/schema', methods=['POST'])
def sql_schema():
    """Pobiera schemat wybranej tabeli."""
    data = request.get_json() or {}
    cfg = data.get("conn", {})
    table = data.get("table", "")
    if not _validate_sql_table_name(table):
        return jsonify({"success": False, "error": "Nieprawidłowa nazwa tabeli"})
    try:
        _require_sql_backend(cfg)
        dialect = _sql_dialect(cfg)
        conn = _get_sql_conn(cfg, readonly=True)
        cur = conn.cursor()
        if dialect == "sqlite":
            cur.execute(f'PRAGMA table_info("{table}")')
            cols = [
                {"name": r[1], "type": r[2], "max_len": None, "nullable": "YES" if r[3] == 0 else "NO"}
                for r in cur.fetchall()
            ]
        else:
            cur.execute("""
                SELECT COLUMN_NAME, DATA_TYPE, CHARACTER_MAXIMUM_LENGTH, IS_NULLABLE
                FROM INFORMATION_SCHEMA.COLUMNS
                WHERE TABLE_NAME = %s
                ORDER BY ORDINAL_POSITION
            """, (table,))
            cols = [
                {"name": r[0], "type": r[1], "max_len": r[2], "nullable": r[3]}
                for r in cur.fetchall()
            ]
        row_count = _sql_table_count(conn, table, dialect)
        conn.close()
        return jsonify({"success": True, "table": table, "columns": cols, "row_count": row_count})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/sql/ask', methods=['POST'])
def sql_ask():
    """Text-to-SQL: pytanie po polsku → LLM → SELECT → wyniki + interpretacja."""
    data = request.get_json() or {}
    cfg = data.get("conn", {})
    question = (data.get("question") or "").strip()
    schema = data.get("schema") or ""

    if not question:
        return jsonify({"success": False, "error": "Brak pytania"})

    try:
        _require_sql_backend(cfg)
        result = _run_text_to_sql_pipeline(
            cfg, question, schema, fetch_limit=500, with_chart=True,
        )
        if not result.get("success"):
            return jsonify(result)

        sql_query = result["sql"]
        result_rows = result["rows"]
        preview = "\n".join([
            ", ".join([f"{k}={v}" for k, v in r.items()])
            for r in result_rows[:10]
        ])
        prompt_interp = (
            f"Pytanie użytkownika: {question}\n"
            f"Zapytanie SQL: {sql_query}\n"
            f"Liczba wyników: {len(result_rows)}\n"
            f"Przykładowe wyniki:\n{preview}\n\n"
            "Odpowiedz po polsku: co wynika z tych danych? Podaj konkretne liczby i wnioski."
        )
        interp = call_llm(
            prompt_interp,
            "Jesteś analitykiem danych. Interpretujesz wyniki SQL po polsku.",
            stream=False,
            max_tokens=1500,
            temperature=0.2,
        )
        result["interpretation"] = _llm_response_text(interp)
        return jsonify(result)
    except Exception as e:
        return jsonify({"success": False, "error": str(e), "sql": ""})


@app.route('/sql/write', methods=['POST'])
def sql_write():
    """Generuje i wykonuje INSERT/UPDATE/DELETE po potwierdzeniu użytkownika (tylko MS SQL)."""
    data = request.get_json() or {}
    cfg = data.get("conn", {})
    if _sql_dialect(cfg) == "sqlite":
        return jsonify({
            "success": False,
            "error": "Tryb zapisu niedostępny dla SQLite — baza jest otwierana tylko do odczytu (mode=ro).",
        })
    if not PYMSSQL_AVAILABLE:
        return jsonify({"success": False, "error": "Moduł pymssql nie jest zainstalowany"})
    question = (data.get("question") or "").strip()
    schema = data.get("schema") or ""
    confirmed = bool(data.get("confirmed", False))
    sql_query = (data.get("sql") or "").strip()

    if not question and not sql_query:
        return jsonify({"success": False, "error": "Brak pytania lub zapytania SQL"})

    try:
        if not sql_query:
            effective_schema, _known = _resolve_sql_schema(cfg, schema)
            system_sql = (
                "Jesteś ekspertem T-SQL (MS SQL Server). "
                "Generujesz zapytania modyfikujące dane: INSERT, UPDATE, DELETE. "
                "Odpowiadasz WYŁĄCZNIE samym SQL — bez wyjaśnień, bez markdown."
            )
            prompt_sql = (
                f"SCHEMAT BAZY:\n{effective_schema}\n\n"
                f"ZADANIE: {question}\n\n"
                "Wygeneruj zapytanie T-SQL (INSERT/UPDATE/DELETE):"
            )
            gen = call_llm(prompt_sql, system_sql, stream=False, max_tokens=2000, temperature=0.1)
            sql_query = _clean_llm_sql(_llm_response_text(gen))

        first_word = sql_query.split()[0].upper() if sql_query.split() else ""
        allowed_write = ("INSERT", "UPDATE", "DELETE", "MERGE")
        if first_word not in allowed_write and first_word not in ("SELECT", "WITH"):
            return jsonify({"success": False, "error": f"Nieznane polecenie: {first_word}", "sql": sql_query})

        if not confirmed:
            impact_info = ""
            if first_word in ("UPDATE", "DELETE"):
                # Sprawdź bezpieczeństwo zapytania PRZED jakimkolwiek wykonaniem
                is_safe_pre, _ = _is_sql_safe(sql_query, ("UPDATE", "DELETE"))
                if is_safe_pre:
                    try:
                        conn_check = _get_sql_conn(cfg)
                        cur_check = conn_check.cursor()
                        if first_word == "UPDATE":
                            where_part = re.search(r"\bWHERE\b(.+?)$", sql_query, re.IGNORECASE | re.DOTALL)
                            table_match = re.search(r"UPDATE\s+\[?(\w+)\]?", sql_query, re.IGNORECASE)
                            if where_part and table_match:
                                table_name = table_match.group(1)
                                # Tylko alphanum/underscore — nie dopuszczamy iniekcji w nazwie tabeli
                                if re.match(r"^\w+$", table_name):
                                    count_sql = f"SELECT COUNT(*) FROM [{table_name}] WHERE {where_part.group(1)}"
                                    cur_check.execute(count_sql)
                                    impact_info = f"Zmieni {cur_check.fetchone()[0]} wierszy"
                        elif first_word == "DELETE":
                            from_match = re.search(r"\bFROM\b(.+)$", sql_query, re.IGNORECASE | re.DOTALL)
                            if from_match:
                                count_sql = f"SELECT COUNT(*) FROM {from_match.group(1)}"
                                cur_check.execute(count_sql)
                                impact_info = f"Usunie {cur_check.fetchone()[0]} wierszy"
                        conn_check.close()
                    except Exception:
                        impact_info = "Nie udało się oszacować wpływu"

            return jsonify({
                "success": True,
                "preview": True,
                "sql": sql_query,
                "impact": impact_info,
                "first_word": first_word,
            })

        is_safe, safe_err = _is_sql_safe(sql_query, allowed_write)
        if not is_safe:
            return jsonify({"success": False, "error": safe_err or "Niedozwolone zapytanie SQL", "sql": sql_query})

        conn = _get_sql_conn(cfg)
        cur = conn.cursor()
        cur.execute(sql_query)
        rows_affected = cur.rowcount
        conn.commit()
        conn.close()
        logger.info("[SQL WRITE] %s | rows=%s | %s", first_word, rows_affected, sql_query[:100])

        return jsonify({
            "success": True,
            "executed": True,
            "sql": sql_query,
            "rows_affected": rows_affected,
            "message": f"Wykonano. Zmieniono/dodano {rows_affected} wierszy.",
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e), "sql": sql_query})


@app.route('/sql/vectorize', methods=['POST'])
def sql_vectorize():
    """Wektoryzuje dane z tabeli SQL do Qdrant (SSE)."""
    data = request.get_json() or {}
    cfg = data.get("conn", {})
    table = data.get("table", "")
    cols = data.get("columns") or []
    label_col = (data.get("label_col") or "").strip()

    if not _validate_sql_table_name(table):
        return jsonify({"success": False, "error": "Nieprawidłowa nazwa tabeli"}), 400

    def generate():
        def sse(event, payload):
            return f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n"

        task_id = ""
        conn = None
        try:
            ok, task = _start_heavy_task("sql_vectorize", f"Wektoryzacja SQL: {table}", {
                "table": table,
                "columns": list(cols),
                "label_col": label_col,
            })
            if not ok:
                yield sse("error", _heavy_task_busy_payload(task))
                return
            task_id = task["id"]

            _require_sql_backend(cfg)
            dialect = _sql_dialect(cfg)
            _ensure_collection_exists()
            conn = _get_sql_conn(cfg, readonly=True)
            data_cur = conn.cursor(as_dict=True) if dialect == "mssql" else conn.cursor()

            if cols:
                active_cols = list(cols)
            else:
                active_cols = _sql_text_columns(conn, table, dialect)
            if not active_cols:
                yield sse("error", {"error": f"Brak kolumn do wektoryzacji w tabeli {table}"})
                _finish_heavy_task(task_id, status="error", error=f"Brak kolumn do wektoryzacji w tabeli {table}")
                task_id = ""
                return

            total = _sql_table_count(conn, table, dialect)
            _update_heavy_task(task_id, total=total, meta={
                "table": table,
                "columns": active_cols,
                "label_col": label_col,
            })
            yield sse("start", {"start": True, "total": total, "table": table, "columns": active_cols, "task_id": task_id})

            qdrant = get_qdrant_client()
            done = 0
            rows_buf: list[str] = []
            BATCH = 50

            for row in _sql_fetch_table_rows(data_cur, table, active_cols, dialect, BATCH):
                rows_buf.append(_sql_row_to_text(table, row, label_col))
                if len(rows_buf) < BATCH:
                    continue

                vectors = get_embeddings_batch(rows_buf, batch_size=6)
                points = []
                for txt, vec in zip(rows_buf, vectors):
                    if not vec or not any(v != 0.0 for v in vec):
                        continue
                    cid = hashlib.md5(txt.encode("utf-8", errors="replace")).hexdigest()
                    points.append(PointStruct(
                        id=cid,
                        vector=vec,
                        payload={"file": f"[SQL] {table}", "text": txt, "full_path": "", "source": "sql", "sql_table": table},
                    ))
                if points:
                    qdrant.upsert(collection_name=ACTIVE_COLLECTION, points=points)
                done += len(rows_buf)
                rows_buf = []
                pct = round(done / total * 100) if total else 0
                _update_heavy_task(task_id, done=done, total=total, progress_pct=pct, current_item=table)
                yield sse("progress", {"progress": True, "done": done, "total": total, "pct": pct})

            if rows_buf:
                vectors = get_embeddings_batch(rows_buf, batch_size=6)
                points = []
                for txt, vec in zip(rows_buf, vectors):
                    if not vec or not any(v != 0.0 for v in vec):
                        continue
                    cid = hashlib.md5(txt.encode("utf-8", errors="replace")).hexdigest()
                    points.append(PointStruct(
                        id=cid,
                        vector=vec,
                        payload={"file": f"[SQL] {table}", "text": txt, "full_path": "", "source": "sql", "sql_table": table},
                    ))
                if points:
                    qdrant.upsert(collection_name=ACTIVE_COLLECTION, points=points)
                done += len(rows_buf)
                pct = round(done / total * 100) if total else 0
                _update_heavy_task(task_id, done=done, total=total, progress_pct=pct, current_item=table)
                yield sse("progress", {"progress": True, "done": done, "total": total, "pct": pct})

            if conn:
                conn.close()
                conn = None
            _docs_cache["data"] = None
            _update_heavy_task(task_id, done=done, total=total, progress_pct=100, current_item=None)
            yield sse("done", {"done": True, "task_id": task_id, "msg": f"Zwektoryzowano {done} wierszy z tabeli [{table}]"})
        except Exception as e:
            logger.exception("sql_vectorize error")
            yield sse("error", {"error": str(e), "task_id": task_id or None})
            if task_id:
                _finish_heavy_task(task_id, status="error", error=str(e)[:500])
                task_id = ""
        finally:
            if conn:
                try:
                    conn.close()
                except Exception:
                    logger.debug("Nie zamknięto połączenia SQL po błędzie", exc_info=True)
            if task_id:
                _finish_heavy_task(task_id)

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route('/sql/vectorize-all', methods=['POST'])
def sql_vectorize_all():
    """Wektoryzuje wszystkie tabele SQL do Qdrant (SSE)."""
    data = request.get_json() or {}
    cfg = data.get("conn", {})

    def generate():
        def sse(event, payload):
            return f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n"

        task_id = ""
        conn = None
        try:
            ok, task = _start_heavy_task("sql_vectorize_all", "Wektoryzacja wszystkich tabel SQL", {})
            if not ok:
                yield sse("error", _heavy_task_busy_payload(task))
                return
            task_id = task["id"]

            _require_sql_backend(cfg)
            dialect = _sql_dialect(cfg)
            _ensure_collection_exists()
            conn = _get_sql_conn(cfg, readonly=True)
            data_cur = conn.cursor(as_dict=True) if dialect == "mssql" else conn.cursor()

            tables = [t["name"] for t in _sql_list_tables(conn, dialect)]
            _update_heavy_task(task_id, total=len(tables), meta={"tables": tables})
            yield sse("start", {"start": True, "total_tables": len(tables), "tables": tables, "task_id": task_id})

            qdrant = get_qdrant_client()
            total_rows = 0

            for table_idx, table in enumerate(tables):
                try:
                    yield sse("table_start", {
                        "table_start": True,
                        "table": table,
                        "table_idx": table_idx,
                        "total_tables": len(tables),
                    })
                    _update_heavy_task(
                        task_id,
                        done=table_idx,
                        total=len(tables),
                        progress_pct=round(table_idx / len(tables) * 100) if tables else 0,
                        current_item=table,
                    )

                    active_cols = _sql_text_columns(conn, table, dialect)
                    if not active_cols:
                        yield sse("table_done", {"table_done": True, "table": table, "rows": 0, "table_idx": table_idx})
                        continue

                    table_total = _sql_table_count(conn, table, dialect)
                    if table_total == 0:
                        yield sse("table_done", {"table_done": True, "table": table, "rows": 0, "table_idx": table_idx})
                        continue

                    table_done = 0
                    rows_buf: list[str] = []
                    BATCH = 50

                    for row in _sql_fetch_table_rows(data_cur, table, active_cols, dialect, BATCH):
                        rows_buf.append(_sql_row_to_text(table, row))
                        if len(rows_buf) < BATCH:
                            continue

                        vectors = get_embeddings_batch(rows_buf, batch_size=6)
                        points = []
                        for txt, vec in zip(rows_buf, vectors):
                            if not vec or not any(v != 0.0 for v in vec):
                                continue
                            cid = hashlib.md5(txt.encode("utf-8", errors="replace")).hexdigest()
                            points.append(PointStruct(
                                id=cid,
                                vector=vec,
                                payload={"file": f"[SQL] {table}", "text": txt, "full_path": "", "source": "sql", "sql_table": table},
                            ))
                        if points:
                            qdrant.upsert(collection_name=ACTIVE_COLLECTION, points=points)

                        table_done += len(rows_buf)
                        total_rows += len(rows_buf)
                        rows_buf = []
                        pct = round(table_done / table_total * 100) if table_total else 0
                        _update_heavy_task(
                            task_id,
                            done=table_idx,
                            total=len(tables),
                            progress_pct=round((table_idx + (pct / 100)) / len(tables) * 100) if tables else 100,
                            current_item=f"{table} ({pct}%)",
                        )
                        yield sse("progress", {
                            "progress": True,
                            "table": table,
                            "done": table_done,
                            "total": table_total,
                            "pct": pct,
                        })

                    if rows_buf:
                        vectors = get_embeddings_batch(rows_buf, batch_size=6)
                        points = []
                        for txt, vec in zip(rows_buf, vectors):
                            if not vec or not any(v != 0.0 for v in vec):
                                continue
                            cid = hashlib.md5(txt.encode("utf-8", errors="replace")).hexdigest()
                            points.append(PointStruct(
                                id=cid,
                                vector=vec,
                                payload={"file": f"[SQL] {table}", "text": txt, "full_path": "", "source": "sql", "sql_table": table},
                            ))
                        if points:
                            qdrant.upsert(collection_name=ACTIVE_COLLECTION, points=points)
                        table_done += len(rows_buf)
                        total_rows += len(rows_buf)

                    yield sse("table_done", {
                        "table_done": True,
                        "table": table,
                        "rows": table_done,
                        "table_idx": table_idx,
                    })
                    _update_heavy_task(
                        task_id,
                        done=table_idx + 1,
                        total=len(tables),
                        progress_pct=round((table_idx + 1) / len(tables) * 100) if tables else 100,
                        current_item=None,
                    )
                except Exception as e:
                    yield sse("table_error", {"table_error": True, "table": table, "error": str(e)})

            if conn:
                conn.close()
                conn = None
            _docs_cache["data"] = None
            _update_heavy_task(task_id, done=len(tables), total=len(tables), progress_pct=100, current_item=None)
            yield sse("done", {
                "done": True,
                "task_id": task_id,
                "total_rows": total_rows,
                "total_tables": len(tables),
                "msg": f"Zwektoryzowano {total_rows} wierszy z {len(tables)} tabel",
            })
        except Exception as e:
            logger.exception("sql_vectorize_all error")
            yield sse("error", {"error": str(e), "task_id": task_id or None})
            if task_id:
                _finish_heavy_task(task_id, status="error", error=str(e)[:500])
                task_id = ""
        finally:
            if conn:
                try:
                    conn.close()
                except Exception:
                    logger.debug("Nie zamknięto połączenia SQL po błędzie", exc_info=True)
            if task_id:
                _finish_heavy_task(task_id)

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ============================================================
# /health — endpoint dla diagnostyki (kolorowe kropki, modal startowy,
# self-update, statusy). Używa istniejących _check_* helpers.
# ============================================================


@app.route('/health', methods=['GET'])
def health():
    """Zwraca bogaty status systemu dla UI (dashboard + modal diagnostyczny)."""
    light = request.args.get("light", "0") == "1"
    try:
        q = _check_qdrant_health()
        ocr = _ocr_health_status()
        parsers = _file_parsers_health() if not light else {}
        ollama_h = {"ok": True} if light else _check_ollama_health()
        or_h = {"ok": True} if light else _check_openrouter_health()

        # LLM status (preferuj OpenRouter jeśli skonfigurowany)
        llm_ok = False
        llm_detail = ""
        if light:
            llm_ok = True
            llm_detail = "light"
        elif OPENROUTER_API_KEY:
            llm_ok = or_h.get("ok", False)
            llm_detail = or_h.get("error") or "OpenRouter"
        else:
            llm_ok = ollama_h.get("ok", False)
            llm_detail = ollama_h.get("error") or "Ollama"

        # SQL
        sql_cfg = _load_sql_config()
        sql_configured = _sql_conn_configured(sql_cfg)
        sql_status = "ok" if (
            sql_configured and (_sql_dialect(sql_cfg) == "sqlite" or PYMSSQL_AVAILABLE)
        ) else ("warn" if sql_configured else "absent")

        # Wektory / kolekcja
        vectors = 0
        coll_name = ACTIVE_COLLECTION
        try:
            if q.get("ok") and q.get("active_collection_exists"):
                vectors = q.get("points_in_active", 0)
        except Exception:
            pass

        # Ogólny status
        critical_ok = q.get("ok", False) and (llm_ok or True)  # LLM fallbacki istnieją
        overall = "ok" if critical_ok else "degraded"

        # Wersja (z git describe lub stałej)
        try:
            ver = _get_app_version()
        except Exception:
            ver = "dev"

        return jsonify({
            "success": True,
            "overall": overall,
            "version": ver,
            "timestamp": int(time.time()),

            # Połączenia (dla kropek i tabeli w modalu)
            "qdrant": q,
            "qdrant_status": "ok" if q.get("ok") else "error",
            "llm": {"ok": llm_ok, "detail": llm_detail},
            "llm_status": "ok" if llm_ok else "warn",
            "embedding": {
                "ok": ollama_h.get("ok", False),  # embeddingi idą przez Ollama
                "model": "nomic-embed-text"
            },
            "ocr": ocr,
            "ocr_available": ocr.get("available", False),
            "file_parsers": parsers,

            # SQL (opcjonalny)
            "sql_available": PYMSSQL_AVAILABLE,
            "sql_configured": sql_configured,
            "sql_status": sql_status,

            # Aktywna kolekcja + wektory (dla topbara)
            "active_collection": {
                "name": coll_name,
                "points": vectors
            },
            "vectors_count": vectors,

            # Inne przydatne dla UI
            "provider": DEFAULT_LLM_PROVIDER,
            "llm_model": _effective_llm_model(DEFAULT_LLM_PROVIDER),
            "gemini_configured": bool(GEMINI_API_KEY),
            "gemini_model": GEMINI_MODEL,
            "gemini": _check_gemini_health(),
            "app_api_key_set": bool(APP_API_KEY),
        })
    except Exception as e:
        logger.exception("health endpoint error")
        return jsonify({
            "success": False,
            "overall": "error",
            "error": str(e)[:200]
        }), 500


@app.route('/api/update/status', methods=['GET'])
def api_update_status():
    """Zwraca status aktualizacji — aktualna wersja, najnowsza, changelog."""
    if not _localhost_only():
        return jsonify({"success": False, "error": "Dostępne tylko z localhost"}), 403
    try:
        local_tag = _get_local_latest_tag()
        remote_tag = _get_remote_latest_tag()
        release_info = _get_latest_github_release() or {}

        if release_info.get("tag_name"):
            remote_tag = release_info["tag_name"]

        def _version_key(v: str) -> tuple:
            try:
                return tuple(int(x) for x in v.lstrip("v").split("."))
            except Exception:
                return (0, 0)

        update_available = bool(
            remote_tag and local_tag and _version_key(remote_tag) > _version_key(local_tag)
        )

        return jsonify({
            "success": True,
            "local_version": local_tag,
            "remote_version": remote_tag or local_tag,
            "current_version": local_tag,
            "latest_version": remote_tag or local_tag,
            "update_available": update_available,
            "release": release_info,
            "changelog": release_info.get("body", ""),
            "release_url": release_info.get("html_url"),
            "message": "Nowa wersja dostępna" if update_available else "Jesteś na najnowszej wersji",
        })
    except Exception as e:
        logger.exception("api_update_status error")
        return jsonify({"success": False, "error": str(e)[:200]}), 500


@app.route('/api/update/pull', methods=['POST'])
def api_update_pull():
    """Pobiera najnowszy kod z GitHub (git pull origin master)."""
    if not _localhost_only():
        return jsonify({"success": False, "error": "Dostępne tylko z localhost"}), 403
    try:
        result = _git_pull()
        ok = result.get("success", False)
        msg = result.get("stdout") or result.get("stderr") or (
            "Kod zaktualizowany. Zalecany restart aplikacji." if ok else "Błąd podczas aktualizacji."
        )
        return jsonify({
            "success": ok,
            "message": msg,
            "output": result.get("stdout", ""),
            "error": result.get("stderr", "") if not ok else "",
            "details": result,
        })
    except Exception as e:
        logger.exception("api_update_pull error")
        return jsonify({"success": False, "error": str(e)[:200]}), 500


@app.route('/api/update/restart', methods=['POST'])
def api_update_restart():
    """Restartuje aplikację (poprzez systemd --user service)."""
    if not _localhost_only():
        return jsonify({"success": False, "error": "Dostępne tylko z localhost"}), 403
    try:
        result = _try_restart_service()
        return jsonify({
            "success": result.get("success", False),
            "message": result.get("message", "Restart initiated"),
            "method": result.get("method", "unknown"),
        })
    except Exception as e:
        logger.exception("api_update_restart error")
        return jsonify({"success": False, "error": str(e)[:200]}), 500


def _localhost_only() -> bool:
    """True gdy żądanie pochodzi z localhost (self-update / service API).
    Nagłówki X-Forwarded-For są ufane tylko gdy TRUST_PROXY=true.
    """
    addr = (request.remote_addr or "").strip()
    if addr in ("127.0.0.1", "::1", "localhost"):
        return True
    if TRUST_PROXY:
        fwd = (request.headers.get("X-Forwarded-For") or "").split(",")[0].strip()
        if fwd in ("127.0.0.1", "::1", "localhost"):
            return True
    return False


def _systemd_service_state() -> str:
    """Stan usługi ai_analiza — najpierw user service, potem systemowa."""
    import shutil
    if not shutil.which("systemctl"):
        return "dev_mode"
    for cmd in (
        ["systemctl", "--user", "is-active", "ai_analiza"],
        ["systemctl", "is-active", "ai_analiza"],
    ):
        try:
            st = subprocess.run(cmd, capture_output=True, text=True, timeout=3)
            val = (st.stdout or "").strip()
            if val:
                return val
        except Exception:
            pass
    if os.environ.get("INVOCATION_ID"):
        return "active"
    return "unknown"


def _systemd_service_logs(limit: int = 30) -> str:
    for cmd in (
        ["journalctl", "--user", "-u", "ai_analiza", "-n", str(limit), "--no-pager", "--output=short"],
        ["journalctl", "-u", "ai_analiza", "-n", str(limit), "--no-pager", "--output=short"],
    ):
        try:
            logs = subprocess.run(cmd, capture_output=True, text=True, timeout=5)
            if logs.stdout:
                return logs.stdout
        except Exception:
            pass
    return ""


@app.route('/api/service/status', methods=['GET'])
def service_status():
    """Status usługi systemd i ostatnie logi — tylko localhost."""
    if not _localhost_only():
        return jsonify({"success": False, "error": "Dostępne tylko z localhost"}), 403

    return jsonify({
        "success": True,
        "is_systemd": bool(os.environ.get("INVOCATION_ID")),
        "active": _systemd_service_state(),
        "logs": _systemd_service_logs(30),
    })


@app.route('/api/service/restart', methods=['POST'])
def service_restart():
    """Restart usługi ai_analiza — tylko localhost."""
    if not _localhost_only():
        return jsonify({"success": False, "error": "Dostępne tylko z localhost"}), 403

    def _do_restart():
        time.sleep(0.6)
        try:
            import shutil
            if shutil.which("systemctl"):
                for cmd in (
                    ["systemctl", "--user", "restart", "ai_analiza"],
                    ["systemctl", "restart", "ai_analiza"],
                ):
                    try:
                        result = subprocess.run(cmd, capture_output=True, text=True, timeout=15)
                        if result.returncode == 0:
                            return
                    except Exception:
                        continue
            os._exit(0)
        except Exception:
            os._exit(0)

    threading.Thread(target=_do_restart, daemon=True).start()
    return jsonify({"success": True, "msg": "Restart zlecony"})
