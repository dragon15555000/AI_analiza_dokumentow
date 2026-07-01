"""
Routes dla audytu finansowego — analiza formuł arkuszy.
"""

import logging
import tempfile
import math
import json
import hashlib
import io
from datetime import datetime
from pathlib import Path
from collections import Counter, defaultdict, deque, OrderedDict
import re
import zipfile
import xml.etree.ElementTree as ET

from flask import Blueprint, request, send_file

from financial_audit import AuditAnalyzer, CellRef, FormulaParser
from llm_client import call_llm, _llm_response_text, LLM_MODEL
from prompts import SEARCH_MODES
from services.financial_forensics import (
    detect_control_total_and_hidden_reference_signals as svc_detect_control_total_and_hidden_reference_signals,
    detect_duplicate_and_numeric_signals as svc_detect_duplicate_and_numeric_signals,
    detect_hardcoded_values_and_pattern_deviations as svc_detect_hardcoded_values_and_pattern_deviations,
)
from services.ocr_cross_check import extract_candidates as svc_extract_ocr_candidates
from services.ocr_cross_check import run_cross_check as svc_run_ocr_cross_check
from utils.http import json_error, json_success

logger = logging.getLogger("ai_analiza")

financial_bp = Blueprint("financial", __name__, url_prefix="/api/audit")
_FINANCIAL_AUDIT_CACHE_MAX = 6
_FINANCIAL_AUDIT_CACHE: OrderedDict[str, dict] = OrderedDict()
_FINANCIAL_AI_OPINION_PROVIDER = "ollama"
_FINANCIAL_AI_OPINION_MAX_FINDINGS = 12


def _normalize_analysis_type(value: str) -> str:
    normalized = (value or "").strip().lower()
    if normalized in {"full", "deep", "complete"}:
        return "full"
    if normalized in {"formulas", "formula"}:
        return "formulas"
    if normalized in {"quick", "fast", "screen", "screening"}:
        return "quick"
    return "targeted"


def _pl_count(value: int, one: str, few: str, many: str) -> str:
    value = abs(int(value))
    if value == 1:
        return one
    if value % 100 in (12, 13, 14):
        return many
    if value % 10 in (2, 3, 4):
        return few
    return many


def _format_dt(value) -> str | None:
    if not value:
        return None
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value)


def _format_size_human(size_bytes: int) -> str:
    units = ["B", "KB", "MB", "GB"]
    value = float(size_bytes)
    for unit in units:
        if value < 1024 or unit == units[-1]:
            if unit == "B":
                return f"{int(value)} {unit}"
            return f"{value:.1f} {unit}"
        value /= 1024
    return f"{size_bytes} B"


def _serialize_excel_value(value):
    if value is None:
        return None
    if hasattr(value, "isoformat"):
        return value.isoformat()
    if isinstance(value, (int, float, bool, str)):
        return value
    return str(value)


def _value_to_text(value) -> str:
    if value is None:
        return ""
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value).strip()


def _coerce_number(value):
    if isinstance(value, bool) or value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = _value_to_text(value)
    if not text:
        return None
    text = text.replace(" ", "").replace("\xa0", "").replace(",", ".")
    try:
        return float(text)
    except ValueError:
        return None


def _coerce_datetime(value):
    if value is None:
        return None
    if hasattr(value, "hour") and hasattr(value, "weekday"):
        return value
    text = _value_to_text(value)
    if not text:
        return None
    for parser in (datetime.fromisoformat,):
        try:
            return parser(text)
        except ValueError:
            continue
    return None


def _extract_sequence_number(value) -> int | None:
    text = _value_to_text(value)
    if not text:
        return None
    matches = re.findall(r"\d+", text)
    if not matches:
        return None
    try:
        return int(matches[-1])
    except ValueError:
        return None


def _header_kind(header: str) -> str:
    h = (header or "").strip().lower()
    if re.search(r"(faktur|invoice|dokument|numer|nr|rachun|pozycj|id)", h):
        return "document"
    if re.search(r"(kontrah|vendor|dostawc|odbior|klient|firma|nazwa)", h):
        return "party"
    if re.search(r"(kwot|amount|wart|netto|brutto|saldo|koszt|suma|payment|przelew)", h):
        return "amount"
    if re.search(r"(data|date|godzin|czas)", h):
        return "datetime"
    return "generic"


def _calc_percentile(sorted_values: list[float], percentile: float) -> float | None:
    if not sorted_values:
        return None
    if len(sorted_values) == 1:
        return sorted_values[0]
    position = (len(sorted_values) - 1) * percentile
    lower = math.floor(position)
    upper = math.ceil(position)
    if lower == upper:
        return sorted_values[lower]
    weight = position - lower
    return sorted_values[lower] * (1 - weight) + sorted_values[upper] * weight


def _benford_expected() -> dict[int, float]:
    return {digit: math.log10(1 + 1 / digit) for digit in range(1, 10)}


def _leading_digit(value) -> int | None:
    number = _coerce_number(value)
    if number is None:
        return None
    number = abs(number)
    if number < 1:
        return None
    while number >= 10:
        number /= 10
    return int(str(number)[0])


def _risk_level_from_anomalies(anomalies: list[dict]) -> str:
    high = sum(1 for anomaly in anomalies if anomaly.get("severity") == "HIGH")
    medium = sum(1 for anomaly in anomalies if anomaly.get("severity") == "MEDIUM")
    low = sum(1 for anomaly in anomalies if anomaly.get("severity") == "LOW")
    if high > 0:
        return "CRITICAL"
    if medium >= 3:
        return "HIGH"
    if medium > 0:
        return "MEDIUM"
    if low > 0:
        return "LOW"
    return "LOW"


def _normalize_formula_pattern(formula: str, current_row: int, current_col: int) -> str:
    from openpyxl.utils import column_index_from_string

    def repl(match):
        sheet = (match.group("sheet") or "").replace("$", "")
        col = match.group("col").replace("$", "")
        row = int(match.group("row"))
        col_offset = column_index_from_string(col) - current_col
        row_offset = row - current_row
        prefix = f"{sheet}!" if sheet else ""
        return f"{prefix}C[{col_offset}]R[{row_offset}]"

    expr = (formula or "").strip().upper().replace(";", ",")
    expr = re.sub(
        r"(?:(?P<sheet>'[^']+'|[A-Z0-9_]+)!)?\$?(?P<col>[A-Z]{1,3})\$?(?P<row>\d+)",
        repl,
        expr,
    )
    return re.sub(r"\s+", "", expr)


def _split_excel_args(args_text: str) -> list[str]:
    args: list[str] = []
    current: list[str] = []
    depth = 0
    in_string = False

    for ch in args_text or "":
        if ch == '"':
            in_string = not in_string
            current.append(ch)
        elif not in_string and ch == "(":
            depth += 1
            current.append(ch)
        elif not in_string and ch == ")":
            depth = max(0, depth - 1)
            current.append(ch)
        elif not in_string and depth == 0 and ch in {",", ";"}:
            args.append("".join(current).strip())
            current = []
        else:
            current.append(ch)

    tail = "".join(current).strip()
    if tail:
        args.append(tail)
    return args


def _load_excel_app_metadata(file_path: Path) -> dict:
    try:
        with zipfile.ZipFile(file_path) as zf:
            with zf.open("docProps/app.xml") as app_file:
                root = ET.fromstring(app_file.read())
    except Exception:
        return {}

    ns = {"ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}

    def get(tag: str):
        node = root.find(f"ep:{tag}", ns)
        return node.text if node is not None else None

    return {
        "application": get("Application"),
        "app_version": get("AppVersion"),
        "company": get("Company"),
        "manager": get("Manager"),
    }


def _append_finding(
    findings: list,
    appendix: list,
    *,
    finding_type: str,
    severity: str,
    sheet: str,
    message: str,
    cell: str = "",
    value=None,
    formula: str = "",
    comment: str = "",
    details: dict | None = None,
):
    finding = {
        "type": finding_type,
        "severity": severity,
        "sheet": sheet,
        "cell": cell,
        "message": message,
        "value": _serialize_excel_value(value),
        "formula": formula or None,
        "comment": comment or message,
    }
    if details:
        finding.update(details)
    findings.append(finding)
    appendix.append(
        {
            "sheet": sheet,
            "cell": cell or "—",
            "value": _serialize_excel_value(value),
            "formula": formula or "",
            "severity": severity,
            "type": finding_type,
            "comment": comment or message,
        }
    )


def _anomaly_details(anomaly: dict) -> dict:
    anomaly_type = anomaly.get("type", "")
    mapping = {
        "circular_reference": {
            "label": "Zależność cykliczna",
            "severity_label": "Wysokie ryzyko",
            "impact": "Wynik może zależeć od samego siebie lub generować niestabilne przeliczenia.",
            "recommendation": "Prześledź łańcuch zależności i rozbij cykl na osobne etapy obliczeń.",
            "fraud_hypothesis": "Logika obliczeń została spięta w pętlę, więc wynik nie ma stabilnej i przejrzystej podstawy.",
            "intent_hypothesis": "Możliwe zaciemnienie sposobu liczenia albo niekontrolowana zmiana modelu.",
            "verification_target": "Prześledź pełny cykl zależności i sprawdź, która komórka wprowadziła zapętlenie.",
            "confidence_basis": "high_direct_logic_break",
        },
        "hidden_formula": {
            "label": "Ukryta formuła",
            "severity_label": "Średnie ryzyko",
            "impact": "Ważna logika może być niewidoczna dla odbiorcy raportu lub kontrolera.",
            "recommendation": "Sprawdź, czy ukrycie było zamierzone i czy komórka nie wpływa na wynik końcowy bez wyjaśnienia.",
        },
        "external_reference": {
            "label": "Odwołanie zewnętrzne",
            "severity_label": "Średnie ryzyko",
            "impact": "Wynik zależy od danych spoza bieżącego pliku, co utrudnia audyt i replikację obliczeń.",
            "recommendation": "Zweryfikuj źródło zewnętrzne i rozważ wciągnięcie danych do kontrolowanego arkusza wejściowego.",
        },
        "cosmetic_check": {
            "label": "Pozorna kontrola",
            "severity_label": "Średnie ryzyko",
            "impact": "Formuła wygląda jak walidacja lub test zgodności, ale logicznie nie ma szans wykryć błędu.",
            "recommendation": "Sprawdź, czy warunek naprawdę porównuje niezależne dane i czy może zwrócić negatywny wynik w realnym scenariuszu.",
        },
        "hardcoded_in_formula_region": {
            "label": "Twarda wartość w obszarze formuł",
            "severity_label": "Wysokie ryzyko",
            "impact": "Wynik w kolumnie obliczeniowej został wpisany ręcznie zamiast wyliczony, co może ukrywać manipulację albo błąd.",
            "recommendation": "Porównaj z sąsiednimi wierszami i ustal, czy komórka nie nadpisuje standardowego wzorca obliczeń.",
            "fraud_hypothesis": "Użytkownik ręcznie nadpisał komórkę w obszarze, który normalnie liczy się formułami.",
            "intent_hypothesis": "Celowe narzucenie z góry założonego wyniku bez oparcia w danych źródłowych.",
            "verification_target": "Sprawdź, czy wpisana wartość daje korzystniejszy wynik niż wynik z sąsiedniego wzorca formuł.",
            "confidence_basis": "high_manual_override_in_formula_region",
        },
        "hardcoded_output_like_value": {
            "label": "Ręcznie wpisany wynik końcowy",
            "severity_label": "Wysokie ryzyko",
            "impact": "Komórka wygląda jak wynik końcowy raportu lub tabeli, ale nie liczy się formułą, tylko została wpisana ręcznie.",
            "recommendation": "Ustal źródło tej wartości i sprawdź, czy nie zastępuje automatycznego wyniku obliczeń.",
            "fraud_hypothesis": "Końcowy wynik został wpisany ręcznie zamiast wyliczony automatycznie.",
            "intent_hypothesis": "Próba zafałszowania wyniku końcowego z pominięciem logiki arkusza.",
            "verification_target": "Załącz dane wejściowe i przelicz, jaki powinien być prawdziwy wynik końcowy.",
            "confidence_basis": "high_manual_final_result",
        },
        "formula_pattern_deviation": {
            "label": "Odstępstwo wzorca formuły",
            "severity_label": "Średnie ryzyko",
            "impact": "Jedna formuła w kolumnie różni się od dominującego wzorca i może liczyć coś innego niż pozostałe wiersze.",
            "recommendation": "Sprawdź, czy odmienna logika jest uzasadniona biznesowo, czy wynika z ręcznej zmiany odwołań.",
        },
        "hidden_rows": {
            "label": "Ukryte wiersze",
            "severity_label": "Średnie ryzyko",
            "impact": "Część danych może być pominięta podczas zwykłego przeglądu arkusza.",
            "recommendation": "Odsłoń wskazane wiersze i potwierdź, że nie zawierają pominiętych transakcji lub wyjątków.",
        },
        "hidden_columns": {
            "label": "Ukryte kolumny",
            "severity_label": "Średnie ryzyko",
            "impact": "Ukryte pola mogą zawierać dodatkowe parametry, korekty albo wewnętrzne komentarze wpływające na interpretację danych.",
            "recommendation": "Przejrzyj zawartość ukrytych kolumn i ich wpływ na obliczenia lub raportowanie.",
        },
        "merged_cells": {
            "label": "Scalone komórki",
            "severity_label": "Niska uwaga",
            "impact": "Scalenia utrudniają filtrowanie, porównywanie i automatyczne mapowanie danych źródłowych.",
            "recommendation": "Sprawdź, czy scalone zakresy nie maskują braków w danych albo ręcznych opisów wyjątków.",
        },
        "very_hidden_sheet": {
            "label": "Arkusz very-hidden",
            "severity_label": "Wysokie ryzyko",
            "impact": "Arkusz jest ukryty głębiej niż standardowe hidden i nie pojawia się w zwykłym interfejsie Excela.",
            "recommendation": "Zweryfikuj zawartość arkusza oraz powód jego ukrycia w konfiguracji skoroszytu.",
            "fraud_hypothesis": "Część logiki lub danych została ukryta w arkuszu niewidocznym w standardowym widoku Excela.",
            "intent_hypothesis": "Możliwe ukrycie kluczowych danych przed zwykłą kontrolą arkusza.",
            "verification_target": "Odkryj arkusz very-hidden i sprawdź, czy zasila wyniki albo przechowuje wyjątkowe dane wejściowe.",
            "confidence_basis": "high_hidden_supporting_sheet",
        },
        "sheet_protection": {
            "label": "Ochrona arkusza",
            "severity_label": "Niska uwaga",
            "impact": "Ochrona może być uzasadniona, ale utrudnia szybkie sprawdzenie logiki i wyjątków.",
            "recommendation": "Sprawdź, które pola są chronione i czy nie blokuje to rewizji kluczowych komórek.",
        },
        "active_filter": {
            "label": "Aktywny filtr",
            "severity_label": "Niska uwaga",
            "impact": "Widok danych może być zawężony i ukrywać część rekordów przed odbiorcą raportu.",
            "recommendation": "Przed analizą źródłową wyczyść filtry i porównaj liczbę rekordów z pełnym zestawem.",
        },
        "duplicate_document": {
            "label": "Duplikat numeru dokumentu",
            "severity_label": "Wysokie ryzyko",
            "impact": "Powielony numer dokumentu może wskazywać na duplikację wpisów, korektę bez śladu albo błąd ewidencji.",
            "recommendation": "Porównaj wskazane wiersze z dokumentami źródłowymi i historią księgowania.",
            "fraud_hypothesis": "Ten sam numer dokumentu może wskazywać na zdublowanie wpisu.",
            "intent_hypothesis": "Próba sztucznego zawyżenia kosztów lub przychodów przez podwójne księgowanie.",
            "verification_target": "Poproś o fizyczne lub źródłowe kopie dokumentów dla wszystkich zduplikowanych wpisów.",
            "confidence_basis": "high_duplicate_document_identifier",
        },
        "duplicate_party": {
            "label": "Powtarzający się kontrahent",
            "severity_label": "Niska uwaga",
            "impact": "Sama powtarzalność nie jest dowodem błędu, ale może pomóc w grupowaniu powiązanych transakcji.",
            "recommendation": "Sprawdź, czy powtarzające się wpisy nie łączą się z duplikatami kwot albo numerów dokumentów.",
        },
        "duplicate_amount": {
            "label": "Powtarzająca się kwota",
            "severity_label": "Średnie ryzyko",
            "impact": "Powtarzające się identyczne kwoty mogą wskazywać na sztuczne dzielenie lub kopiowanie zapisów.",
            "recommendation": "Zweryfikuj, czy te same kwoty nie pojawiają się z tym samym kontrahentem lub dokumentem.",
        },
        "numbering_gap": {
            "label": "Luka w numeracji",
            "severity_label": "Średnie ryzyko",
            "impact": "Braki w ciągu numerów mogą oznaczać pominięte lub usunięte dokumenty.",
            "recommendation": "Porównaj ciąg numerów z rejestrem źródłowym i ustal, czy brakujące pozycje są udokumentowane.",
        },
        "near_threshold": {
            "label": "Kwota tuż pod progiem",
            "severity_label": "Średnie ryzyko",
            "impact": "Transakcja znajduje się bardzo blisko typowego progu akceptacji, co może oznaczać obchodzenie limitu.",
            "recommendation": "Sprawdź politykę akceptacji i porównaj z dokumentem źródłowym oraz osobą zatwierdzającą.",
        },
        "round_amount": {
            "label": "Kwota zaokrąglona",
            "severity_label": "Niska uwaga",
            "impact": "Idealnie okrągłe kwoty bywają uzasadnione, ale często wymagają dodatkowego kontekstu biznesowego.",
            "recommendation": "Zweryfikuj, czy zaokrąglenie wynika z umowy, budżetu lub ręcznej korekty wartości.",
        },
        "amount_outlier": {
            "label": "Odstająca kwota",
            "severity_label": "Wysokie ryzyko",
            "impact": "Kwota wyraźnie odstaje od reszty populacji i może wymagać osobnego uzasadnienia biznesowego.",
            "recommendation": "Porównaj transakcję z podobnymi pozycjami i sprawdź dokument źródłowy oraz ścieżkę akceptacji.",
            "fraud_hypothesis": "Kwota skrajnie odbiega od populacji i może być podstawiona poza normalnym wzorcem transakcji.",
            "intent_hypothesis": "Wprowadzenie jednorazowej, nieuprawnionej operacji na dużą kwotę.",
            "verification_target": "Zweryfikuj uzasadnienie biznesowe i dokument źródłowy dla tej konkretnej transakcji.",
            "confidence_basis": "high_statistical_outlier",
        },
        "benford_deviation": {
            "label": "Odchylenie od Benforda",
            "severity_label": "Wysokie ryzyko",
            "impact": "Rozkład pierwszych cyfr odbiega od wzorca Benforda, co może wskazywać na nienaturalną strukturę danych liczbowych.",
            "recommendation": "Traktuj to jako sygnał statystyczny — zestaw wynik z dokumentami źródłowymi i innymi czerwonymi flagami.",
            "fraud_hypothesis": "Rozkład cyfr wygląda nienaturalnie, co może sugerować ręczne konstruowanie wartości.",
            "intent_hypothesis": "Możliwe dopisywanie fikcyjnych wartości tworzonych z pominięciem naturalnego wzorca danych.",
            "verification_target": "Przeprowadź pełny audyt dokumentacji źródłowej dla tej kolumny i zestaw wynik z innymi czerwonymi flagami.",
            "confidence_basis": "high_statistical_distribution_shift",
        },
        "weekend_activity": {
            "label": "Aktywność weekendowa",
            "severity_label": "Niska uwaga",
            "impact": "Rekord został oznaczony datą weekendową, co bywa nietypowe w niektórych procesach księgowych.",
            "recommendation": "Sprawdź, czy transakcja mogła zostać zarejestrowana automatycznie lub po godzinach pracy.",
        },
        "night_activity": {
            "label": "Aktywność nocna",
            "severity_label": "Niska uwaga",
            "impact": "Rekord ma znacznik czasu w godzinach nocnych, co może wymagać wyjaśnienia operacyjnego.",
            "recommendation": "Zweryfikuj źródło zapisu i użytkownika lub system, który nadał godzinę operacji.",
        },
        "control_total_mismatch": {
            "label": "Niezgodność sumy kontrolnej",
            "severity_label": "Wysokie ryzyko",
            "impact": "Wynik formuły sumującej nie zgadza się z wyliczeniem opartym na wskazanym zakresie danych.",
            "recommendation": "Porównaj zakres sumowania, wartości źródłowe i ewentualne ręczne nadpisania wyniku.",
            "fraud_hypothesis": "Suma kontrolna nie zgadza się z zakresem wejściowym, więc wynik mógł zostać zniekształcony.",
            "intent_hypothesis": "Ukrycie manka albo manipulacja wynikiem agregacji.",
            "verification_target": "Sprawdź, które dokładnie komórki składowe zostały pominięte lub zmienione względem sumy kontrolnej.",
            "confidence_basis": "high_control_total_break",
        },
        "cross_sheet_hidden_reference": {
            "label": "Odwołanie do ukrytego arkusza",
            "severity_label": "Wysokie ryzyko",
            "impact": "Wynik zależy od danych z arkusza ukrytego lub very-hidden, co utrudnia audyt ścieżki obliczeń.",
            "recommendation": "Otwórz wskazany arkusz źródłowy i sprawdź, czy ukrycie jest uzasadnione oraz czy dane są spójne.",
            "fraud_hypothesis": "Wynik zależy od ukrytego źródła, co może maskować rzeczywisty mechanizm obliczeń.",
            "intent_hypothesis": "Ukrycie źródła wejścia albo obejście jawnej kontroli wyniku.",
            "verification_target": "Sprawdź zawartość ukrytego arkusza i wszystkie komórki zasilające to odwołanie.",
            "confidence_basis": "high_hidden_cross_sheet_dependency",
        },
    }
    meta = mapping.get(
        anomaly_type,
        {
            "label": anomaly_type or "Anomalia",
            "severity_label": "Do weryfikacji",
            "impact": "Wykryto nietypowy wzorzec wymagający ręcznej oceny.",
            "recommendation": "Przejrzyj wskazaną komórkę i zweryfikuj, czy logika jest zgodna z założeniami modelu.",
            "fraud_hypothesis": "Nietypowe zjawisko wymaga weryfikacji i może wynikać z ręcznej ingerencji albo błędu.",
            "intent_hypothesis": "Brak wystarczających dowodów intencji",
            "verification_target": "Zweryfikuj wskazaną komórkę, jej zależności i źródło danych wejściowych.",
            "confidence_basis": "manual_review_required",
        },
    )
    return {**anomaly, **meta}


def _severity_rank(severity: str) -> int:
    return {
        "CRITICAL": 4,
        "HIGH": 3,
        "MEDIUM": 2,
        "LOW": 1,
    }.get((severity or "").upper(), 0)


def _is_high_priority_anomaly(anomaly: dict) -> bool:
    return str(anomaly.get("severity", "")).upper() in {"HIGH", "CRITICAL"}


def _compact_text(value, limit: int = 220) -> str:
    text = _value_to_text(value)
    if len(text) <= limit:
        return text
    return text[: limit - 1] + "…"


def _build_ai_finding_id(sheet_name: str, anomaly: dict) -> str:
    cell = _format_evidence_cell(sheet_name, str(anomaly.get("cell", "") or "").strip()) or "sheet"
    return f"{sheet_name}:{cell}:{anomaly.get('type', 'anomaly')}"


def _build_ai_evidence_lines(sheet_name: str, anomaly: dict) -> list[str]:
    lines: list[str] = []
    cell = _format_evidence_cell(sheet_name, anomaly.get("cell", ""))
    if cell:
        lines.append(f"Komórka: {cell}")
    if anomaly.get("header"):
        lines.append(f"Kolumna: {anomaly['header']}")
    if anomaly.get("value") not in (None, ""):
        lines.append(f"Wartość: {_compact_text(_serialize_excel_value(anomaly.get('value')))}")
    if anomaly.get("formula"):
        lines.append(f"Formuła: {_compact_text(anomaly.get('formula'), limit=180)}")
    if anomaly.get("message"):
        lines.append(f"Fakt techniczny: {_compact_text(anomaly.get('message'))}")
    if anomaly.get("impact"):
        lines.append(f"Wpływ: {_compact_text(anomaly.get('impact'))}")
    rows = anomaly.get("rows", []) or []
    if rows:
        lines.append(f"Wiersze: {', '.join(map(str, rows[:12]))}")
    return lines[:6]


def _build_ai_evidence_item(sheet_name: str, anomaly: dict) -> dict:
    fact = {
        "sheet": sheet_name,
        "cell": _format_evidence_cell(sheet_name, anomaly.get("cell", "")) or None,
        "value": _serialize_excel_value(anomaly.get("value")),
        "formula": anomaly.get("formula") or None,
        "header": anomaly.get("header") or None,
        "rows": (anomaly.get("rows", []) or [])[:25] or None,
        "message": anomaly.get("message") or "",
    }
    return {
        "finding_id": _build_ai_finding_id(sheet_name, anomaly),
        "type": anomaly.get("type", ""),
        "label": anomaly.get("label", anomaly.get("type", "Anomalia")),
        "severity": anomaly.get("severity", ""),
        "fact": fact,
        "evidence": _build_ai_evidence_lines(sheet_name, anomaly),
        "fraud_hypothesis": anomaly.get("fraud_hypothesis", ""),
        "intent_hypothesis": anomaly.get("intent_hypothesis", "Brak wystarczających dowodów intencji"),
        "verification_target": anomaly.get("verification_target", ""),
        "confidence_basis": anomaly.get("confidence_basis", "manual_review_required"),
    }


def _build_sheet_ai_evidence_pack(
    sheet_name: str,
    sheet_risk_level: str,
    logic_description: str,
    logic_summary: dict,
    analysis: dict,
    anomalies: list[dict],
) -> dict | None:
    findings = [
        _build_ai_evidence_item(sheet_name, anomaly)
        for anomaly in anomalies
        if _is_high_priority_anomaly(anomaly)
    ]
    if not findings:
        return None

    return {
        "sheet": sheet_name,
        "priority": sheet_risk_level,
        "investigation_goal": "Oceń, czy zestaw sygnałów wskazuje na celową manipulację logiką lub danymi, bez przesądzania oszustwa.",
        "logic_context": {
            "description": logic_description,
            "output_cells": (logic_summary.get("output_cells", []) or [])[:8],
            "hub_cells": [item.get("cell") for item in (logic_summary.get("hub_cells", []) or [])[:5]],
            "focus_rows": (analysis.get("focus_rows", []) or [])[:12],
            "focus_columns": (analysis.get("focus_columns", []) or [])[:12],
        },
        "findings": findings[:_FINANCIAL_AI_OPINION_MAX_FINDINGS],
    }


def _build_ai_evidence_pack(file_name: str, sheet_packs: list[dict]) -> dict | None:
    valid_sheets = [sheet for sheet in sheet_packs if sheet and sheet.get("findings")]
    if not valid_sheets:
        return None
    findings_count = sum(len(sheet.get("findings", [])) for sheet in valid_sheets)
    return {
        "version": 1,
        "source": "financial_audit",
        "file_name": file_name,
        "investigation_goal": "Oceń, czy zestaw sygnałów wskazuje na poszlaki celowej manipulacji logiką lub danymi, bez przesądzania winy.",
        "summary": {
            "sheet_count": len(valid_sheets),
            "findings_count": findings_count,
            "allowed_severities": ["HIGH", "CRITICAL"],
        },
        "sheets": valid_sheets,
    }


def _build_financial_forensics_prompt(evidence_pack: dict) -> tuple[str, str]:
    mode = SEARCH_MODES["financial_forensics"]
    prompt = (
        "PAKIET DOWODOWY Z AUDYTU XLSX (jedyne źródło faktów):\n"
        f"{json.dumps(evidence_pack, ensure_ascii=False, indent=2)}\n\n"
        f"{mode['prompt_suffix']}"
    )
    return mode["system"], prompt


def _extract_json_object(raw_text: str) -> dict | None:
    text = (raw_text or "").strip()
    if not text:
        return None
    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        return None
    try:
        parsed = json.loads(match.group(0))
    except json.JSONDecodeError:
        return None
    return parsed if isinstance(parsed, dict) else None


def _normalize_ai_forensic_opinion(raw_opinion: dict | None) -> dict | None:
    if not isinstance(raw_opinion, dict):
        return None

    findings_payload = raw_opinion.get("findings")
    if not isinstance(findings_payload, list):
        findings_payload = []

    findings: list[dict] = []
    for item in findings_payload[:_FINANCIAL_AI_OPINION_MAX_FINDINGS]:
        if not isinstance(item, dict):
            continue
        fact = _compact_text(item.get("fact", ""))
        expert_comment = _compact_text(item.get("expert_comment", ""))
        if not fact and not expert_comment:
            continue
        findings.append(
            {
                "finding_id": _compact_text(item.get("finding_id", ""), 120),
                "fact": fact,
                "intent": _compact_text(
                    item.get("intent", "Brak wystarczających dowodów intencji")
                ),
                "expert_comment": expert_comment,
                "confidence": _compact_text(item.get("confidence", "")),
                "next_check": _compact_text(item.get("next_check", "")),
            }
        )

    overall = _compact_text(raw_opinion.get("overall_assessment", ""), 400)
    sheet_comment = _compact_text(raw_opinion.get("sheet_comment", ""), 400)
    limitations = _compact_text(raw_opinion.get("limitations", ""), 400)
    if not overall and not findings and not sheet_comment:
        return None
    return {
        "sheet_comment": sheet_comment,
        "overall_assessment": overall,
        "limitations": limitations,
        "findings": findings,
    }


def _risk_label_pl(level: str) -> str:
    return {
        "LOW": "niskie",
        "MEDIUM": "umiarkowane",
        "HIGH": "wysokie",
        "CRITICAL": "krytyczne",
    }.get((level or "").upper(), (level or "brak").lower())


def _report_export_filename(file_name: str, suffix: str = "docx") -> str:
    stem = Path(file_name or "audyt_xlsx").stem
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._") or "audyt_xlsx"
    return f"raport_audytu_finansowego_{safe}_{datetime.now().strftime('%Y%m%d_%H%M')}.{suffix}"


def _iter_high_priority_anomalies(report: dict) -> list[tuple[str, dict]]:
    items: list[tuple[str, dict]] = []
    for sheet_name, sheet_data in (report.get("sheets") or {}).items():
        if not isinstance(sheet_data, dict):
            continue
        for anomaly in sheet_data.get("anomalies", []) or []:
            if not isinstance(anomaly, dict):
                continue
            if (anomaly.get("severity") or "").upper() not in {"HIGH", "CRITICAL"}:
                continue
            items.append((sheet_name, anomaly))
    return items


def _build_financial_docx(report: dict) -> io.BytesIO:
    import docx as _docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = _docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Raport z audytu finansowego XLSX", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    file_meta = report.get("file_metadata") or {}
    file_name = file_meta.get("name") or report.get("file_name") or "brak"
    summary = report.get("summary") or {}
    risk_level = _risk_label_pl(summary.get("risk_level", ""))
    generated_at = datetime.now().strftime("%d.%m.%Y %H:%M")
    doc.add_paragraph(f"Data wygenerowania: {generated_at}")
    doc.add_paragraph(f"Nazwa pliku źródłowego: {file_name}")
    doc.add_paragraph(
        f"Priorytet ogólny: {risk_level} | Formuł: {summary.get('total_formulas', 0)} | "
        f"Silnych sygnałów: {sum(1 for _, anomaly in _iter_high_priority_anomalies(report) if anomaly)}"
    )

    doc.add_heading("Podsumowanie ryzyka", level=1)
    logic_overview = report.get("logic_overview") or {}
    doc.add_paragraph(
        logic_overview.get("description")
        or "Raport prezentuje tylko silne poszlaki wymagające dalszej weryfikacji."
    )

    high_priority = _iter_high_priority_anomalies(report)
    doc.add_heading("Findings HIGH/CRITICAL", level=1)
    if high_priority:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Arkusz"
        hdr[1].text = "Komórka"
        hdr[2].text = "Typ"
        hdr[3].text = "Severity"
        hdr[4].text = "Opis"
        for sheet_name, anomaly in high_priority[:20]:
            row = table.add_row().cells
            row[0].text = str(sheet_name)
            row[1].text = str(anomaly.get("cell") or "—")
            row[2].text = str(anomaly.get("label") or anomaly.get("type") or "anomaly")
            row[3].text = str(anomaly.get("severity") or "—")
            row[4].text = str(anomaly.get("message") or anomaly.get("impact") or "—")
    else:
        doc.add_paragraph("Brak findings HIGH/CRITICAL w przekazanym raporcie.")

    ai_opinion = report.get("ai_forensic_opinion")
    if isinstance(ai_opinion, dict):
        doc.add_heading("Wniosek śledczy AI", level=1)
        if ai_opinion.get("overall_assessment"):
            doc.add_paragraph(str(ai_opinion["overall_assessment"]))
        if ai_opinion.get("sheet_comment"):
            doc.add_paragraph(str(ai_opinion["sheet_comment"]))
        for item in ai_opinion.get("findings", []) or []:
            fact = str(item.get("fact") or "").strip()
            if not fact:
                continue
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"Fakt: {fact}").bold = True
            if item.get("intent"):
                doc.add_paragraph(f"Możliwa intencja: {item['intent']}")
            if item.get("expert_comment"):
                doc.add_paragraph(f"Komentarz ekspercki: {item['expert_comment']}")
            if item.get("next_check"):
                doc.add_paragraph(f"Co sprawdzić dalej: {item['next_check']}")

    doc.add_heading("Graf powiązań", level=1)
    lineage_included = False
    for sheet_name, sheet_data in (report.get("sheets") or {}).items():
        if not isinstance(sheet_data, dict):
            continue
        lineage = sheet_data.get("lineage_graph") or {}
        edges = lineage.get("edges") or []
        nodes = lineage.get("nodes") or []
        if not edges:
            continue
        lineage_included = True
        doc.add_heading(str(sheet_name), level=2)
        summary_block = lineage.get("summary") or {}
        doc.add_paragraph(
            f"Relacje: {summary_block.get('relation_count', len(edges))} | "
            f"Węzły ukryte: {summary_block.get('hidden_node_count', 0)} | "
            f"Silne findings: {summary_block.get('high_priority_findings', 0)}"
        )
        edge_table = doc.add_table(rows=1, cols=4)
        edge_table.style = "Table Grid"
        hdr = edge_table.rows[0].cells
        hdr[0].text = "Źródło"
        hdr[1].text = "Cel"
        hdr[2].text = "Typ relacji"
        hdr[3].text = "Powód"
        for edge in edges[:12]:
            row = edge_table.add_row().cells
            row[0].text = str(edge.get("source") or "—")
            row[1].text = str(edge.get("target") or "—")
            row[2].text = str(edge.get("type") or "dependency")
            row[3].text = str(edge.get("reason") or "—")
        if nodes:
            doc.add_paragraph("Węzły objęte grafem:")
            node_table = doc.add_table(rows=1, cols=5)
            node_table.style = "Table Grid"
            hdr = node_table.rows[0].cells
            hdr[0].text = "Węzeł"
            hdr[1].text = "Typ"
            hdr[2].text = "Ukryty"
            hdr[3].text = "Severity"
            hdr[4].text = "Wartość"
            for node in nodes[:12]:
                row = node_table.add_row().cells
                row[0].text = str(node.get("label") or node.get("id") or "—")
                row[1].text = str(node.get("type") or "node")
                row[2].text = "tak" if node.get("hidden") else "nie"
                row[3].text = str(node.get("severity") or "NONE")
                row[4].text = str(node.get("value") if node.get("value") is not None else "—")
    if not lineage_included:
        doc.add_paragraph("Brak relacji do pokazania w grafie powiązań.")

    doc.add_heading("Co sprawdzić dalej", level=1)
    next_checks: list[str] = []
    if isinstance(ai_opinion, dict):
        for item in ai_opinion.get("findings", []) or []:
            next_check = str(item.get("next_check") or "").strip()
            if next_check and next_check not in next_checks:
                next_checks.append(next_check)
    for _, anomaly in high_priority:
        recommendation = str(
            anomaly.get("verification_target")
            or anomaly.get("recommendation")
            or ""
        ).strip()
        if recommendation and recommendation not in next_checks:
            next_checks.append(recommendation)
        if len(next_checks) >= 8:
            break
    if next_checks:
        for item in next_checks[:8]:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("Brak dodatkowych kroków weryfikacyjnych w przekazanym payloadzie.")

    doc.add_heading("Zastrzeżenie", level=1)
    doc.add_paragraph(
        "Raport wskazuje poszlaki i nietypowe wzorce wymagające weryfikacji przez człowieka. "
        "Nie przesądza winy, oszustwa ani odpowiedzialności prawnej."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _format_evidence_cell(sheet_name: str, cell: str) -> str:
    cell_text = (cell or "").strip()
    if not cell_text:
        return ""
    if "!" in cell_text or "," in cell_text:
        return cell_text
    return f"{sheet_name}!{cell_text}"


def _short_formula(formula: str, limit: int = 120) -> str:
    text = (formula or "").strip()
    if len(text) <= limit:
        return text
    return text[: limit - 1] + "…"


def _build_sheet_evidence_summary(sheet_name: str, anomalies: list[dict], logic_summary: dict) -> dict:
    ranked = sorted(
        anomalies,
        key=lambda item: (
            -_severity_rank(item.get("severity", "")),
            str(item.get("cell", "")),
            str(item.get("message", "")),
        ),
    )
    evidence_items: list[str] = []
    seen: set[tuple[str, str, str]] = set()

    for anomaly in ranked:
        anomaly_type = anomaly.get("type", "")
        cell = _format_evidence_cell(sheet_name, anomaly.get("cell", ""))
        formula = _short_formula(anomaly.get("formula", ""))
        header = anomaly.get("header", "")
        value = _serialize_excel_value(anomaly.get("value"))
        rows = anomaly.get("rows", []) or []
        columns = anomaly.get("columns", []) or []
        missing_numbers = anomaly.get("missing_numbers", []) or []
        threshold = anomaly.get("threshold")

        if anomaly_type == "cosmetic_check":
            detail = anomaly.get("message", "")
            if ": " in detail:
                detail = detail.split(": ", 1)[1]
            text = f"{cell}: kontrola {formula} jest pozorna — {detail}."
        elif anomaly_type == "hardcoded_in_formula_region":
            text = (
                f"{cell}: w kolumnie „{header}” wpisano ręcznie wartość {value} w miejscu, "
                "gdzie sąsiednie pozycje liczą się formułami."
            )
        elif anomaly_type == "hardcoded_output_like_value":
            text = (
                f"{cell}: wynik w kolumnie „{header}” został wpisany ręcznie, "
                "mimo że wygląda jak końcowy rezultat obliczeń."
            )
        elif anomaly_type == "formula_pattern_deviation":
            text = (
                f"{cell}: formuła {formula} odbiega od dominującego wzorca w kolumnie „{header}”, "
                "więc ten wiersz może liczyć coś innego niż reszta."
            )
        elif anomaly_type == "duplicate_document":
            text = (
                f"{cell or sheet_name}: numer dokumentu „{value}” powtarza się w wierszach "
                f"{', '.join(map(str, rows[:6]))}."
            )
        elif anomaly_type == "duplicate_amount":
            text = (
                f"{cell or sheet_name}: kwota {value} powtarza się w kolumnie „{header}” "
                f"w wierszach {', '.join(map(str, rows[:6]))}."
            )
        elif anomaly_type == "numbering_gap":
            text = (
                f"{sheet_name}: w kolumnie „{header}” brakuje numerów "
                f"{', '.join(map(str, missing_numbers[:8]))}."
            )
        elif anomaly_type == "near_threshold":
            text = f"{cell}: kwota {value} jest tuż poniżej progu {threshold}."
        elif anomaly_type == "amount_outlier":
            z_score = anomaly.get("z_score")
            z_text = f" (z-score {z_score})" if z_score is not None else ""
            text = f"{cell}: kwota {value} odstaje od reszty populacji{z_text}."
        elif anomaly_type == "benford_deviation":
            sample_size = anomaly.get("sample_size")
            text = (
                f"{sheet_name}: kolumna „{header}” ma nienaturalny rozkład cyfr wiodących"
                + (f" dla próby {sample_size} wartości." if sample_size else ".")
            )
        elif anomaly_type == "hidden_rows":
            text = f"{sheet_name}: ukryto wiersze {', '.join(map(str, rows[:10]))}."
        elif anomaly_type == "hidden_columns":
            text = f"{sheet_name}: ukryto kolumny {', '.join(columns[:10])}."
        elif anomaly_type == "control_total_mismatch":
            text = f"{cell}: suma kontrolna nie zgadza się z zakresem wejściowym."
        elif anomaly_type == "cross_sheet_hidden_reference":
            text = f"{cell}: wynik odwołuje się do ukrytego arkusza, więc ścieżka obliczeń nie jest jawna."
        else:
            message = anomaly.get("message", "")
            if cell:
                text = f"{cell}: {message}"
            else:
                text = f"{sheet_name}: {message}"

        fingerprint = (anomaly_type, cell, text)
        if fingerprint in seen:
            continue
        seen.add(fingerprint)
        evidence_items.append(text)
        if len(evidence_items) >= 4:
            break

    if evidence_items:
        return {
            "headline": "Najmocniejsze sygnały do sprawdzenia:",
            "items": evidence_items,
        }

    return {
        "headline": "Opis logiki arkusza:",
        "items": [logic_summary.get("description_pl") or "Nie wykryto konkretnych sygnałów wymagających wyjaśnienia."],
    }


def _extract_anomaly_formula_keys(sheet_name: str, anomalies: list[dict], formula_keys: set[str]) -> list[str]:
    keys: list[str] = []
    for anomaly in anomalies:
        raw_cell = str(anomaly.get("cell", "") or "")
        if not raw_cell:
            continue
        for token in raw_cell.split(","):
            coord = token.strip()
            if not coord:
                continue
            full_key = coord if "!" in coord else f"{sheet_name}!{coord}"
            if full_key in formula_keys and full_key not in keys:
                keys.append(full_key)
    return keys


def _build_sheet_flow_graph(sheet_name: str, analyzer: AuditAnalyzer, logic_summary: dict, anomalies: list[dict]) -> dict:
    formula_keys = set(analyzer.cells.keys())
    if not formula_keys:
        return {
            "nodes": [],
            "edges": [],
            "narratives": [],
            "hidden_formula_nodes": 0,
            "high_risk_relation_count": 0,
        }

    high_risk_anomalies = [
        anomaly for anomaly in anomalies if (anomaly.get("severity") or "").upper() == "HIGH"
    ]
    if not high_risk_anomalies:
        return {
            "nodes": [],
            "edges": [],
            "narratives": [],
            "hidden_formula_nodes": len(formula_keys),
            "high_risk_relation_count": 0,
        }

    output_keys = [
        key for key in logic_summary.get("output_cells", []) or [] if key in formula_keys
    ]
    hub_keys = [
        item.get("cell")
        for item in logic_summary.get("hub_cells", []) or []
        if item.get("cell") in formula_keys
    ]
    anomaly_keys = _extract_anomaly_formula_keys(sheet_name, high_risk_anomalies, formula_keys)

    anchor_keys: list[str] = []
    for group in (anomaly_keys, output_keys, hub_keys):
        for key in group:
            if key and key not in anchor_keys:
                anchor_keys.append(key)
            if len(anchor_keys) >= 10:
                break
        if len(anchor_keys) >= 10:
            break

    if not anchor_keys:
        return {
            "nodes": [],
            "edges": [],
            "narratives": [],
            "hidden_formula_nodes": len(formula_keys),
            "high_risk_relation_count": len(high_risk_anomalies),
        }

    selected_formula_keys: set[str] = set()
    queue = deque((key, 0) for key in anchor_keys)
    seen_depth: dict[str, int] = {}
    max_formula_nodes = 28
    max_depth = 2

    while queue and len(selected_formula_keys) < max_formula_nodes:
        key, depth = queue.popleft()
        if key in seen_depth and seen_depth[key] <= depth:
            continue
        seen_depth[key] = depth
        if key not in formula_keys:
            continue
        selected_formula_keys.add(key)
        if depth >= max_depth:
            continue
        cell = analyzer.cells[key]
        related_formula_keys = [
            str(dep) for dep in cell.dependencies if str(dep) in formula_keys
        ] + [
            str(dep) for dep in cell.dependents if str(dep) in formula_keys
        ]
        for related_key in related_formula_keys:
            if related_key not in seen_depth or seen_depth[related_key] > depth + 1:
                queue.append((related_key, depth + 1))

    if not selected_formula_keys:
        selected_formula_keys = set(sorted(formula_keys)[: min(12, len(formula_keys))])

    source_nodes: dict[str, dict] = {}
    edges: list[dict] = []
    reverse_edges: dict[str, list[str]] = defaultdict(list)

    for cell_key in selected_formula_keys:
        cell = analyzer.cells[cell_key]
        for dep in cell.dependencies:
            dep_key = str(dep)
            if dep_key in selected_formula_keys:
                edges.append({"source": dep_key, "target": cell_key})
                reverse_edges[cell_key].append(dep_key)
                continue
            source_nodes.setdefault(
                dep_key,
                {
                    "id": dep_key,
                    "sheet": dep.sheet or sheet_name,
                    "label": dep_key.split("!", 1)[-1] if dep.sheet == sheet_name else dep_key,
                    "kind": "input" if dep.sheet == sheet_name else "external",
                    "note": "Wejście z tego samego arkusza" if dep.sheet == sheet_name else "Wejście z innego arkusza",
                },
            )
            edges.append({"source": dep_key, "target": cell_key})
            reverse_edges[cell_key].append(dep_key)

    if not output_keys:
        output_keys = sorted(
            selected_formula_keys,
            key=lambda key: len(analyzer.cells[key].dependents),
            reverse=True,
        )[:3]

    distance_to_output: dict[str, int] = {}
    queue = deque()
    for key in output_keys:
        if key in selected_formula_keys:
            distance_to_output[key] = 0
            queue.append(key)

    while queue:
        target_key = queue.popleft()
        current_distance = distance_to_output[target_key]
        for source_key in reverse_edges.get(target_key, []):
            next_distance = current_distance + 1
            if source_key not in distance_to_output or next_distance < distance_to_output[source_key]:
                distance_to_output[source_key] = next_distance
                queue.append(source_key)

    if not distance_to_output:
        for index, key in enumerate(sorted(selected_formula_keys)[:8]):
            distance_to_output[key] = index % 3

    max_distance = max(distance_to_output.values(), default=0)
    selected_anomaly_cells = set(anomaly_keys)
    selected_hubs = set(hub_keys)
    selected_outputs = set(output_keys)

    nodes: list[dict] = []
    for node_key, source in source_nodes.items():
        distance = distance_to_output.get(node_key, max_distance + 1)
        nodes.append(
            {
                "id": node_key,
                "label": source["label"],
                "kind": source["kind"],
                "layer": max(max_distance + 1 - distance, 0),
                "note": source["note"],
                "formula": "",
                "dependencies_count": 0,
                "dependents_count": 1,
            }
        )

    for node_key in selected_formula_keys:
        cell = analyzer.cells[node_key]
        if node_key in selected_outputs:
            kind = "output"
            note = "Komórka wynikowa"
        elif node_key in selected_anomaly_cells:
            kind = "anomaly"
            note = "Wymaga sprawdzenia"
        elif node_key in selected_hubs:
            kind = "hub"
            note = "Centralny węzeł obliczeń"
        else:
            kind = "formula"
            note = "Etap obliczeń"
        distance = distance_to_output.get(node_key, max_distance)
        nodes.append(
            {
                "id": node_key,
                "label": node_key.split("!", 1)[-1],
                "kind": kind,
                "layer": max(max_distance + 1 - distance, 0),
                "note": note,
                "formula": _short_formula(cell.formula, limit=90),
                "dependencies_count": len(cell.dependencies),
                "dependents_count": len(cell.dependents),
            }
        )

    nodes.sort(key=lambda item: (item["layer"], item["kind"], item["label"]))
    node_ids = {node["id"] for node in nodes}
    edges = [edge for edge in edges if edge["source"] in node_ids and edge["target"] in node_ids]

    narratives: list[str] = []
    for output_key in output_keys[:3]:
        if output_key not in selected_formula_keys:
            continue
        deps = [edge["source"] for edge in edges if edge["target"] == output_key]
        if not deps:
            continue
        dep_labels = [dep.split("!", 1)[-1] if dep.startswith(f"{sheet_name}!") else dep for dep in deps[:4]]
        extra = ""
        if len(deps) > 4:
            extra = f" i {len(deps) - 4} kolejnych wejść"
        narratives.append(
            f"{output_key.split('!', 1)[-1]} powstaje z {', '.join(dep_labels)}{extra}."
        )

    for anomaly in high_risk_anomalies:
        cell_key = _format_evidence_cell(sheet_name, anomaly.get("cell", ""))
        if cell_key in selected_anomaly_cells:
            narratives.append(anomaly.get("message", ""))
        if len(narratives) >= 6:
            break

    unique_narratives: list[str] = []
    for text in narratives:
        clean = (text or "").strip()
        if clean and clean not in unique_narratives:
            unique_narratives.append(clean)

    return {
        "nodes": nodes,
        "edges": edges,
        "narratives": unique_narratives[:6],
        "hidden_formula_nodes": max(0, len(formula_keys) - len(selected_formula_keys)),
        "high_risk_relation_count": len(high_risk_anomalies),
    }


def _empty_lineage_graph() -> dict:
    return {
        "nodes": [],
        "edges": [],
        "summary": {
            "relation_count": 0,
            "hidden_node_count": 0,
            "high_priority_findings": 0,
        },
    }


def _graph_cell_hidden(workbook, sheet_name: str, coord: str) -> bool:
    if not workbook or sheet_name not in workbook.sheetnames or not coord:
        return False
    sheet = workbook[sheet_name]
    row_idx, col_idx = _parse_excel_coord(coord)
    if row_idx is None or col_idx is None:
        return False
    from openpyxl.utils import get_column_letter

    col_letter = get_column_letter(col_idx)
    return bool(
        sheet.sheet_state != "visible"
        or sheet.row_dimensions[row_idx].hidden
        or sheet.column_dimensions[col_letter].hidden
    )


def _graph_cell_formula(workbook, sheet_name: str, coord: str) -> str:
    if not workbook or sheet_name not in workbook.sheetnames or not coord:
        return ""
    row_idx, col_idx = _parse_excel_coord(coord)
    if row_idx is None or col_idx is None:
        return ""
    value = workbook[sheet_name].cell(row_idx, col_idx).value
    return value if isinstance(value, str) and value.startswith("=") else ""


def _graph_cell_value(values_wb, sheet_name: str, coord: str):
    if not values_wb or sheet_name not in values_wb.sheetnames or not coord:
        return None
    row_idx, col_idx = _parse_excel_coord(coord)
    if row_idx is None or col_idx is None:
        return None
    return values_wb[sheet_name].cell(row_idx, col_idx).value


def _append_lineage_node(nodes_by_id: dict[str, dict], **node_fields):
    node_id = node_fields["id"]
    existing = nodes_by_id.get(node_id)
    if existing is None:
        nodes_by_id[node_id] = node_fields
        return
    for key, value in node_fields.items():
        if key == "severity":
            current_rank = _severity_rank(existing.get("severity", ""))
            new_rank = _severity_rank(value or "")
            if new_rank > current_rank:
                existing[key] = value
            continue
        if existing.get(key) in (None, "", False) and value not in (None, ""):
            existing[key] = value
    existing["hidden"] = bool(existing.get("hidden") or node_fields.get("hidden"))


def _append_lineage_edge(edges: list[dict], seen: set[tuple[str, str, str]], **edge_fields):
    fingerprint = (
        edge_fields["source"],
        edge_fields["target"],
        edge_fields.get("type", ""),
    )
    if fingerprint in seen:
        return
    seen.add(fingerprint)
    edges.append(edge_fields)


def _build_sheet_lineage_graph(
    sheet_name: str,
    analyzer: AuditAnalyzer,
    anomalies: list[dict],
    workbook,
    values_wb=None,
) -> dict:
    try:
        high_priority_anomalies = [
            anomaly
            for anomaly in anomalies
            if (anomaly.get("severity") or "").upper() in {"HIGH", "CRITICAL"}
        ]
        if not high_priority_anomalies:
            return _empty_lineage_graph()

        nodes_by_id: dict[str, dict] = {}
        edges: list[dict] = []
        seen_edges: set[tuple[str, str, str]] = set()

        def add_node(
            *,
            node_id: str,
            node_sheet: str,
            node_cell: str,
            label: str,
            node_type: str,
            severity: str = "",
            hidden: bool = False,
            value=None,
            formula: str = "",
        ):
            _append_lineage_node(
                nodes_by_id,
                id=node_id,
                sheet=node_sheet,
                cell=node_cell,
                label=label,
                type=node_type,
                severity=severity or "NONE",
                hidden=bool(hidden),
                value=_serialize_excel_value(value),
                formula=formula or "",
            )

        for anomaly in high_priority_anomalies:
            raw_cell = str(anomaly.get("cell", "") or "").strip()
            if not raw_cell:
                continue
            target_coord = raw_cell.split(",")[0].strip().split("!")[-1]
            target_id = f"{sheet_name}!{target_coord}"
            target_formula = anomaly.get("formula") or _graph_cell_formula(
                workbook, sheet_name, target_coord
            )
            target_hidden = _graph_cell_hidden(workbook, sheet_name, target_coord)
            target_value = anomaly.get("value")
            if target_value is None:
                target_value = _graph_cell_value(values_wb, sheet_name, target_coord)
            target_type = "formula"
            if target_id in analyzer.cells and not analyzer.cells[target_id].dependents:
                target_type = "output"
            elif not target_formula:
                target_type = "flagged_cell"

            add_node(
                node_id=target_id,
                node_sheet=sheet_name,
                node_cell=target_coord,
                label=target_id,
                node_type=target_type,
                severity=anomaly.get("severity", ""),
                hidden=target_hidden,
                value=target_value,
                formula=target_formula,
            )

            if target_id in analyzer.cells:
                cell = analyzer.cells[target_id]
                for dep in list(cell.dependencies)[:6]:
                    dep_id = str(dep)
                    dep_coord = f"{dep.col}{dep.row}"
                    dep_hidden = _graph_cell_hidden(workbook, dep.sheet or sheet_name, dep_coord)
                    dep_formula = (
                        analyzer.cells[dep_id].formula if dep_id in analyzer.cells else _graph_cell_formula(workbook, dep.sheet or sheet_name, dep_coord)
                    )
                    dep_value = _graph_cell_value(values_wb, dep.sheet or sheet_name, dep_coord)
                    dep_type = "input"
                    edge_type = "formula_dependency"
                    edge_reason = f"Formuła w {target_id} pobiera dane z {dep_id}."
                    if dep.sheet and dep.sheet != sheet_name:
                        dep_type = "hidden_source" if dep_hidden else "external_source"
                        edge_type = "cross_sheet_reference"
                        edge_reason = (
                            f"Formuła w {target_id} odwołuje się do {'ukrytego ' if dep_hidden else ''}arkusza {dep.sheet}."
                        )
                    add_node(
                        node_id=dep_id,
                        node_sheet=dep.sheet or sheet_name,
                        node_cell=dep_coord,
                        label=dep_id,
                        node_type=dep_type,
                        severity=anomaly.get("severity", "") if dep.sheet and dep.sheet != sheet_name else "",
                        hidden=dep_hidden,
                        value=dep_value,
                        formula=dep_formula,
                    )
                    _append_lineage_edge(
                        edges,
                        seen_edges,
                        source=dep_id,
                        target=target_id,
                        type=edge_type,
                        reason=edge_reason,
                    )

                for dependent in list(cell.dependents)[:4]:
                    dependent_id = str(dependent)
                    dependent_coord = f"{dependent.col}{dependent.row}"
                    dependent_formula = (
                        analyzer.cells[dependent_id].formula if dependent_id in analyzer.cells else ""
                    )
                    dependent_type = "output"
                    if dependent_id in analyzer.cells and analyzer.cells[dependent_id].dependents:
                        dependent_type = "formula"
                    add_node(
                        node_id=dependent_id,
                        node_sheet=sheet_name,
                        node_cell=dependent_coord,
                        label=dependent_id,
                        node_type=dependent_type,
                        hidden=_graph_cell_hidden(workbook, sheet_name, dependent_coord),
                        value=_graph_cell_value(values_wb, sheet_name, dependent_coord),
                        formula=dependent_formula,
                    )
                    _append_lineage_edge(
                        edges,
                        seen_edges,
                        source=target_id,
                        target=dependent_id,
                        type="feeds_result",
                        reason=f"Wynik z {target_id} zasila dalsze obliczenie w {dependent_id}.",
                    )

            if anomaly.get("type") == "cross_sheet_hidden_reference":
                source_sheet = anomaly.get("source_sheet")
                source_cell = anomaly.get("source_cell")
                if source_sheet and source_cell:
                    source_id = f"{source_sheet}!{source_cell}"
                    add_node(
                        node_id=source_id,
                        node_sheet=source_sheet,
                        node_cell=source_cell,
                        label=source_id,
                        node_type="hidden_source",
                        severity=anomaly.get("severity", ""),
                        hidden=_graph_cell_hidden(workbook, source_sheet, source_cell),
                        value=_graph_cell_value(values_wb, source_sheet, source_cell),
                        formula=_graph_cell_formula(workbook, source_sheet, source_cell),
                    )
                    _append_lineage_edge(
                        edges,
                        seen_edges,
                        source=source_id,
                        target=target_id,
                        type="cross_sheet_reference",
                        reason=(
                            f"Ukryta komórka {source_id} zasila wynik {target_id}."
                        ),
                    )

        if not nodes_by_id and not edges:
            return _empty_lineage_graph()

        nodes = sorted(
            nodes_by_id.values(),
            key=lambda item: (
                0 if item["type"] == "hidden_source" else 1,
                -_severity_rank(item.get("severity", "")),
                item["label"],
            ),
        )
        hidden_node_count = sum(1 for node in nodes if node.get("hidden"))
        return {
            "nodes": nodes,
            "edges": edges,
            "summary": {
                "relation_count": len(edges),
                "hidden_node_count": hidden_node_count,
                "high_priority_findings": len(high_priority_anomalies),
            },
        }
    except Exception as exc:
        logger.warning("Lineage graph build failed for %s: %s", sheet_name, exc)
        return _empty_lineage_graph()


def _empty_forensic_result(formula_count: int, structure: dict) -> dict:
    return {
        "findings": [],
        "technical_appendix": [],
        "structure": structure,
        "formula_signals": {
            "hardcoded_values_in_formula_regions": [],
            "formula_pattern_deviations": [],
            "hardcoded_output_like_values": [],
        },
        "data_signals": {
            "duplicates": [],
            "amount_outliers": [],
            "benford_deviations": [],
        },
        "control_signals": {
            "control_total_mismatches": [],
            "cross_sheet_hidden_references": [],
        },
        "formula_cells_count": formula_count,
    }


def _analysis_mode_rank(mode: str) -> int:
    normalized = _normalize_analysis_type(mode)
    return {
        "quick": 1,
        "targeted": 2,
        "full": 3,
        "formulas": 3,
    }.get(normalized, 0)


def _file_sha256(file_path: Path) -> str:
    digest = hashlib.sha256()
    with file_path.open("rb") as handle:
        while True:
            chunk = handle.read(1024 * 1024)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def _remember_financial_cache_entry(fingerprint: str, entry: dict) -> None:
    _FINANCIAL_AUDIT_CACHE[fingerprint] = entry
    _FINANCIAL_AUDIT_CACHE.move_to_end(fingerprint)
    while len(_FINANCIAL_AUDIT_CACHE) > _FINANCIAL_AUDIT_CACHE_MAX:
        _FINANCIAL_AUDIT_CACHE.popitem(last=False)


def _build_base_sheet_cache(formulas_result: dict) -> dict:
    workbook = formulas_result["workbook"]
    base_sheets: dict[str, dict] = {}

    for sheet_name, sheet_info in formulas_result.get("sheets", {}).items():
        formulas = sheet_info.get("formulas", [])
        analyzer = AuditAnalyzer()
        ws = workbook[sheet_name]

        for formula_info in formulas:
            cell_ref = CellRef(
                sheet=sheet_name,
                col="".join([c for c in formula_info["cell"] if c.isalpha()]),
                row=int("".join([c for c in formula_info["cell"] if c.isdigit()])),
            )
            analyzer.add_cell(
                cell_ref,
                formula_info["formula"],
                value=formula_info.get("value"),
            )
            if formula_info.get("is_hidden"):
                analyzer.mark_hidden(cell_ref)

        analyzer.build_dependency_graph()
        analyzer.detect_anomalies()
        logic_summary = analyzer.get_logic_summary()
        screening_findings: list = []
        screening_appendix: list = []
        structure = _sheet_structure_signals(
            sheet_name, ws, screening_findings, screening_appendix, record_findings=True
        )
        screening_raw = analyzer.anomalies + screening_findings
        scope = _build_targeted_scope(ws, formulas, screening_raw)

        base_sheets[sheet_name] = {
            "sheet_info": sheet_info,
            "formulas": formulas,
            "analyzer": analyzer,
            "logic_summary": logic_summary,
            "screening_findings": screening_findings,
            "screening_appendix": screening_appendix,
            "screening_raw": screening_raw,
            "structure": structure,
            "scope": scope,
            "forensic_cache": {},
        }

    return base_sheets


def _get_financial_cache_entry(tmp_path: Path) -> tuple[str, dict, bool]:
    fingerprint = _file_sha256(tmp_path)
    cached = _FINANCIAL_AUDIT_CACHE.get(fingerprint)
    if cached:
        _FINANCIAL_AUDIT_CACHE.move_to_end(fingerprint)
        return fingerprint, cached, True

    formulas_result = _extract_formulas_from_excel(tmp_path)
    if "error" in formulas_result:
        return fingerprint, formulas_result, False

    entry = {
        "fingerprint": fingerprint,
        "formulas_result": formulas_result,
        "base_sheets": _build_base_sheet_cache(formulas_result),
        "values_wb": None,
    }
    _remember_financial_cache_entry(fingerprint, entry)
    return fingerprint, entry, False


def _get_cached_values_workbook(cache_entry: dict, tmp_path: Path):
    if cache_entry.get("values_wb") is None:
        cache_entry["values_wb"] = _load_values_workbook(tmp_path)
    return cache_entry["values_wb"]


def _get_sheet_forensic_result(cache_entry: dict, tmp_path: Path, sheet_name: str, analysis_type: str) -> tuple[dict, str, bool]:
    normalized = _normalize_analysis_type(analysis_type)
    base_sheet = cache_entry["base_sheets"][sheet_name]
    formulas = base_sheet["formulas"]
    structure = base_sheet["structure"]
    scope = base_sheet["scope"]
    forensic_cache = base_sheet["forensic_cache"]

    if normalized == "quick":
        return _empty_forensic_result(len(formulas), structure), "screening_only", False

    if normalized == "targeted" and not scope.get("should_deep_scan"):
        return _empty_forensic_result(len(formulas), structure), "screening_only", False

    cache_key = "full" if normalized in {"full", "formulas"} else "targeted"
    cached = forensic_cache.get(cache_key)
    if cached:
        return cached["forensic"], cached["sheet_mode"], True

    workbook = cache_entry["formulas_result"]["workbook"]
    ws = workbook[sheet_name]
    values_ws = _get_cached_values_workbook(cache_entry, tmp_path)[sheet_name]

    if cache_key == "full":
        forensic = _analyze_workbook_forensics(
            sheet_name,
            ws,
            values_ws,
            formulas,
            include_structure=False,
        )
        sheet_mode = "full_deep_scan"
    else:
        forensic = _analyze_workbook_forensics(
            sheet_name,
            ws,
            values_ws,
            formulas,
            include_structure=False,
            selected_rows=scope["selected_rows"],
            selected_cols=scope["selected_cols"],
        )
        sheet_mode = "targeted_deep_scan"

    forensic_cache[cache_key] = {
        "forensic": forensic,
        "sheet_mode": sheet_mode,
    }
    return forensic, sheet_mode, False


def _describe_sheet_logic(sheet_name: str, logic_summary: dict) -> str:
    outputs = logic_summary.get("output_like_cells", 0)
    inputs = logic_summary.get("input_like_cells", 0)
    depth = logic_summary.get("max_dependency_depth", 0)
    cross_refs = logic_summary.get("cross_sheet_references", 0)
    text = (
        f"Arkusz „{sheet_name}” zawiera {outputs} {_pl_count(outputs, 'komórkę wynikową', 'komórki wynikowe', 'komórek wynikowych')} i {inputs} {_pl_count(inputs, 'formułę bazową', 'formuły bazowe', 'formuł bazowych')}."
    )
    if depth >= 4:
        text += f" Logika jest wieloetapowa — najdłuższy łańcuch zależności ma {depth} {_pl_count(depth, 'poziom', 'poziomy', 'poziomów')}."
    elif depth >= 2:
        text += f" Obliczenia przebiegają maksymalnie przez {depth} {_pl_count(depth, 'poziom', 'poziomy', 'poziomów')} zależności."
    else:
        text += " Logika jest płytka i stosunkowo łatwa do prześledzenia."
    if cross_refs:
        text += f" Dodatkowo wykryto {cross_refs} {_pl_count(cross_refs, 'odwołanie', 'odwołania', 'odwołań')} między arkuszami."
    return text


def _sheet_structure_signals(
    sheet_name: str, ws, findings: list, appendix: list, *, record_findings: bool = True
) -> dict:
    hidden_rows = sorted(idx for idx, dim in ws.row_dimensions.items() if dim.hidden)
    hidden_cols = sorted(str(col) for col, dim in ws.column_dimensions.items() if dim.hidden)
    filters_active = bool(ws.auto_filter and ws.auto_filter.ref)
    protection_enabled = bool(getattr(ws.protection, "sheet", False))
    merged_ranges = [str(rng) for rng in ws.merged_cells.ranges]

    if hidden_rows and record_findings:
        _append_finding(
            findings,
            appendix,
            finding_type="hidden_rows",
            severity="MEDIUM",
            sheet=sheet_name,
            message=f"W arkuszu ukryto {len(hidden_rows)} {_pl_count(len(hidden_rows), 'wiersz', 'wiersze', 'wierszy')}.",
            comment=f"Ukryte wiersze: {', '.join(map(str, hidden_rows[:12]))}",
            details={"rows": hidden_rows[:25]},
        )
    if hidden_cols and record_findings:
        _append_finding(
            findings,
            appendix,
            finding_type="hidden_columns",
            severity="MEDIUM",
            sheet=sheet_name,
            message=f"W arkuszu ukryto {len(hidden_cols)} {_pl_count(len(hidden_cols), 'kolumnę', 'kolumny', 'kolumn')} danych.",
            comment=f"Ukryte kolumny: {', '.join(hidden_cols[:12])}",
            details={"columns": hidden_cols[:25]},
        )
    if merged_ranges and record_findings:
        _append_finding(
            findings,
            appendix,
            finding_type="merged_cells",
            severity="LOW",
            sheet=sheet_name,
            message=f"W arkuszu wykryto {len(merged_ranges)} {_pl_count(len(merged_ranges), 'scalony zakres', 'scalone zakresy', 'scalonych zakresów')}.",
            comment=f"Scalone zakresy: {', '.join(merged_ranges[:10])}",
            details={"merged_ranges": merged_ranges[:25]},
        )
    if ws.sheet_state == "veryHidden" and record_findings:
        _append_finding(
            findings,
            appendix,
            finding_type="very_hidden_sheet",
            severity="HIGH",
            sheet=sheet_name,
            message="Arkusz ma stan very-hidden i nie jest widoczny w zwykłym interfejsie Excela.",
            comment="Arkusz wymaga weryfikacji w źródle, bo został ukryty głębiej niż standardowe hidden.",
        )
    if protection_enabled and record_findings:
        _append_finding(
            findings,
            appendix,
            finding_type="sheet_protection",
            severity="LOW",
            sheet=sheet_name,
            message="Arkusz jest objęty ochroną.",
            comment="Sprawdź, czy ochrona nie blokuje rewizji kluczowych komórek.",
        )
    if filters_active and record_findings:
        _append_finding(
            findings,
            appendix,
            finding_type="active_filter",
            severity="LOW",
            sheet=sheet_name,
            message=f"Na arkuszu aktywny jest filtr w zakresie {ws.auto_filter.ref}.",
            comment="Widok danych może nie pokazywać wszystkich rekordów jednocześnie.",
        )

    return {
        "sheet_state": ws.sheet_state,
        "hidden_rows": hidden_rows,
        "hidden_columns": hidden_cols,
        "merged_ranges": merged_ranges,
        "merged_range_count": len(merged_ranges),
        "filters_active": filters_active,
        "filter_range": ws.auto_filter.ref if filters_active else None,
        "protection_enabled": protection_enabled,
    }


def _parse_excel_coord(cell_ref: str) -> tuple[int | None, int | None]:
    from openpyxl.utils.cell import column_index_from_string, coordinate_from_string

    coord = (cell_ref or "").strip()
    if not coord:
        return None, None
    coord = coord.split("!")[-1].strip()
    try:
        col, row = coordinate_from_string(coord)
        return int(row), int(column_index_from_string(col))
    except Exception:
        return None, None


def _expand_window(values: set[int], minimum: int, maximum: int, radius: int) -> set[int]:
    expanded: set[int] = set()
    for value in values:
        if value is None:
            continue
        for current in range(max(minimum, value - radius), min(maximum, value + radius) + 1):
            expanded.add(current)
    return expanded


def _top_formula_columns(formulas: list[dict], limit: int = 6) -> list[int]:
    counts = Counter()
    for formula in formulas:
        _, col_idx = _parse_excel_coord(formula.get("cell", ""))
        if col_idx is not None:
            counts[col_idx] += 1
    return [col_idx for col_idx, _ in counts.most_common(limit)]


def _build_targeted_scope(ws, formulas: list[dict], screening_findings: list[dict]) -> dict:
    from openpyxl.utils import get_column_letter

    suspicious_rows: set[int] = set()
    suspicious_cols: set[int] = set()
    escalation_types: list[str] = []
    escalation_severities: set[str] = set()

    for finding in screening_findings:
        severity = (finding.get("severity") or "").upper()
        if severity not in {"MEDIUM", "HIGH"}:
            continue
        escalation_severities.add(severity)
        finding_type = finding.get("type", "signal")
        if finding_type not in escalation_types:
            escalation_types.append(finding_type)

        raw_cell = finding.get("cell")
        if raw_cell:
            for token in str(raw_cell).split(","):
                row_idx, col_idx = _parse_excel_coord(token.strip())
                if row_idx is not None:
                    suspicious_rows.add(row_idx)
                if col_idx is not None:
                    suspicious_cols.add(col_idx)

        for row_idx in finding.get("rows", []) or []:
            if isinstance(row_idx, int):
                suspicious_rows.add(row_idx)
        for row_idx in finding.get("missing_numbers", []) or []:
            if isinstance(row_idx, int) and 2 <= row_idx <= ws.max_row:
                suspicious_rows.add(row_idx)

        column_value = finding.get("column")
        if isinstance(column_value, str) and column_value:
            _, col_idx = _parse_excel_coord(f"{column_value}1")
            if col_idx is not None:
                suspicious_cols.add(col_idx)
        for column_value in finding.get("columns", []) or []:
            _, col_idx = _parse_excel_coord(f"{column_value}1")
            if col_idx is not None:
                suspicious_cols.add(col_idx)

    should_deep_scan = bool(escalation_types)
    if not should_deep_scan:
        return {
            "should_deep_scan": False,
            "selected_rows": set(),
            "selected_cols": set(),
            "focus_rows": [],
            "focus_columns": [],
            "reasons": [],
        }

    if not suspicious_cols:
        suspicious_cols.update(_top_formula_columns(formulas, limit=6))
    if not suspicious_rows and suspicious_cols:
        suspicious_rows.update(range(2, min(ws.max_row, 40) + 1))

    expanded_rows = _expand_window(suspicious_rows, 2, max(ws.max_row, 2), radius=4)
    expanded_cols = _expand_window(suspicious_cols, 1, max(ws.max_column, 1), radius=1)

    if expanded_rows and not expanded_cols and ws.max_column > 12:
        expanded_cols.update(_top_formula_columns(formulas, limit=6))
    if not expanded_rows and ws.max_column <= 12:
        expanded_rows = set(range(2, ws.max_row + 1))

    return {
        "should_deep_scan": True,
        "selected_rows": expanded_rows,
        "selected_cols": expanded_cols,
        "focus_rows": sorted(expanded_rows)[:30],
        "focus_columns": [get_column_letter(col_idx) for col_idx in sorted(expanded_cols)[:12]],
        "reasons": escalation_types,
        "highest_severity": "HIGH" if "HIGH" in escalation_severities else "MEDIUM",
    }


def _sheet_column_profiles(
    ws, values_ws, selected_rows: set[int] | None = None, selected_cols: set[int] | None = None
):
    profiles = []
    narrowed_rows = sorted(row for row in (selected_rows or set()) if 2 <= row <= ws.max_row)
    full_row_range = range(2, ws.max_row + 1)
    for col_idx in range(1, ws.max_column + 1):
        if selected_cols is not None and selected_rows is None and col_idx not in selected_cols:
            continue
        header = _value_to_text(ws.cell(1, col_idx).value) or f"Kolumna {col_idx}"
        entries = []
        if selected_rows is None:
            row_iter = full_row_range
        elif selected_cols is not None and col_idx in selected_cols:
            row_iter = full_row_range
        else:
            row_iter = narrowed_rows
            if not row_iter:
                continue

        for row_idx in row_iter:
            cell = ws.cell(row_idx, col_idx)
            raw_value = cell.value
            is_formula = isinstance(raw_value, str) and raw_value.startswith("=")
            display_value = _get_formula_display_value(values_ws, row_idx, col_idx) if is_formula else raw_value
            if raw_value in (None, "") and display_value in (None, ""):
                continue
            entries.append(
                {
                    "row": row_idx,
                    "coord": cell.coordinate,
                    "formula": raw_value if is_formula else "",
                    "is_formula": is_formula,
                    "display_value": display_value,
                    "raw_value": raw_value,
                }
            )
        if (
            not entries
            and selected_rows is not None
            and selected_cols is not None
            and col_idx not in selected_cols
        ):
            continue
        profiles.append(
            {
                "col_idx": col_idx,
                "header": header,
                "kind": _header_kind(header),
                "entries": entries,
            }
        )
    return profiles


def _iter_formula_entries(ws, values_ws, *, selected_rows: set[int] | None = None, selected_cols: set[int] | None = None):
    for cell in ws._cells.values():
        raw_value = cell.value
        if not (isinstance(raw_value, str) and raw_value.startswith("=")):
            continue
        row_idx = cell.row
        col_idx = cell.column
        if selected_rows is not None and row_idx not in selected_rows:
            continue
        if selected_cols is not None and col_idx not in selected_cols:
            continue
        yield {
            "coord": cell.coordinate,
            "row": row_idx,
            "col_idx": col_idx,
            "formula": raw_value,
            "display_value": _get_formula_display_value(values_ws, row_idx, col_idx),
        }


def _resolve_sum_like_references(args_text: str, values_ws, current_sheet: str):
    from openpyxl.utils.cell import get_column_letter, range_boundaries

    coords: list[str] = []
    total = 0.0
    had_numeric = False

    for token in _split_excel_args(args_text):
        cleaned = token.strip()
        if not cleaned or "(" in cleaned:
            return None

        if "!" in cleaned:
            sheet_part, ref_part = cleaned.split("!", 1)
            ref_sheet = sheet_part.strip("'")
            if ref_sheet != current_sheet:
                return None
            cleaned = ref_part

        if ":" in cleaned:
            try:
                min_col, min_row, max_col, max_row = range_boundaries(cleaned)
            except Exception:
                return None
            for row_idx in range(min_row, max_row + 1):
                for col_idx in range(min_col, max_col + 1):
                    value = values_ws.cell(row_idx, col_idx).value
                    number = _coerce_number(value)
                    if number is None:
                        continue
                    had_numeric = True
                    total += number
                    coords.append(f"{get_column_letter(col_idx)}{row_idx}")
            continue

        row_idx, col_idx = _parse_excel_coord(cleaned)
        if row_idx is None or col_idx is None:
            return None
        value = values_ws.cell(row_idx, col_idx).value
        number = _coerce_number(value)
        if number is None:
            continue
        had_numeric = True
        total += number
        coords.append(cleaned.replace("$", ""))

    if not had_numeric:
        return None
    return total, coords


def _detect_control_total_and_hidden_reference_signals(
    sheet_name: str,
    ws,
    values_ws,
    findings: list,
    appendix: list,
    *,
    selected_rows: set[int] | None = None,
    selected_cols: set[int] | None = None,
):
    workbook = ws.parent
    summary = {
        "control_total_mismatches": [],
        "cross_sheet_hidden_references": [],
    }
    hidden_ref_seen: set[tuple[str, str]] = set()

    for entry in _iter_formula_entries(
        ws,
        values_ws,
        selected_rows=selected_rows,
        selected_cols=selected_cols,
    ):
        formula = entry["formula"]
        coord = entry["coord"]
        display_value = entry["display_value"]

        for dep in FormulaParser.parse_cell_refs(formula, default_sheet=sheet_name):
            dep_sheet = dep.sheet or sheet_name
            if dep_sheet == sheet_name or dep_sheet not in workbook.sheetnames:
                continue
            dep_state = getattr(workbook[dep_sheet], "sheet_state", "visible") or "visible"
            if dep_state == "visible":
                continue
            fingerprint = (coord, dep_sheet)
            if fingerprint in hidden_ref_seen:
                continue
            hidden_ref_seen.add(fingerprint)
            _append_finding(
                findings,
                appendix,
                finding_type="cross_sheet_hidden_reference",
                severity="MEDIUM",
                sheet=sheet_name,
                cell=coord,
                value=display_value,
                formula=formula,
                message=f"Formuła w {coord} odwołuje się do ukrytego arkusza „{dep_sheet}”.",
                comment=f"Źródłowy arkusz ma stan {dep_state} i nie jest jawny w zwykłym widoku.",
                details={"source_sheet": dep_sheet, "source_sheet_state": dep_state},
            )
            summary["cross_sheet_hidden_references"].append(
                {"cell": coord, "source_sheet": dep_sheet, "source_sheet_state": dep_state}
            )

        sum_match = re.match(r"^=SUM\((.*)\)$", formula.strip(), re.IGNORECASE)
        if not sum_match:
            continue
        stored_total = _coerce_number(display_value)
        if stored_total is None:
            continue

        resolved = _resolve_sum_like_references(sum_match.group(1), values_ws, sheet_name)
        if not resolved:
            continue
        calculated_total, source_cells = resolved
        diff = abs(stored_total - calculated_total)
        base = max(abs(calculated_total), 1.0)
        if diff <= 0.02 or diff / base <= 0.001:
            continue

        _append_finding(
            findings,
            appendix,
            finding_type="control_total_mismatch",
            severity="HIGH",
            sheet=sheet_name,
            cell=coord,
            value=display_value,
            formula=formula,
            message=f"Suma kontrolna w {coord} nie zgadza się z przeliczeniem zakresu źródłowego.",
            comment=(
                f"Formuła zwraca {_fmt_decimal(stored_total)}, a z danych źródłowych wychodzi "
                f"{_fmt_decimal(calculated_total)}."
            ),
            details={
                "source_cells": source_cells[:30],
                "calculated_total": round(calculated_total, 6),
                "difference": round(diff, 6),
            },
        )
        summary["control_total_mismatches"].append(
            {
                "cell": coord,
                "source_cells": source_cells[:30],
                "calculated_total": round(calculated_total, 6),
                "stored_total": round(stored_total, 6),
            }
        )

    return summary


def _fmt_decimal(value: float) -> str:
    text = f"{value:.6f}".rstrip("0").rstrip(".")
    return text or "0"


def _detect_hardcoded_values_and_pattern_deviations(
    sheet_name: str,
    ws,
    values_ws,
    findings: list,
    appendix: list,
    *,
    selected_rows: set[int] | None = None,
    selected_cols: set[int] | None = None,
):
    from openpyxl.utils import get_column_letter

    column_profiles = _sheet_column_profiles(
        ws, values_ws, selected_rows=selected_rows, selected_cols=selected_cols
    )
    formula_counts_by_row = Counter(
        entry["row"]
        for profile in column_profiles
        for entry in profile["entries"]
        if entry["is_formula"]
    )
    summary = {
        "hardcoded_values_in_formula_regions": [],
        "formula_pattern_deviations": [],
        "hardcoded_output_like_values": [],
    }

    for profile in column_profiles:
        entries = profile["entries"]
        formula_entries = [entry for entry in entries if entry["is_formula"]]
        if len(formula_entries) < 2:
            continue

        formula_rows = sorted(entry["row"] for entry in formula_entries)
        min_row = formula_rows[0]
        max_row = formula_rows[-1]
        region_entries = [entry for entry in entries if min_row <= entry["row"] <= max_row]
        formula_ratio = len(formula_entries) / max(1, len(region_entries))

        patterns = Counter(
            _normalize_formula_pattern(entry["formula"], entry["row"], profile["col_idx"])
            for entry in formula_entries
        )
        dominant_pattern, dominant_count = patterns.most_common(1)[0]

        if len(formula_entries) >= 3 and dominant_count >= 2:
            for entry in formula_entries:
                normalized = _normalize_formula_pattern(entry["formula"], entry["row"], profile["col_idx"])
                if normalized != dominant_pattern:
                    msg = (
                        f"Formuła w {entry['coord']} odbiega od dominującego wzorca kolumny "
                        f"„{profile['header']}”."
                    )
                    _append_finding(
                        findings,
                        appendix,
                        finding_type="formula_pattern_deviation",
                        severity="MEDIUM",
                        sheet=sheet_name,
                        cell=entry["coord"],
                        value=entry["display_value"],
                        formula=entry["formula"],
                        message=msg,
                        comment="Sprawdź odwołania względem sąsiednich wierszy; to może być ręczna zmiana logiki.",
                        details={"header": profile["header"]},
                    )
                    summary["formula_pattern_deviations"].append(
                        {
                            "cell": entry["coord"],
                            "header": profile["header"],
                            "formula": entry["formula"],
                            "dominant_pattern": dominant_pattern,
                        }
                    )

        if formula_ratio >= 0.6:
            for entry in region_entries:
                if entry["is_formula"]:
                    continue
                if entry["display_value"] in (None, ""):
                    continue
                has_formula_above = any(row < entry["row"] for row in formula_rows)
                has_formula_below = any(row > entry["row"] for row in formula_rows)
                if not (has_formula_above and has_formula_below):
                    continue
                msg = (
                    f"W obszarze formuł kolumny „{profile['header']}” wykryto twardą wartość w {entry['coord']}."
                )
                _append_finding(
                    findings,
                    appendix,
                    finding_type="hardcoded_in_formula_region",
                    severity="HIGH",
                    sheet=sheet_name,
                    cell=entry["coord"],
                    value=entry["display_value"],
                    formula="",
                    message=msg,
                    comment="Większość komórek w tej części kolumny zawiera formuły, więc wpis ręczny wymaga wyjaśnienia.",
                    details={"header": profile["header"], "column": get_column_letter(profile["col_idx"])},
                )
                summary["hardcoded_values_in_formula_regions"].append(
                    {
                        "cell": entry["coord"],
                        "header": profile["header"],
                        "value": _serialize_excel_value(entry["display_value"]),
                    }
                )

        for entry in entries:
            if entry["is_formula"] or entry["display_value"] in (None, ""):
                continue
            if entry["row"] <= max_row or entry["row"] > max_row + 2:
                continue
            if formula_counts_by_row.get(entry["row"], 0) == 0:
                continue
            msg = (
                f"Komórka {entry['coord']} wygląda jak wynik końcowy, ale zawiera wpis ręczny "
                f"zamiast formuły w kolumnie „{profile['header']}”."
            )
            _append_finding(
                findings,
                appendix,
                finding_type="hardcoded_output_like_value",
                severity="HIGH",
                sheet=sheet_name,
                cell=entry["coord"],
                value=entry["display_value"],
                formula="",
                message=msg,
                comment="W tym obszarze wynik powinien być liczony automatycznie, a nie dopisany ręcznie.",
                details={"header": profile["header"], "row": entry["row"]},
            )
            summary["hardcoded_output_like_values"].append(
                {
                    "cell": entry["coord"],
                    "header": profile["header"],
                    "value": _serialize_excel_value(entry["display_value"]),
                }
            )

    return summary, column_profiles


def _detect_duplicate_and_numeric_signals(sheet_name: str, column_profiles, findings: list, appendix: list):
    thresholds = [1000, 5000, 10000, 20000, 50000, 100000]
    summary = {
        "duplicates": [],
        "numbering_gaps": [],
        "near_thresholds": [],
        "round_amounts": [],
        "amount_outliers": [],
        "benford_deviations": [],
        "weekend_activity": [],
        "night_activity": [],
    }

    for profile in column_profiles:
        entries = profile["entries"]
        kind = profile["kind"]
        header = profile["header"]

        if kind in {"document", "party", "amount"}:
            grouped = defaultdict(list)
            for entry in entries:
                if entry["is_formula"]:
                    continue
                norm = _value_to_text(entry["display_value"]).casefold()
                if norm:
                    grouped[norm].append(entry)
            for norm_value, rows in grouped.items():
                if len(rows) < 2:
                    continue
                finding_type = {
                    "document": "duplicate_document",
                    "party": "duplicate_party",
                    "amount": "duplicate_amount",
                }[kind]
                severity = "HIGH" if kind == "document" else ("MEDIUM" if kind == "amount" else "LOW")
                msg = (
                    f"W kolumnie „{header}” wykryto powtarzającą się wartość: {_value_to_text(rows[0]['display_value'])}."
                )
                row_list = [row["row"] for row in rows]
                _append_finding(
                    findings,
                    appendix,
                    finding_type=finding_type,
                    severity=severity,
                    sheet=sheet_name,
                    cell=", ".join(row["coord"] for row in rows[:6]),
                    value=rows[0]["display_value"],
                    message=msg,
                    comment=f"Wystąpienia w wierszach: {', '.join(map(str, row_list[:12]))}",
                    details={"header": header, "rows": row_list[:25]},
                )
                summary["duplicates"].append(
                    {
                        "header": header,
                        "value": _serialize_excel_value(rows[0]["display_value"]),
                        "rows": row_list[:25],
                        "kind": kind,
                    }
                )

        if kind == "document":
            seq_values = []
            for entry in entries:
                if entry["is_formula"]:
                    continue
                seq = _extract_sequence_number(entry["display_value"])
                if seq is not None:
                    seq_values.append((seq, entry))
            unique_numbers = sorted({seq for seq, _ in seq_values})
            if len(unique_numbers) >= 4 and unique_numbers[-1] - unique_numbers[0] <= len(unique_numbers) + 20:
                missing = [n for n in range(unique_numbers[0], unique_numbers[-1] + 1) if n not in set(unique_numbers)]
                if missing:
                    _append_finding(
                        findings,
                        appendix,
                        finding_type="numbering_gap",
                        severity="MEDIUM",
                        sheet=sheet_name,
                        message=f"W kolumnie „{header}” wykryto luki w numeracji dokumentów.",
                        comment=f"Brakujące numery: {', '.join(map(str, missing[:15]))}",
                        details={"header": header, "missing_numbers": missing[:25]},
                    )
                    summary["numbering_gaps"].append({"header": header, "missing_numbers": missing[:25]})

        if kind == "amount":
            numeric_entries = []
            for entry in entries:
                if entry["is_formula"]:
                    continue
                amount = _coerce_number(entry["display_value"])
                if amount is None or amount <= 0:
                    continue
                numeric_entries.append((amount, entry))

                if any(abs(amount % threshold) < 0.005 for threshold in (10000, 50000, 100000)):
                    _append_finding(
                        findings,
                        appendix,
                        finding_type="round_amount",
                        severity="LOW",
                        sheet=sheet_name,
                        cell=entry["coord"],
                        value=entry["display_value"],
                        message=f"Kwota w {entry['coord']} jest idealnie zaokrąglona.",
                        comment="Okrągła kwota wymaga potwierdzenia, czy wynika z umowy, limitu budżetowego albo ręcznej korekty.",
                        details={"header": header, "thresholds": [10000, 50000, 100000]},
                    )
                    summary["round_amounts"].append(
                        {"cell": entry["coord"], "header": header, "value": _serialize_excel_value(entry["display_value"])}
                    )

                for threshold in thresholds:
                    diff = threshold - amount
                    window = max(50, threshold * 0.01)
                    if 0 < diff <= window:
                        _append_finding(
                            findings,
                            appendix,
                            finding_type="near_threshold",
                            severity="MEDIUM",
                            sheet=sheet_name,
                            cell=entry["coord"],
                            value=entry["display_value"],
                            message=f"Kwota w {entry['coord']} znajduje się tuż pod progiem {threshold:,.0f}.",
                            comment="Sprawdź, czy nie jest to dzielenie lub ustawienie kwoty pod limit akceptacji.",
                            details={"header": header, "threshold": threshold},
                        )
                        summary["near_thresholds"].append(
                            {
                                "cell": entry["coord"],
                                "header": header,
                                "value": _serialize_excel_value(entry["display_value"]),
                                "threshold": threshold,
                            }
                        )
                        break

            if len(numeric_entries) >= 6:
                values = [amount for amount, _ in numeric_entries]
                mean = sum(values) / len(values)
                variance = sum((value - mean) ** 2 for value in values) / len(values)
                stddev = math.sqrt(variance)
                if stddev > 0:
                    for amount, entry in numeric_entries:
                        z_score = abs((amount - mean) / stddev)
                        if z_score < 3:
                            continue
                        _append_finding(
                            findings,
                            appendix,
                            finding_type="amount_outlier",
                            severity="MEDIUM" if z_score < 4.5 else "HIGH",
                            sheet=sheet_name,
                            cell=entry["coord"],
                            value=entry["display_value"],
                            message=f"Kwota w {entry['coord']} wyraźnie odstaje od reszty kolumny „{header}”.",
                            comment=f"Średnia kolumny to {_fmt_decimal(mean)}, odchylenie standardowe {_fmt_decimal(stddev)}, z-score={z_score:.2f}.",
                            details={"header": header, "z_score": round(z_score, 2), "mean": round(mean, 6), "stddev": round(stddev, 6)},
                        )
                        summary["amount_outliers"].append(
                            {
                                "cell": entry["coord"],
                                "header": header,
                                "value": _serialize_excel_value(entry["display_value"]),
                                "z_score": round(z_score, 2),
                            }
                        )

                leading_digits = [digit for digit in (_leading_digit(amount) for amount, _ in numeric_entries) if digit]
                if len(leading_digits) >= 25:
                    expected = _benford_expected()
                    observed_counts = Counter(leading_digits)
                    sample_size = len(leading_digits)
                    observed = {digit: observed_counts.get(digit, 0) / sample_size for digit in range(1, 10)}
                    mad = sum(abs(observed[digit] - expected[digit]) for digit in range(1, 10)) / 9
                    if mad >= 0.03:
                        top_digit = max(
                            range(1, 10),
                            key=lambda digit: abs(observed[digit] - expected[digit]),
                        )
                        severity = "HIGH" if mad >= 0.05 else "MEDIUM"
                        _append_finding(
                            findings,
                            appendix,
                            finding_type="benford_deviation",
                            severity=severity,
                            sheet=sheet_name,
                            message=f"Kolumna „{header}” odbiega od oczekiwanego rozkładu cyfr wiodących Benforda.",
                            comment=(
                                f"Największe odchylenie dotyczy cyfry {top_digit}: obserwacja "
                                f"{observed[top_digit] * 100:.1f}% vs oczekiwane {expected[top_digit] * 100:.1f}% (MAD={mad:.3f})."
                            ),
                            details={
                                "header": header,
                                "sample_size": sample_size,
                                "mad": round(mad, 4),
                                "dominant_digit": top_digit,
                            },
                        )
                        summary["benford_deviations"].append(
                            {
                                "header": header,
                                "sample_size": sample_size,
                                "mad": round(mad, 4),
                                "dominant_digit": top_digit,
                            }
                        )

        if kind == "datetime":
            for entry in entries:
                if entry["is_formula"]:
                    continue
                dt_value = _coerce_datetime(entry["display_value"])
                if dt_value is None:
                    continue
                if dt_value.weekday() >= 5:
                    _append_finding(
                        findings,
                        appendix,
                        finding_type="weekend_activity",
                        severity="LOW",
                        sheet=sheet_name,
                        cell=entry["coord"],
                        value=entry["display_value"],
                        message=f"Data w {entry['coord']} przypada na weekend.",
                        comment="Sprawdź, czy operacja weekendowa jest zgodna z procesem biznesowym.",
                        details={"header": header},
                    )
                    summary["weekend_activity"].append(
                        {"cell": entry["coord"], "header": header, "value": _serialize_excel_value(entry["display_value"])}
                    )
                if getattr(dt_value, "hour", 0) < 6:
                    _append_finding(
                        findings,
                        appendix,
                        finding_type="night_activity",
                        severity="LOW",
                        sheet=sheet_name,
                        cell=entry["coord"],
                        value=entry["display_value"],
                        message=f"Znacznik czasu w {entry['coord']} wskazuje aktywność nocną.",
                        comment="Sprawdź, czy wpis nie został wygenerowany poza standardowym obiegiem pracy.",
                        details={"header": header},
                    )
                    summary["night_activity"].append(
                        {"cell": entry["coord"], "header": header, "value": _serialize_excel_value(entry["display_value"])}
                    )

    return summary


def _analyze_workbook_forensics(
    sheet_name: str,
    ws,
    values_ws,
    formula_cells: list,
    *,
    include_structure: bool = True,
    selected_rows: set[int] | None = None,
    selected_cols: set[int] | None = None,
):
    findings: list = []
    appendix: list = []

    structure = _sheet_structure_signals(
        sheet_name, ws, findings, appendix, record_findings=include_structure
    )
    formula_signals, column_profiles = svc_detect_hardcoded_values_and_pattern_deviations(
        sheet_name,
        ws,
        values_ws,
        findings,
        appendix,
        selected_rows=selected_rows,
        selected_cols=selected_cols,
    )
    numeric_signals = svc_detect_duplicate_and_numeric_signals(
        sheet_name, column_profiles, findings, appendix
    )
    control_signals = svc_detect_control_total_and_hidden_reference_signals(
        sheet_name,
        ws,
        values_ws,
        findings,
        appendix,
        selected_rows=selected_rows,
        selected_cols=selected_cols,
    )

    return {
        "findings": findings,
        "technical_appendix": appendix,
        "structure": structure,
        "formula_signals": formula_signals,
        "data_signals": numeric_signals,
        "control_signals": control_signals,
        "formula_cells_count": len(formula_cells),
    }


def _get_formula_display_value(values_ws, row_idx: int, col_idx: int):
    if values_ws is None:
        return None
    return values_ws.cell(row_idx, col_idx).value


def _column_index(col_name: str) -> int:
    from openpyxl.utils.cell import column_index_from_string

    return column_index_from_string(col_name)


def _build_file_metadata(file_name: str, file_path: Path, workbook, upload_meta: dict) -> dict:
    stat = file_path.stat()
    props = getattr(workbook, "properties", None)
    hidden_sheets = [name for name in workbook.sheetnames if workbook[name].sheet_state != "visible"]
    app_meta = _load_excel_app_metadata(file_path)
    client_modified_at = None
    raw_client_modified = (upload_meta.get("client_last_modified") or "").strip()
    if raw_client_modified:
        try:
            client_modified_at = datetime.fromtimestamp(int(raw_client_modified) / 1000).isoformat()
        except (TypeError, ValueError, OSError):
            client_modified_at = None
    return {
        "name": file_name,
        "extension": file_path.suffix.lower(),
        "mime_type": (upload_meta.get("client_type") or "").strip() or None,
        "size_bytes": int(upload_meta.get("client_size") or stat.st_size),
        "size_human": _format_size_human(int(upload_meta.get("client_size") or stat.st_size)),
        "client_modified_at": client_modified_at,
        "received_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "uploaded_at": datetime.now().isoformat(timespec="seconds"),
        "workbook": {
            "sheet_count": len(workbook.sheetnames),
            "sheet_names": workbook.sheetnames,
            "hidden_sheet_count": len(hidden_sheets),
            "hidden_sheets": hidden_sheets,
            "creator": getattr(props, "creator", None) if props else None,
            "last_modified_by": getattr(props, "lastModifiedBy", None) if props else None,
            "created": _format_dt(getattr(props, "created", None)) if props else None,
            "modified": _format_dt(getattr(props, "modified", None)) if props else None,
            "title": getattr(props, "title", None) if props else None,
            "description": getattr(props, "description", None) if props else None,
            "company": app_meta.get("company") or (getattr(props, "company", None) if props else None),
            "application": app_meta.get("application"),
            "app_version": app_meta.get("app_version"),
            "manager": app_meta.get("manager"),
        },
    }


def _extract_formulas_from_excel(file_path: Path) -> dict:
    """Ekstraktuj formuły z pliku Excel."""
    try:
        import openpyxl
    except ImportError:
        return {"error": "openpyxl nie zainstalowany"}

    formulas_by_sheet = {}

    try:
        wb = openpyxl.load_workbook(file_path, data_only=False)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_formulas = []

            for cell in ws._cells.values():
                if cell.value and isinstance(cell.value, str) and cell.value.startswith("="):
                    sheet_formulas.append(
                        {
                            "cell": f"{cell.column_letter}{cell.row}",
                            "formula": cell.value,
                            "value": None,
                            "is_hidden": bool(
                                ws.row_dimensions[cell.row].hidden
                                or ws.column_dimensions[cell.column_letter].hidden
                            ),
                        }
                    )

            formulas_by_sheet[sheet_name] = {
                "formulas": sheet_formulas,
                "sheet_state": ws.sheet_state,
                "max_row": ws.max_row,
                "max_column": ws.max_column,
            }

        return {
            "success": True,
            "sheets": formulas_by_sheet,
            "workbook": wb,
        }

    except Exception as e:
        logger.error(f"Błąd parsowania Excel: {e}")
        return {"error": str(e)[:200]}


def _load_values_workbook(file_path: Path):
    import openpyxl

    return openpyxl.load_workbook(file_path, data_only=True)


@financial_bp.route("/financial", methods=["POST"])
def audit_financial():
    """
    POST /api/audit/financial
    Upload dokumentu i analiza formuł arkusza.
    """

    if "file" not in request.files:
        return json_error("Brak pliku", status=400)

    file = request.files["file"]
    if not file.filename:
        return json_error("Nazwa pliku pusta", status=400)

    analysis_type = _normalize_analysis_type(request.form.get("analysis_type", "targeted"))
    upload_meta = {
        "client_last_modified": request.form.get("client_last_modified", ""),
        "client_size": request.form.get("client_size", ""),
        "client_type": request.form.get("client_type", ""),
    }

    with tempfile.NamedTemporaryFile(suffix=Path(file.filename).suffix, delete=False) as tmp:
        file.save(tmp.name)
        tmp_path = Path(tmp.name)

    try:
        fingerprint, cache_entry, cache_hit = _get_financial_cache_entry(tmp_path)
        if "error" in cache_entry:
            return json_error(cache_entry["error"], status=400)
        formulas_result = cache_entry["formulas_result"]

        result = {
            "file_name": file.filename,
            "file_metadata": _build_file_metadata(
                file.filename,
                tmp_path,
                formulas_result["workbook"],
                upload_meta,
            ),
            "sheets": {},
            "summary": {
                "total_sheets": len(formulas_result.get("sheets", {})),
                "total_formulas": 0,
                "risk_level": "LOW",
                "anomalies_count": 0,
                "anomalies": [],
                "severity_counts": {"LOW": 0, "MEDIUM": 0, "HIGH": 0},
            },
            "analysis": {
                "requested_mode": analysis_type,
                "strategy": "screen_then_trace" if analysis_type in {"quick", "targeted"} else "full_pass",
                "deep_scan_sheets": 0,
                "screening_only_sheets": 0,
                "cache_hit": cache_hit,
                "screening_reused": cache_hit,
                "deep_scan_reused_sheets": 0,
                "cache_key": fingerprint[:12],
            },
            "logic_overview": {
                "sheet_count": len(formulas_result.get("sheets", {})),
                "formula_sheets": 0,
                "cross_sheet_references": 0,
                "max_dependency_depth": 0,
                "hub_cells": [],
                "description": "",
            },
            "technical_appendix": [],
            "ai_evidence_pack": None,
            "forensic_overview": {
                "analysis_mode": analysis_type,
                "deep_scan_sheets": 0,
                "screening_only_sheets": 0,
                "hidden_rows_total": 0,
                "hidden_columns_total": 0,
                "very_hidden_sheets": [],
                "protected_sheets": [],
                "filtered_sheets": [],
                "hardcoded_formula_region_hits": 0,
                "formula_pattern_deviations": 0,
                "duplicates": 0,
                "amount_outlier_hits": 0,
                "benford_deviation_hits": 0,
                "control_total_mismatches": 0,
                "cross_sheet_hidden_reference_hits": 0,
            },
        }

        all_anomalies = []
        logic_descriptions = []
        all_hubs = []
        ai_sheet_packs: list[dict] = []
        values_wb = None

        for sheet_name, base_sheet in cache_entry["base_sheets"].items():
            sheet_info = base_sheet["sheet_info"]
            formulas = base_sheet["formulas"]
            analyzer = base_sheet["analyzer"]
            logic_summary = base_sheet["logic_summary"]
            screening_appendix = base_sheet["screening_appendix"]
            screening_raw = base_sheet["screening_raw"]
            scope = base_sheet["scope"]
            forensic, sheet_mode, deep_scan_reused = _get_sheet_forensic_result(
                cache_entry,
                tmp_path,
                sheet_name,
                analysis_type,
            )
            if analysis_type in {"full", "formulas"} and values_wb is None:
                values_wb = _get_cached_values_workbook(cache_entry, tmp_path)
            if deep_scan_reused and sheet_mode != "screening_only":
                result["analysis"]["deep_scan_reused_sheets"] += 1

            enriched_anomalies = [_anomaly_details(a) for a in (screening_raw + forensic["findings"])]
            logic_description = _describe_sheet_logic(sheet_name, logic_summary)
            sheet_risk_level = _risk_level_from_anomalies(enriched_anomalies)

            sheet_data = {
                "formula_count": len(formulas),
                "cell_count": len(analyzer.cells),
                "anomalies": enriched_anomalies,
                "risk_level": sheet_risk_level,
                "sheet_state": sheet_info.get("sheet_state", "visible"),
                "dimensions": {
                    "max_row": sheet_info.get("max_row", 0),
                    "max_column": sheet_info.get("max_column", 0),
                },
                "forensic_signals": {
                    "structure": forensic["structure"],
                    "formula_signals": forensic["formula_signals"],
                    "data_signals": forensic["data_signals"],
                    "control_signals": forensic["control_signals"],
                },
                "logic_summary": {
                    **logic_summary,
                    "description_pl": logic_description,
                },
                "evidence_summary": _build_sheet_evidence_summary(
                    sheet_name,
                    enriched_anomalies,
                    {
                        **logic_summary,
                        "description_pl": logic_description,
                    },
                ),
                "flow_graph": _build_sheet_flow_graph(
                    sheet_name,
                    analyzer,
                    logic_summary,
                    enriched_anomalies,
                ),
                "lineage_graph": _build_sheet_lineage_graph(
                    sheet_name,
                    analyzer,
                    enriched_anomalies,
                    formulas_result["workbook"],
                    values_wb,
                ),
                "analysis": {
                    "mode": sheet_mode,
                    "screening_findings": len(screening_raw),
                    "deep_scan_triggered": sheet_mode != "screening_only",
                    "focus_rows": scope.get("focus_rows", []),
                    "focus_columns": scope.get("focus_columns", []),
                    "reasons": scope.get("reasons", []),
                },
            }
            sheet_ai_pack = _build_sheet_ai_evidence_pack(
                sheet_name,
                sheet_risk_level,
                logic_description,
                logic_summary,
                sheet_data["analysis"],
                enriched_anomalies,
            )
            sheet_data["ai_evidence_pack"] = sheet_ai_pack

            if analysis_type in ["full", "formulas"]:
                values_ws = values_wb[sheet_name]
                sheet_data["formulas"] = [
                    {
                        "cell": str(cell.cell_ref),
                        "formula": cell.formula,
                        "dependencies": [str(d) for d in cell.dependencies],
                        "value": _get_formula_display_value(
                            values_ws,
                            cell.cell_ref.row,
                            _column_index(cell.cell_ref.col),
                        ),
                    }
                    for cell in analyzer.cells.values()
                ]

            result["sheets"][sheet_name] = sheet_data
            result["summary"]["total_formulas"] += len(formulas)
            result["logic_overview"]["formula_sheets"] += 1 if formulas else 0
            result["logic_overview"]["cross_sheet_references"] += logic_summary.get(
                "cross_sheet_references", 0
            )
            result["logic_overview"]["max_dependency_depth"] = max(
                result["logic_overview"]["max_dependency_depth"],
                logic_summary.get("max_dependency_depth", 0),
            )
            logic_descriptions.append(sheet_data["logic_summary"]["description_pl"])
            all_hubs.extend(logic_summary.get("hub_cells", []))
            all_anomalies.extend(enriched_anomalies)
            if sheet_ai_pack:
                ai_sheet_packs.append(sheet_ai_pack)
            if sheet_mode == "screening_only":
                result["analysis"]["screening_only_sheets"] += 1
                result["forensic_overview"]["screening_only_sheets"] += 1
            else:
                result["analysis"]["deep_scan_sheets"] += 1
                result["forensic_overview"]["deep_scan_sheets"] += 1
            result["technical_appendix"].extend([_anomaly_details(item) for item in screening_appendix])
            result["technical_appendix"].extend(
                [_anomaly_details(item) for item in forensic["technical_appendix"]]
            )
            result["forensic_overview"]["hidden_rows_total"] += len(
                forensic["structure"].get("hidden_rows", [])
            )
            result["forensic_overview"]["hidden_columns_total"] += len(
                forensic["structure"].get("hidden_columns", [])
            )
            if forensic["structure"].get("sheet_state") == "veryHidden":
                result["forensic_overview"]["very_hidden_sheets"].append(sheet_name)
            if forensic["structure"].get("protection_enabled"):
                result["forensic_overview"]["protected_sheets"].append(sheet_name)
            if forensic["structure"].get("filters_active"):
                result["forensic_overview"]["filtered_sheets"].append(sheet_name)
            result["forensic_overview"]["hardcoded_formula_region_hits"] += len(
                forensic["formula_signals"].get("hardcoded_values_in_formula_regions", [])
            )
            result["forensic_overview"]["formula_pattern_deviations"] += len(
                forensic["formula_signals"].get("formula_pattern_deviations", [])
            )
            result["forensic_overview"]["duplicates"] += len(
                forensic["data_signals"].get("duplicates", [])
            )
            result["forensic_overview"]["amount_outlier_hits"] += len(
                forensic["data_signals"].get("amount_outliers", [])
            )
            result["forensic_overview"]["benford_deviation_hits"] += len(
                forensic["data_signals"].get("benford_deviations", [])
            )
            result["forensic_overview"]["control_total_mismatches"] += len(
                forensic["control_signals"].get("control_total_mismatches", [])
            )
            result["forensic_overview"]["cross_sheet_hidden_reference_hits"] += len(
                forensic["control_signals"].get("cross_sheet_hidden_references", [])
            )

        high_count = sum(1 for a in all_anomalies if a.get("severity") == "HIGH")
        medium_count = sum(1 for a in all_anomalies if a.get("severity") == "MEDIUM")
        low_count = sum(1 for a in all_anomalies if a.get("severity") == "LOW")

        if high_count > 0:
            result["summary"]["risk_level"] = "CRITICAL"
        elif medium_count >= 3:
            result["summary"]["risk_level"] = "HIGH"
        elif medium_count > 0:
            result["summary"]["risk_level"] = "MEDIUM"
        elif low_count > 0:
            result["summary"]["risk_level"] = "LOW"

        result["summary"]["anomalies_count"] = len(all_anomalies)
        result["summary"]["anomalies"] = all_anomalies
        result["summary"]["severity_counts"] = {
            "LOW": low_count,
            "MEDIUM": medium_count,
            "HIGH": high_count,
        }
        result["logic_overview"]["hub_cells"] = sorted(
            all_hubs,
            key=lambda cell: (cell.get("dependents_count", 0), cell.get("dependencies_count", 0)),
            reverse=True,
        )[:5]
        result["logic_overview"]["description"] = " ".join(logic_descriptions) or (
            "Nie wykryto formuł, więc arkusz pełni raczej rolę danych wejściowych niż modelu obliczeniowego."
        )
        result["ai_evidence_pack"] = _build_ai_evidence_pack(file.filename, ai_sheet_packs)

        return json_success(**result)

    except Exception as e:
        logger.error(f"Błąd audytu: {e}", exc_info=True)
        return json_error(f"Błąd: {str(e)[:200]}", status=500)

    finally:
        try:
            tmp_path.unlink()
        except Exception:
            pass


@financial_bp.route("/financial/opinion", methods=["POST"])
def audit_financial_opinion():
    payload = request.get_json(silent=True) or {}
    evidence_pack = payload.get("evidence_pack")
    if not isinstance(evidence_pack, dict):
        report = payload.get("report") or {}
        if isinstance(report, dict):
            evidence_pack = report.get("ai_evidence_pack")

    if not isinstance(evidence_pack, dict):
        return json_error("Brak ai_evidence_pack do analizy.", status=400, opinion_available=False)

    sheets = evidence_pack.get("sheets")
    if not isinstance(sheets, list) or not any(
        isinstance(sheet, dict) and sheet.get("findings") for sheet in sheets
    ):
        return json_error(
            "Brak silnych anomalii do opinii AI.",
            status=400,
            opinion_available=False,
        )

    system_prompt, user_prompt = _build_financial_forensics_prompt(evidence_pack)
    try:
        llm_result = call_llm(
            user_prompt,
            system=system_prompt,
            stream=False,
            provider=_FINANCIAL_AI_OPINION_PROVIDER,
            model=LLM_MODEL,
            max_tokens=1200,
            temperature=0.1,
        )
        raw_text = _llm_response_text(llm_result)
    except Exception as exc:
        logger.warning("AI forensic opinion failed: %s", exc)
        return json_success(
            opinion_available=False,
            ai_forensic_opinion=None,
            warning=f"Nie udało się wygenerować opinii AI: {str(exc)[:200]}",
        )

    normalized_opinion = _normalize_ai_forensic_opinion(_extract_json_object(raw_text))
    if normalized_opinion is None:
        logger.warning("AI forensic opinion returned invalid JSON: %s", raw_text[:300])
        return json_success(
            opinion_available=False,
            ai_forensic_opinion=None,
            warning="Model zwrócił niepoprawny JSON, więc opinia została odrzucona.",
        )

    return json_success(
        opinion_available=True,
        ai_forensic_opinion=normalized_opinion,
        evidence_summary=evidence_pack.get("summary", {}),
        provider=_FINANCIAL_AI_OPINION_PROVIDER,
    )


@financial_bp.route("/financial/ocr-cross-check", methods=["POST"])
def audit_financial_ocr_cross_check():
    """Weryfikacja krzyżowa sygnałów HIGH/CRITICAL względem dokumentów OCR w Qdrant.

    Sekcja dodatkowa — nie warunkuje wyniku bazowego audytu XLSX. Gdy Qdrant/OCR
    niedostępny, zwraca status "unavailable" per kandydat (fail-closed), a nie błąd 5xx.
    """
    payload = request.get_json(silent=True) or {}
    evidence_pack = payload.get("evidence_pack")
    if not isinstance(evidence_pack, dict):
        report = payload.get("report") or {}
        if isinstance(report, dict):
            evidence_pack = report.get("ai_evidence_pack")

    if not isinstance(evidence_pack, dict):
        return json_error("Brak ai_evidence_pack do weryfikacji OCR.", status=400, cross_check_available=False)

    candidates, truncated = svc_extract_ocr_candidates(evidence_pack)
    if not candidates:
        return json_success(cross_check_available=True, cross_check=[], truncated=False)

    import app as _app

    def _search(query_text: str) -> list[dict]:
        client = _app.get_qdrant_client()
        vector = _app.get_embedding(query_text)
        res = client.query_points(collection_name=_app.ACTIVE_COLLECTION, query=vector, limit=3)
        return [
            {"file": p.payload.get("file", ""), "text": p.payload.get("text", ""), "score": float(p.score)}
            for p in res.points
        ]

    results = svc_run_ocr_cross_check(candidates, _search)

    return json_success(
        cross_check_available=True,
        cross_check=results,
        checked=len(results),
        truncated=truncated,
    )


@financial_bp.route("/financial/export", methods=["POST"])
def export_financial_report_docx():
    payload = request.get_json(silent=True) or {}
    report = payload.get("report") if isinstance(payload.get("report"), dict) else payload
    if not isinstance(report, dict) or not isinstance(report.get("sheets"), dict):
        return json_error("Brak payloadu raportu audytu finansowego.", status=400)

    try:
        buf = _build_financial_docx(report)
    except Exception as exc:
        logger.error("Financial DOCX export failed: %s", exc, exc_info=True)
        return json_error(f"Błąd eksportu DOCX: {str(exc)[:200]}", status=500)

    return send_file(
        buf,
        as_attachment=True,
        download_name=_report_export_filename(report.get("file_name") or (report.get("file_metadata") or {}).get("name") or "audyt_xlsx"),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
