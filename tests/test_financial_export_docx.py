from copy import deepcopy
from io import BytesIO

from docx import Document
from openpyxl import Workbook

from app import app


def _post_financial_audit(workbook: Workbook, filename: str = "audit.xlsx"):
    payload = BytesIO()
    workbook.save(payload)
    payload.seek(0)
    client = app.test_client()
    return client.post(
        "/api/audit/financial",
        data={"file": (payload, filename), "analysis_type": "full"},
        content_type="multipart/form-data",
    )


def _export_docx(report: dict):
    client = app.test_client()
    return client.post("/api/audit/financial/export", json={"report": report})


def _docx_text(blob: bytes) -> str:
    doc = Document(BytesIO(blob))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _report_workbook() -> Workbook:
    wb = Workbook()
    ws = wb.active
    ws.title = "Jawny"
    ws["A1"] = "Opis"
    ws["B1"] = "Wynik"
    ws["A2"] = "rekord kontrolny"
    hidden = wb.create_sheet("UkryteDane")
    hidden.sheet_state = "hidden"
    hidden["A2"] = 123
    ws["B2"] = "=UkryteDane!A2"
    return wb


def test_financial_export_endpoint_returns_docx_file():
    audit_response = _post_financial_audit(_report_workbook(), "raport.xlsx")
    report = audit_response.get_json()

    export_response = _export_docx(report)

    assert export_response.status_code == 200
    assert (
        export_response.headers["Content-Type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment;" in export_response.headers.get("Content-Disposition", "")


def test_financial_export_docx_contains_title_and_high_priority_findings():
    audit_response = _post_financial_audit(_report_workbook(), "raport.xlsx")
    report = audit_response.get_json()

    export_response = _export_docx(report)
    text = _docx_text(export_response.data)

    assert "Raport z audytu finansowego XLSX" in text
    assert "Findings HIGH/CRITICAL" in text
    assert "Odwołanie do ukrytego arkusza" in text
    assert "Formuła w B2 odwołuje się do ukrytego arkusza" in text


def test_financial_export_docx_includes_ai_opinion_section_when_present():
    audit_response = _post_financial_audit(_report_workbook(), "raport.xlsx")
    report = audit_response.get_json()
    report["ai_forensic_opinion"] = {
        "overall_assessment": "Silna poszlaka ukrycia źródła danych wejściowych.",
        "sheet_comment": "Najmocniejszy sygnał dotyczy arkusza Jawny.",
        "findings": [
            {
                "finding_id": "Jawny:Jawny!B2:cross_sheet_hidden_reference",
                "fact": "Komórka B2 pobiera wartość z ukrytego arkusza.",
                "intent": "Możliwe obejście jawnej kontroli wyniku.",
                "expert_comment": "Źródło obliczeń nie jest jawne dla zwykłego odbiorcy.",
                "confidence": "wysoka",
                "next_check": "Otwórz arkusz UkryteDane i sprawdź komórkę A2.",
            }
        ],
    }

    export_response = _export_docx(report)
    text = _docx_text(export_response.data)

    assert "Wniosek śledczy AI" in text
    assert "Silna poszlaka ukrycia źródła danych wejściowych." in text
    assert "Możliwa intencja: Możliwe obejście jawnej kontroli wyniku." in text
    assert "Co sprawdzić dalej: Otwórz arkusz UkryteDane i sprawdź komórkę A2." in text


def test_financial_export_docx_works_without_ai_opinion_and_lineage_graph():
    audit_response = _post_financial_audit(_report_workbook(), "raport.xlsx")
    report = deepcopy(audit_response.get_json())
    report.pop("ai_forensic_opinion", None)
    for sheet in report["sheets"].values():
        sheet.pop("lineage_graph", None)

    export_response = _export_docx(report)
    text = _docx_text(export_response.data)

    assert export_response.status_code == 200
    assert "Graf powiązań" in text
    assert "Brak relacji do pokazania w grafie powiązań." in text
    assert "Zastrzeżenie" in text
